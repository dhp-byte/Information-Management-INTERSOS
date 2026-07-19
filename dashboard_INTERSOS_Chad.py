import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import folium
from folium.plugins import MarkerCluster, HeatMap, MiniMap
from streamlit_folium import st_folium
from datetime import datetime, timedelta
import io, warnings, time, json
warnings.filterwarnings("ignore")

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

# ── ReportLab (PDF) ──────────────────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors as rl_colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable, PageBreak)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

st.set_page_config(
    page_title="SI Tchad · IM Dashboard",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
# SI BRAND  (from intersos.org)
# ══════════════════════════════════════════════════════════════════════════════
SI_RED  = "#C1121F"
SI_RED2 = "#8B0D14"

SI_IMAGES = [
    "https://images.unsplash.com/photo-1488521787991-ed7bbaae773c?w=1400&q=80",
    "https://images.unsplash.com/photo-1531206715517-5c0ba140b2b8?w=1400&q=80",
    "https://images.unsplash.com/photo-1509099836639-18ba1795216d?w=1400&q=80",
    "https://images.unsplash.com/photo-1542810634-71277d95dcbb?w=1400&q=80",
    "https://images.unsplash.com/photo-1469571486292-0ba58a3f068b?w=1400&q=80",
]

DARK = dict(
    bg="#12151C", surface="#1C2130", panel="#222840",
    border="rgba(227,0,27,0.18)", border2="rgba(227,0,27,0.35)",
    text="#F1F5F9", muted="#94A3B8", dim="#2A3050",
    navbg="#0E1120", navborder="rgba(227,0,27,0.3)",
    card="#1C2130", input="#161B2E",
    plotbg="rgba(0,0,0,0)", gridc="rgba(255,255,255,0.05)",
    fontc="#64748B", title="#E2E8F0",
    accent="#C1121F", accent2="#FF3355",
    sbg="#12151C",
)
LIGHT = dict(
    bg="#F8F9FB", surface="#FFFFFF", panel="#FFFFFF",
    border="rgba(227,0,27,0.12)", border2="rgba(227,0,27,0.28)",
    text="#0F172A", muted="#64748B", dim="#F1F5F9",
    navbg="#FFFFFF", navborder="rgba(227,0,27,0.2)",
    card="#FFFFFF", input="#F8FAFC",
    plotbg="rgba(0,0,0,0)", gridc="rgba(0,0,0,0.04)",
    fontc="#94A3B8", title="#1E293B",
    accent="#C1121F", accent2="#8B0D14",
    sbg="#F8F9FB",
)
COLORS  = ["#C1121F","#3B82F6","#F59E0B","#10B981","#8B5CF6","#14B8A6","#EC4899","#F97316"]
SC = {"WASH":"#3B82F6","FSL":"#10B981","Shelter & NFI":"#F59E0B","Cash & Voucher Assistance":"#C1121F"}
CREDENTIALS = {"im_manager": "intersos2026"}

NAV_ITEMS = [
    ("\U0001f3e0", "Overview"),("\U0001f5fa\ufe0f", "Map"),("\U0001f4a7", "WASH"),
    ("\U0001f33e", "FSL"),("\U0001f4b5", "CVA"),("\U0001f4ca", "Indicators"),
    ("\U0001f5c2\ufe0f", "Raw Data"),("\U0001f4c4", "Report"),
]

# ══════════════════════════════════════════════════════════════════════════════
# CSS INJECTOR  —  SI brand identity
# ══════════════════════════════════════════════════════════════════════════════
def inject_css(T):
    is_dark = T.get("bg","") == DARK.get("bg","")
    shadow     = "0 4px 24px rgba(0,0,0,0.4)" if is_dark else "0 4px 24px rgba(0,0,0,0.08)"
    nav_shadow = "0 2px 20px rgba(0,0,0,0.5)" if is_dark else "0 2px 12px rgba(0,0,0,0.1)"
    img_filter = "brightness(0.65) saturate(1.2)" if is_dark else "brightness(0.8) saturate(1.1)"
    login_sh   = "0 40px 80px rgba(0,0,0,0.6)" if is_dark else "0 40px 80px rgba(0,0,0,0.15)"
    ac  = T["accent"]
    ac2 = T["accent2"]
    css = (
        "@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800;900"
        "&family=JetBrains+Mono:wght@400;500&display=swap');\n"
        ":root{"
        "--bg:" + T["bg"] + ";--surface:" + T["surface"] + ";--panel:" + T["panel"] + ";"
        "--border:" + T["border"] + ";--border2:" + T["border2"] + ";"
        "--text:" + T["text"] + ";--muted:" + T["muted"] + ";--dim:" + T["dim"] + ";"
        "--navbg:" + T["navbg"] + ";--navbdr:" + T["navborder"] + ";"
        "--card:" + T["card"] + ";--input:" + T["input"] + ";"
        "--accent:" + ac + ";--accent2:" + ac2 + ";"
        "--si-red:" + SI_RED + ";--si-red2:" + SI_RED2 + ";"
        "--font:'Outfit',sans-serif;--mono:'JetBrains Mono',monospace;"
        "--r:12px;--shadow:" + shadow + ";}"
        "*,*::before,*::after{box-sizing:border-box;}"
        "html,body{font-family:var(--font);background:var(--bg);color:var(--text);margin:0;}"
        "[data-testid='stAppViewContainer'],[data-testid='stMain'],.main .block-container{background:var(--bg)!important;}"
        "[data-testid='stHeader']{background:var(--navbg)!important;border-bottom:3px solid " + SI_RED + "!important;}"
        "[data-testid='block-container']{padding:0!important;max-width:100%!important;}"
        "::-webkit-scrollbar{width:5px;height:5px;}"
        "::-webkit-scrollbar-thumb{background:" + SI_RED + ";border-radius:3px;}"

        # HERO
        ".si-hero{position:relative;overflow:hidden;border-radius:14px;margin-bottom:1.5rem;min-height:200px;"
        "background:linear-gradient(135deg,#1A0004 0%,#2A0008 50%,#1A0004 100%);}"
        ".si-hero img{position:absolute;inset:0;width:100%;height:100%;object-fit:cover;"
        "opacity:0.3;filter:" + img_filter + ";}"
        ".si-hero-content{position:relative;z-index:2;padding:2rem 2.5rem;min-height:200px;"
        "display:flex;flex-direction:column;justify-content:flex-end;}"
        ".si-tag{display:inline-block;background:" + SI_RED + ";color:#fff;font-size:.68rem;"
        "font-weight:800;letter-spacing:.12em;text-transform:uppercase;"
        "padding:3px 10px;border-radius:3px;margin-bottom:.7rem;width:fit-content;}"
        ".si-hero-title{font-size:1.9rem;font-weight:900;color:#fff;margin:0 0 .4rem;"
        "letter-spacing:-.02em;text-shadow:0 2px 20px rgba(0,0,0,0.5);}"
        ".si-hero-sub{font-size:.85rem;color:rgba(255,255,255,.7);margin:0;}"
        ".si-hero-bar{position:absolute;bottom:0;left:0;right:0;height:4px;"
        "background:linear-gradient(90deg," + SI_RED + "," + SI_RED2 + "," + SI_RED + ");"
        "animation:redpulse 3s ease-in-out infinite;}"
        "@keyframes redpulse{0%,100%{opacity:1}50%{opacity:.6}}"

        # IMAGE SLIDER
        ".img-slider{overflow:hidden;border-radius:10px;margin:1rem 0;}"
        ".img-track{display:flex;gap:10px;animation:slide 32s linear infinite;width:max-content;}"
        ".img-track img{height:145px;width:240px;object-fit:cover;border-radius:8px;"
        "flex-shrink:0;filter:" + img_filter + ";transition:filter .3s;"
        "border:1px solid rgba(227,0,27,0.2);}"
        ".img-track img:hover{filter:brightness(1) saturate(1.3);}"
        "@keyframes slide{0%{transform:translateX(0)}100%{transform:translateX(-50%)}}"

        # TOPBAR
        ".si-topbar{background:var(--navbg);border-bottom:3px solid " + SI_RED + ";"
        "padding:.55rem 1.8rem;display:flex;align-items:center;gap:14px;"
        "position:sticky;top:0;z-index:1000;backdrop-filter:blur(16px);"
        "box-shadow:" + nav_shadow + ";}"
        ".si-logo-box{width:36px;height:36px;background:" + SI_RED + ";border-radius:6px;"
        "display:flex;align-items:center;justify-content:center;font-size:1.15rem;"
        "color:#fff;font-weight:900;flex-shrink:0;}"
        ".si-brand{font-size:.95rem;font-weight:900;color:#fff;font-family:var(--font);}"
        ".si-brand small{display:block;font-size:.59rem;color:rgba(255,255,255,.4);"
        "font-weight:400;letter-spacing:.09em;text-transform:uppercase;}"
        ".si-chip{background:rgba(255,255,255,.08);border:1px solid rgba(255,255,255,.12);"
        "border-radius:20px;padding:3px 12px;font-size:.71rem;color:rgba(255,255,255,.55);}"

        # NAV BUTTONS  (active index injected dynamically in top_nav)
        "div[data-testid='stHorizontalBlock']>div>[data-testid='stButton'] button{"
        "background:transparent!important;border:none!important;"
        "border-bottom:3px solid transparent!important;border-radius:0!important;"
        "color:rgba(255,255,255,0.5)!important;font-family:var(--font)!important;"
        "font-size:.82rem!important;font-weight:500!important;"
        "padding:.55rem .35rem!important;width:100%!important;"
        "transition:all .18s!important;white-space:nowrap!important;}"
        "div[data-testid='stHorizontalBlock']>div>[data-testid='stButton'] button:hover{"
        "color:#fff!important;background:rgba(227,0,27,0.1)!important;"
        "border-bottom-color:rgba(227,0,27,0.5)!important;}"
        "div[data-testid='stHorizontalBlock']{gap:2px!important;margin:0!important;padding:0!important;}"

        # KPI CARDS
        ".kpi-card{background:var(--card);border:1px solid var(--border);border-radius:var(--r);"
        "padding:1.1rem 1.2rem .9rem;position:relative;overflow:hidden;"
        "transition:transform .2s,box-shadow .2s;box-shadow:var(--shadow);}"
        ".kpi-card:hover{transform:translateY(-2px);box-shadow:0 10px 36px rgba(227,0,27,0.2);}"
        ".kpi-card::before{content:'';position:absolute;top:0;left:0;right:0;height:3px;}"
        ".kpi-card::after{content:'';position:absolute;bottom:0;left:0;width:0;height:2px;"
        "background:" + SI_RED + ";transition:width .4s;}"
        ".kpi-card:hover::after{width:100%;}"
        ".kpi-red::before{background:linear-gradient(90deg," + SI_RED + "," + SI_RED2 + ");}"
        ".kpi-blue::before{background:linear-gradient(90deg,#3B82F6,#60A5FA);}"
        ".kpi-green::before{background:linear-gradient(90deg,#10B981,#34D399);}"
        ".kpi-amber::before{background:linear-gradient(90deg,#F59E0B,#FCD34D);}"
        ".kpi-purple::before{background:linear-gradient(90deg,#8B5CF6,#A78BFA);}"
        ".kpi-teal::before{background:linear-gradient(90deg,#14B8A6,#5EEAD4);}"
        ".kpi-label{font-size:.63rem;font-weight:700;letter-spacing:.12em;text-transform:uppercase;"
        "color:var(--muted);margin-bottom:.4rem;}"
        ".kpi-value{font-size:1.8rem;font-weight:900;color:var(--text);line-height:1;font-family:var(--font);}"
        ".kpi-sub{font-size:.7rem;color:var(--muted);margin-top:.3rem;}"
        ".kpi-icon{position:absolute;top:.85rem;right:.95rem;font-size:1.3rem;opacity:.12;}"
        ".kpi-delta{font-size:.69rem;font-weight:700;margin-top:.35rem;"
        "display:inline-flex;align-items:center;gap:3px;padding:1px 7px;border-radius:20px;}"
        ".d-up{background:rgba(16,185,129,.15);color:#34D399;}"
        ".d-down{background:rgba(227,0,27,.15);color:#FF6680;}"
        ".d-mid{background:rgba(245,158,11,.15);color:#FCD34D;}"

        # SECTION HEADER
        ".sh{font-size:.66rem;font-weight:800;letter-spacing:.14em;text-transform:uppercase;"
        "color:" + SI_RED + ";margin:1.8rem 0 1rem;display:flex;align-items:center;gap:10px;}"
        ".sh::before{content:'';width:3px;height:14px;background:" + SI_RED + ";border-radius:2px;flex-shrink:0;}"
        ".sh::after{content:'';flex:1;height:1px;background:var(--border);}"

        # PAGE HEADER
        ".ph{border-bottom:1px solid var(--border);margin-bottom:1.5rem;padding-bottom:.9rem;}"
        ".ph h1{font-size:1.55rem;font-weight:900;color:var(--text);margin:0 0 3px;"
        "font-family:var(--font);letter-spacing:-.02em;}"
        ".ph p{font-size:.81rem;color:var(--muted)!important;margin:0;}"

        # BADGES
        ".badge{display:inline-block;padding:2px 9px;border-radius:3px;"
        "font-size:.66rem;font-weight:700;letter-spacing:.04em;}"
        ".bg{background:rgba(16,185,129,.18);color:#34D399;}"
        ".ba{background:rgba(245,158,11,.18);color:#FCD34D;}"
        ".br{background:rgba(227,0,27,.18);color:#FF6680;}"
        ".bn{background:rgba(100,116,139,.18);color:#94A3B8;}"
        ".pbar-wrap{background:rgba(128,128,128,.12);border-radius:3px;height:5px;overflow:hidden;margin:5px 0 3px;}"
        ".pbar{height:100%;border-radius:3px;}"

        # IND CARD
        ".ind-card{background:var(--card);border:1px solid var(--border);border-left:3px solid " + SI_RED + ";"
        "border-radius:var(--r);padding:.9rem 1.1rem;margin-bottom:.5rem;transition:border-color .2s;}"
        ".ind-card:hover{border-left-color:" + SI_RED2 + ";}"
        ".ind-name{font-size:.87rem;font-weight:600;color:var(--text);}"
        ".ind-vals{font-size:.71rem;color:var(--muted);font-family:var(--mono);}"

        # ALERT
        ".alert-banner{background:rgba(227,0,27,.07);border:1px solid rgba(227,0,27,.22);"
        "border-left:3px solid " + SI_RED + ";border-radius:var(--r);padding:.72rem 1rem;margin-bottom:.55rem;}"
        ".alert-banner span{font-size:.82rem;color:#FF9AAA;}"

        # INPUTS
        "[data-testid='stSelectbox']>div>div{background:var(--input)!important;"
        "border:1px solid var(--border2)!important;border-radius:var(--r)!important;color:var(--text)!important;}"
        "[data-testid='stTextInput']>div>div>input{background:var(--input)!important;"
        "border:1px solid var(--border2)!important;color:var(--text)!important;border-radius:var(--r)!important;}"
        "[data-baseweb='select']*{color:var(--text)!important;}"
        "[data-baseweb='menu']{background:var(--input)!important;}"
        "label{color:var(--muted)!important;font-size:.81rem!important;}"
        "p{color:var(--muted)!important;}"

        # BUTTONS
        ".stButton>button{background:" + SI_RED + "!important;color:#fff!important;"
        "border:none!important;border-radius:6px!important;"
        "font-family:var(--font)!important;font-weight:700!important;"
        "font-size:.88rem!important;padding:.55rem 1.4rem!important;"
        "transition:all .2s!important;letter-spacing:.01em!important;}"
        ".stButton>button:hover{background:" + SI_RED2 + "!important;"
        "transform:translateY(-1px)!important;box-shadow:0 4px 16px rgba(227,0,27,.4)!important;}"
        "[data-testid='stDownloadButton']>button{background:var(--card)!important;"
        "border:1px solid " + SI_RED + "!important;color:" + SI_RED + "!important;border-radius:6px!important;}"
        "[data-testid='stDownloadButton']>button:hover{background:" + SI_RED + "!important;color:#fff!important;}"
        "[data-testid='stFileUploader']{background:var(--card)!important;"
        "border:1px dashed var(--border2)!important;border-radius:var(--r)!important;}"
        "[data-testid='stDataFrame']{border-radius:var(--r);overflow:hidden;box-shadow:var(--shadow);}"
        "[data-testid='stVerticalBlock']{gap:0!important;}"
        "div[data-testid='stHorizontalBlock']{gap:12px!important;}"

        # DS CARDS
        ".ds-card{background:var(--card);border:1px solid var(--border);border-radius:var(--r);"
        "padding:2rem 1.8rem;text-align:center;transition:all .22s;position:relative;overflow:hidden;}"
        ".ds-card::before{content:'';position:absolute;top:0;left:0;right:0;height:3px;"
        "background:" + SI_RED + ";transform:scaleX(0);transition:transform .3s;transform-origin:left;}"
        ".ds-card:hover{border-color:" + SI_RED + ";transform:translateY(-3px);"
        "box-shadow:0 12px 40px rgba(227,0,27,.18);}"
        ".ds-card:hover::before{transform:scaleX(1);}"
        ".ds-icon{font-size:2.5rem;margin-bottom:.8rem;}"
        ".ds-title{font-size:1.05rem;font-weight:800;color:var(--text);margin-bottom:.4rem;}"
        ".ds-desc{font-size:.82rem;color:var(--muted);line-height:1.6;}"
        ".step-badge{display:inline-block;background:rgba(227,0,27,.15);color:" + SI_RED + ";"
        "font-size:.69rem;font-weight:800;letter-spacing:.09em;text-transform:uppercase;"
        "padding:3px 12px;border-radius:4px;margin-bottom:1.2rem;}"
        ".kf-card{background:var(--card);border:1px solid var(--border);border-left:3px solid " + SI_RED + ";"
        "border-radius:var(--r);padding:.8rem 1.1rem;margin-bottom:.45rem;"
        "display:flex;align-items:center;gap:12px;}"
        ".kf-name{font-size:.87rem;font-weight:600;color:var(--text);}"
        ".kf-meta{font-size:.72rem;color:var(--muted);margin-top:2px;}"
        ".kf-badge{font-size:.7rem;font-weight:700;padding:2px 9px;border-radius:4px;"
        "background:rgba(16,185,129,.15);color:#34D399;margin-left:auto;white-space:nowrap;}"
        ".step-row{display:flex;align-items:center;gap:8px;font-size:.8rem;color:var(--muted);margin-bottom:1.5rem;}"
        ".sdot{width:26px;height:26px;border-radius:50%;display:flex;align-items:center;"
        "justify-content:center;font-size:.72rem;font-weight:800;flex-shrink:0;}"
        ".sa{background:" + SI_RED + ";color:#fff;} .sd{background:#10B981;color:#fff;}"
        ".sn{background:var(--dim);color:var(--muted);}"
        ".sline{flex:1;height:1px;background:var(--border);}"

        # LOGIN
        ".login-card{background:var(--card);border:1px solid var(--border2);"
        "border-top:4px solid " + SI_RED + ";border-radius:16px;padding:2.8rem 2.5rem;"
        "box-shadow:" + login_sh + ";}"

        # SIDEBAR PROFILE
        "[data-testid='stSidebar']{background:var(--navbg)!important;"
        "border-right:1px solid var(--border)!important;}"
        "[data-testid='stSidebar'] *{font-family:var(--font)!important;}"
        ".si-avatar{width:56px;height:56px;background:" + SI_RED + ";border-radius:50%;"
        "display:flex;align-items:center;justify-content:center;"
        "font-size:1.4rem;font-weight:900;color:#fff;margin:0 auto 1rem;}"
        ".si-profile-name{font-size:1rem;font-weight:700;color:var(--text);"
        "text-align:center;margin-bottom:.2rem;}"
        ".si-profile-role{font-size:.72rem;color:" + SI_RED + ";text-align:center;"
        "font-weight:700;letter-spacing:.08em;text-transform:uppercase;margin-bottom:1.2rem;}"
        ".si-stat{background:var(--panel);border:1px solid var(--border);"
        "border-radius:8px;padding:.6rem .9rem;margin-bottom:.4rem;"
        "display:flex;justify-content:space-between;align-items:center;}"
        ".si-stat-label{font-size:.72rem;color:var(--muted);}"
        ".si-stat-val{font-size:.85rem;font-weight:700;color:var(--text);}"
        ".si-sb-sep{border:none;border-top:1px solid var(--border);margin:.9rem 0;}"
        ".si-threshold-card{background:var(--panel);border:1px solid var(--border);"
        "border-left:3px solid " + SI_RED + ";border-radius:8px;"
        "padding:.7rem .9rem;margin-bottom:.5rem;}"
        ".si-thr-label{font-size:.72rem;font-weight:700;color:var(--text);margin-bottom:.2rem;}"
        ".si-thr-val{font-size:.8rem;color:" + SI_RED + ";font-weight:700;}"

        # ALERT TOAST SYSTEM
        ".alert-toast{"
        "border-radius:10px;padding:.75rem 1rem;margin-bottom:.55rem;"
        "display:flex;align-items:flex-start;gap:10px;animation:slideIn .3s ease;}"
        "@keyframes slideIn{from{opacity:0;transform:translateX(-10px)}to{opacity:1;transform:translateX(0)}}"
        ".alert-critical{background:rgba(227,0,27,.10);border:1px solid rgba(227,0,27,.3);"
        "border-left:4px solid " + SI_RED + ";}"
        ".alert-warning{background:rgba(245,158,11,.10);border:1px solid rgba(245,158,11,.3);"
        "border-left:4px solid #F59E0B;}"
        ".alert-info{background:rgba(59,130,246,.10);border:1px solid rgba(59,130,246,.3);"
        "border-left:4px solid #3B82F6;}"
        ".alert-ok{background:rgba(16,185,129,.10);border:1px solid rgba(16,185,129,.3);"
        "border-left:4px solid #10B981;}"
        ".alert-icon{font-size:1.1rem;flex-shrink:0;margin-top:1px;}"
        ".alert-body{flex:1;}"
        ".alert-title{font-size:.84rem;font-weight:700;color:var(--text);margin-bottom:.2rem;}"
        ".alert-desc{font-size:.76rem;color:var(--muted);line-height:1.5;}"
        ".alert-time{font-size:.68rem;color:var(--muted);margin-top:.2rem;font-family:var(--mono);}"

        # DELTA PILL (period comparison)
        ".delta-pill{display:inline-flex;align-items:center;gap:4px;padding:2px 8px;"
        "border-radius:20px;font-size:.69rem;font-weight:700;margin-left:6px;}"
        ".dp-up{background:rgba(16,185,129,.15);color:#34D399;}"
        ".dp-down{background:rgba(227,0,27,.15);color:#FF6680;}"
        ".dp-flat{background:rgba(100,116,139,.15);color:#94A3B8;}"

        # CHART EXPORT BUTTON
        ".chart-export-wrap{position:relative;}"
        ".export-btn-area{display:flex;justify-content:flex-end;margin-bottom:4px;gap:6px;}"
        ".stDownloadButton>button{font-size:.72rem!important;padding:.28rem .75rem!important;"
        "border-radius:6px!important;}"
    )
    st.markdown(f"<style>{css}</style>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 1 — CHART EXPORT  (PNG download button next to every chart)
# ══════════════════════════════════════════════════════════════════════════════
def pc(fig, title="chart", h=340, leg=True):
    """Render Plotly chart with PNG export button."""
    fig = T(fig, h=h, leg=leg)
    col_chart, col_btn = st.columns([1, 0.001])
    with col_chart:
        st.plotly_chart(fig, use_container_width=True,
                        config={"displayModeBar": True,
                                "modeBarButtonsToRemove": ["select2d","lasso2d"],
                                "toImageButtonOptions": {
                                    "format": "png", "filename": title,
                                    "height": h*2, "width": 1200, "scale": 2
                                }})
    # Also provide a Streamlit download button as fallback
    try:
        img_bytes = fig.to_image(format="png", width=1200, height=h*2, scale=2)
        st.download_button(
            f"⬇ PNG",
            data=img_bytes,
            file_name=f"{title.replace(' ','_')}.png",
            mime="image/png",
            key=f"dl_{title}_{id(fig)}",
            help=f"Download '{title}' as PNG"
        )
    except Exception:
        pass   # kaleido not available — toolbar still works


# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 2 — PERIOD COMPARISON HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def _split_periods(df, date_col, ref_date=None):
    """
    Split df into current period (last 90 days) and previous period (90 days before that).
    Returns (df_current, df_previous).
    """
    if not has(df, date_col) or df.empty:
        return df, pd.DataFrame()
    df2 = df.copy()
    df2[date_col] = pd.to_datetime(df2[date_col], errors="coerce")
    df2 = df2.dropna(subset=[date_col])
    if df2.empty:
        return df, pd.DataFrame()
    ref = ref_date or df2[date_col].max()
    cut = ref - timedelta(days=90)
    cut2 = cut - timedelta(days=90)
    cur = df2[df2[date_col] >= cut]
    prv = df2[(df2[date_col] >= cut2) & (df2[date_col] < cut)]
    return cur, prv

def delta_pill(current, previous, fmt=None, inverse=False):
    """
    Return HTML delta pill showing % change from previous to current.
    inverse=True means higher is worse (e.g. failures).
    """
    if previous == 0:
        return ""
    pct = (current - previous) / previous * 100
    if abs(pct) < 1:
        cls, arrow = "dp-flat", "→"
    elif pct > 0:
        cls = "dp-down" if inverse else "dp-up"
        arrow = "↑"
    else:
        cls = "dp-up" if inverse else "dp-down"
        arrow = "↓"
    val = fmt(abs(pct)) if fmt else f"{abs(pct):.1f}%"
    return f"<span class='delta-pill {cls}'>{arrow} {val}</span>"

def kpi_with_delta(label, value_str, value_num, prev_num,
                   sub="", color="blue", icon="",
                   inverse=False, period_label="vs prev 90 days"):
    """KPI card with period comparison delta pill."""
    dp = delta_pill(value_num, prev_num, inverse=inverse) if prev_num is not None else ""
    return f"""<div class='kpi-card kpi-{color}'>
      <div class='kpi-icon'>{icon}</div>
      <div class='kpi-label'>{label}</div>
      <div class='kpi-value'>{value_str}{dp}</div>
      <div class='kpi-sub'>{sub}</div>
      <div style='font-size:.67rem;color:var(--muted);margin-top:.2rem;'>{period_label}</div>
    </div>"""


# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 3 — ALERT ENGINE  (configurable thresholds)
# ══════════════════════════════════════════════════════════════════════════════
DEFAULT_THRESHOLDS = {
    "wash_coverage_min":    75,   # % min WASH coverage
    "fsl_coverage_min":     70,   # % min FSL HH coverage
    "cva_failure_max":       5,   # max % CVA transfers failed
    "pipeline_break_max":    2,   # max pipeline breaks
    "indicator_offtrack_max":20,  # max % off-track indicators
    "beneficiary_suspended_max": 10,  # max % suspended beneficiaries
}

def get_thresholds():
    return st.session_state.get("thresholds", DEFAULT_THRESHOLDS.copy())

def set_thresholds(t):
    st.session_state["thresholds"] = t

def run_alert_engine(dfs):
    """
    Evaluate all data against configurable thresholds.
    Returns list of alert dicts: {level, icon, title, desc}.
    """
    thr    = get_thresholds()
    alerts = []
    now    = datetime.now().strftime("%H:%M")

    df_b = dfs.get("Beneficiary_Registration", pd.DataFrame())
    df_w = dfs.get("WASH_Activities",           pd.DataFrame())
    df_f = dfs.get("Protection_Monitoring",          pd.DataFrame())
    df_c = dfs.get("Health_Activities",        pd.DataFrame())
    df_i = dfs.get("Indicator_Tracker",         pd.DataFrame())

    # ── WASH coverage ─────────────────────────────────────────────────────────
    if not df_w.empty:
        reached = ssum(df_w,"Beneficiaires_Atteints")
        target  = ssum(df_w,"Target_Beneficiaries") or 1
        pct     = reached / target * 100
        if pct < thr["wash_coverage_min"]:
            alerts.append({"level":"critical","icon":"🚨","title":f"WASH Coverage Low — {pct:.0f}%",
                "desc":f"Below threshold of {thr['wash_coverage_min']}%. "
                       f"{reached:,.0f} reached of {target:,.0f} targeted.",
                "time":now})
        elif pct < thr["wash_coverage_min"] + 10:
            alerts.append({"level":"warning","icon":"⚠️","title":f"WASH Coverage At Risk — {pct:.0f}%",
                "desc":f"Approaching minimum threshold of {thr['wash_coverage_min']}%.",
                "time":now})

    # ── FSL coverage ──────────────────────────────────────────────────────────
    if not df_f.empty:
        hhr = ssum(df_f,"Beneficiaires_Atteints"); hht = ssum(df_f,"Cible_Beneficiaires") or 1
        fsl_pct = hhr / hht * 100
        if fsl_pct < thr["fsl_coverage_min"]:
            alerts.append({"level":"critical","icon":"🚨","title":f"FSL HH Coverage Low — {fsl_pct:.0f}%",
                "desc":f"Only {hhr:,.0f} of {hht:,.0f} targeted households reached.",
                "time":now})

        # Pipeline breaks
        pb = slen(df_f,"Type_Programme","Pipeline break")
        if pb > thr["pipeline_break_max"]:
            alerts.append({"level":"critical","icon":"⛓️","title":f"FSL Supply Chain — {pb} Pipeline Breaks",
                "desc":f"Exceeds maximum threshold of {thr['pipeline_break_max']} breaks. "
                       "Immediate logistics coordination required.",
                "time":now})

    # ── CVA failures ──────────────────────────────────────────────────────────
    if not df_c.empty and len(df_c) > 0:
        failed  = slen(df_c,"Statut","Failed")
        fail_pct = failed / len(df_c) * 100
        if fail_pct > thr["cva_failure_max"]:
            alerts.append({"level":"critical","icon":"💳","title":f"CVA Transfer Failure Rate — {fail_pct:.1f}%",
                "desc":f"{failed} failed transfers ({fail_pct:.1f}% of total). "
                       f"Threshold: {thr['cva_failure_max']}%. Verify payment agent credentials.",
                "time":now})
        pending = slen(df_c,"Statut","Pending")
        if pending > 0:
            alerts.append({"level":"warning","icon":"⏳","title":f"CVA — {pending} Pending Transfers",
                "desc":"Payment cycle not yet completed. Verify before next distribution round.",
                "time":now})

    # ── Off-track indicators ───────────────────────────────────────────────────
    if not df_i.empty and len(df_i) > 0:
        off   = slen(df_i,"Status","Hors trajectoire")
        at_r  = slen(df_i,"Status","À risque")
        off_pct = off / len(df_i) * 100
        if off_pct > thr["indicator_offtrack_max"]:
            alerts.append({"level":"critical","icon":"📉","title":f"Indicators — {off} Off Track ({off_pct:.0f}%)",
                "desc":f"Exceeds maximum threshold of {thr['indicator_offtrack_max']}% off-track. "
                       "Corrective measures required in next MPR.",
                "time":now})
        if at_r > 0:
            alerts.append({"level":"warning","icon":"📊","title":f"Indicators — {at_r} At Risk",
                "desc":"Monitor closely and prepare contingency measures.",
                "time":now})

    # ── Suspended beneficiaries ────────────────────────────────────────────────
    if not df_b.empty and len(df_b) > 0:
        susp = slen(df_b,"Statut_Enregistrement","Suspended")
        susp_pct = susp / len(df_b) * 100
        if susp_pct > thr["beneficiary_suspended_max"]:
            alerts.append({"level":"warning","icon":"🔴","title":f"{susp} Beneficiaries Suspended ({susp_pct:.1f}%)",
                "desc":"Review suspended cases and resolve eligibility issues.",
                "time":now})

    # ── All clear ─────────────────────────────────────────────────────────────
    if not alerts:
        alerts.append({"level":"ok","icon":"✅","title":"All systems nominal",
            "desc":"No threshold violations detected at this time.",
            "time":now})

    return alerts

def render_alert(alert):
    """Render a single alert toast."""
    level_css = {"critical":"alert-critical","warning":"alert-warning",
                 "info":"alert-info","ok":"alert-ok"}.get(alert["level"],"alert-info")
    st.markdown(f"""
    <div class='alert-toast {level_css}'>
      <div class='alert-icon'>{alert['icon']}</div>
      <div class='alert-body'>
        <div class='alert-title'>{alert['title']}</div>
        <div class='alert-desc'>{alert['desc']}</div>
        <div class='alert-time'>{alert['time']}</div>
      </div>
    </div>""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# FEATURE 4 — SIDEBAR with profile, alerts & threshold settings
# ══════════════════════════════════════════════════════════════════════════════
def render_sidebar(dfs):
    """Render profile + alerts + threshold settings in Streamlit sidebar."""
    th         = TH()
    user       = st.session_state.get("user", "im_manager")
    src        = st.session_state.get("data_source", "excel")
    dark       = st.session_state.get("dark", True)
    login_time = st.session_state.get("login_time",
                                      datetime.now().strftime("%d %b %Y %H:%M"))

    # Sidebar-specific styling
    st.sidebar.markdown(f"""
    <style>
    [data-testid="stSidebar"] {{
        background: {"#0E1120" if dark else "#FFFFFF"} !important;
        border-right: 1px solid {SI_RED}33 !important;
    }}
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] label {{
        color: {"#94A3B8" if dark else "#64748B"} !important;
        font-family: Outfit, sans-serif !important;
    }}
    [data-testid="stSidebar"] .stButton > button {{
        background: {"rgba(227,0,27,0.1)" if dark else "#FEF2F2"} !important;
        border: 1px solid {SI_RED}44 !important;
        color: {SI_RED} !important;
        border-radius: 8px !important;
        font-size: .82rem !important;
        font-weight: 600 !important;
    }}
    [data-testid="stSidebar"] .stButton > button:hover {{
        background: {SI_RED} !important;
        color: #fff !important;
    }}
    </style>
    """, unsafe_allow_html=True)

    # ── Avatar & identity ─────────────────────────────────────────────────────
    bg_card  = "#1C2130" if dark else "#F8F9FB"
    tc       = "#F1F5F9" if dark else "#0F172A"
    mc       = "#94A3B8" if dark else "#64748B"
    brd      = "rgba(227,0,27,0.2)"

    st.sidebar.markdown(f"""
    <div style='text-align:center;padding:1.5rem 0 1rem;'>
      <div style='width:58px;height:58px;background:{SI_RED};border-radius:50%;
           display:flex;align-items:center;justify-content:center;
           font-size:1.5rem;font-weight:900;color:#fff;margin:0 auto .9rem;
           box-shadow:0 4px 16px rgba(227,0,27,0.4);'>👤</div>
      <div style='font-size:1rem;font-weight:700;color:{tc};
           font-family:Outfit,sans-serif;'>{user.replace("_"," ").title()}</div>
      <div style='font-size:.68rem;font-weight:800;color:{SI_RED};
           letter-spacing:.1em;text-transform:uppercase;margin-top:3px;'>
           IM Manager</div>
    </div>
    <hr style='border:none;border-top:1px solid {brd};margin:.5rem 0 1rem;'>
    """, unsafe_allow_html=True)

    # ── Session statistics ────────────────────────────────────────────────────
    sheets     = list(dfs.keys()) if dfs else []
    total_rows = sum(len(v) for v in dfs.values()) if dfs else 0
    src_label  = "🔗 KoBoToolbox" if src == "kobo" else "📊 Excel"

    def stat_row(label, val):
        return (f"<div style='display:flex;justify-content:space-between;"
                f"align-items:center;padding:.45rem .8rem;"
                f"background:{bg_card};border-radius:8px;margin-bottom:.35rem;"
                f"border:1px solid {brd};'>"
                f"<span style='font-size:.72rem;color:{mc};'>{label}</span>"
                f"<span style='font-size:.82rem;font-weight:700;color:{tc};'>{val}</span>"
                f"</div>")

    st.sidebar.markdown(
        stat_row("Data source", src_label) +
        stat_row("Sheets loaded", str(len(sheets))) +
        stat_row("Total records", f"{total_rows:,}") +
        stat_row("Connected", login_time) +
        stat_row("Last refresh", datetime.now().strftime("%H:%M:%S")),
        unsafe_allow_html=True
    )

    st.sidebar.markdown(f"<hr style='border:none;border-top:1px solid {brd};margin:.9rem 0;'>",
                        unsafe_allow_html=True)

    # ── Quick actions ─────────────────────────────────────────────────────────
    st.sidebar.markdown(f"<div style='font-size:.67rem;font-weight:800;color:{SI_RED};"
                        f"letter-spacing:.12em;text-transform:uppercase;margin-bottom:.6rem;'>"
                        f"Quick Actions</div>", unsafe_allow_html=True)

    c1, c2 = st.sidebar.columns(2)
    with c1:
        if st.button("🔄 Refresh", use_container_width=True,
                     key="sb_refresh", help="Clear cache and reload"):
            st.cache_data.clear(); st.rerun()
    with c2:
        th_lbl = "☀️ Light" if dark else "🌙 Dark"
        if st.button(th_lbl, use_container_width=True,
                     key="sb_theme", help="Toggle theme"):
            st.session_state["dark"] = not dark; st.rerun()

    if st.sidebar.button("🔌 Change Data Source", use_container_width=True, key="sb_src"):
        st.session_state["kobo_step"] = 0
        st.session_state.pop("dfs", None); st.rerun()

    if st.sidebar.button("🚪 Logout", use_container_width=True, key="sb_logout"):
        st.session_state.clear(); st.rerun()

    st.sidebar.markdown(f"<hr style='border:none;border-top:1px solid {brd};margin:.9rem 0;'>",
                        unsafe_allow_html=True)

    # ── Alert threshold configurator ──────────────────────────────────────────
    st.sidebar.markdown(f"<div style='font-size:.67rem;font-weight:800;color:{SI_RED};"
                        f"letter-spacing:.12em;text-transform:uppercase;margin-bottom:.8rem;'>"
                        f"⚙️ Alert Thresholds</div>", unsafe_allow_html=True)


    # ── Branding footer ───────────────────────────────────────────────────────
    st.sidebar.markdown(f"""
    <div style='text-align:center;margin-top:1.5rem;padding:.8rem;
         border-top:1px solid {brd};'>
      <div style='font-size:.68rem;font-weight:800;letter-spacing:.06em;
           color:{SI_RED};'>INTERSOS</div>
      <div style='font-size:.63rem;color:{mc};margin-top:2px;'>
        Mission Tchad · IM Platform © 2026
      </div>
    </div>
    """, unsafe_allow_html=True)


def render_alert_sidebar(alert, dark=True):
    """Compact alert for sidebar display."""
    bg_map  = {"critical": "rgba(227,0,27,0.10)",
               "warning":  "rgba(245,158,11,0.10)",
               "info":     "rgba(59,130,246,0.10)",
               "ok":       "rgba(16,185,129,0.10)"}
    bd_map  = {"critical": SI_RED, "warning": "#F59E0B",
               "info": "#3B82F6", "ok": "#10B981"}
    tc      = "#F1F5F9" if dark else "#0F172A"
    mc      = "#94A3B8" if dark else "#64748B"
    bg      = bg_map.get(alert["level"], "rgba(100,116,139,0.10)")
    bd      = bd_map.get(alert["level"], "#64748B")

    st.sidebar.markdown(f"""
    <div style='background:{bg};border-left:3px solid {bd};border-radius:8px;
         padding:.6rem .8rem;margin-bottom:.4rem;'>
      <div style='font-size:.78rem;font-weight:700;color:{tc};'>
        {alert['icon']} {alert['title']}</div>
      <div style='font-size:.7rem;color:{mc};margin-top:2px;line-height:1.4;'>
        {alert['desc']}</div>
    </div>
    """, unsafe_allow_html=True)



# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def N(n, d=0):
    try:
        n = float(n)
        if pd.isna(n): return "—"
        if n >= 1_000_000: return f"{n/1_000_000:.1f}M"
        if n >= 1_000:     return f"{n/1_000:.1f}K"
        return f"{n:,.{d}f}"
    except: return "—"

def has(df, col): return col in df.columns
def ssum(df, col): return df[col].sum() if has(df, col) else 0
def slen(df, col, val): return len(df[df[col]==val]) if has(df, col) else 0
def suniq(df, col): return df[col].nunique() if has(df, col) else 0
def sopts(df, col): return ["All"] + sorted(df[col].dropna().unique().tolist()) if has(df, col) else ["All"]
def sfilt(df, col, val):
    if val == "All" or not has(df, col): return df
    return df[df[col]==val]

def TH(): return DARK if st.session_state.get("dark", True) else LIGHT

def bdg(status):
    s = str(status).lower()
    if any(x in s for x in ["on track","paid","completed","above 80","in stock","fully"]):
        return f"<span class='badge bg'>{status}</span>"
    if any(x in s for x in ["risk","pending","partial","low stock","60–80","awaiting"]):
        return f"<span class='badge ba'>{status}</span>"
    if any(x in s for x in ["off","fail","non-func","pipeline break","cancel","below 60","not"]):
        return f"<span class='badge br'>{status}</span>"
    return f"<span class='badge bn'>{status}</span>"

def kpi(label, value, sub="", color="blue", icon="", delta=None, ddir="up"):
    dc = {"up":"d-up","down":"d-down","mid":"d-mid"}.get(ddir,"d-mid")
    da = f"<div class='kpi-delta {dc}'>{delta}</div>" if delta else ""
    return f"""<div class='kpi-card kpi-{color}'>
      <div class='kpi-icon'>{icon}</div>
      <div class='kpi-label'>{label}</div>
      <div class='kpi-value'>{value}</div>
      <div class='kpi-sub'>{sub}</div>{da}
    </div>"""

def sh(t): st.markdown(f"<div class='sh'>{t}</div>", unsafe_allow_html=True)
def ph(title, sub=""): st.markdown(f"<div class='ph'><h1>{title}</h1><p>{sub}</p></div>", unsafe_allow_html=True)

def T(fig, h=340, leg=True):
    th = TH()
    fig.update_layout(
        height=h, paper_bgcolor=th["plotbg"], plot_bgcolor=th["plotbg"],
        font=dict(family="Outfit", color=th["fontc"], size=11),
        margin=dict(l=12,r=12,t=36,b=12), showlegend=leg,
        legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(color=th["muted"],size=10),
                    borderwidth=0, orientation="h", yanchor="bottom", y=1.02, xanchor="left", x=0),
        colorway=COLORS,
        title_font=dict(size=12, color=th["title"], family="Outfit"),
    )
    fig.update_xaxes(showgrid=False, zeroline=False, color=th["fontc"],
                     tickfont=dict(size=10,color=th["fontc"]),
                     linecolor=th["gridc"])
    fig.update_yaxes(gridcolor=th["gridc"], zeroline=False, color=th["fontc"],
                     tickfont=dict(size=10,color=th["fontc"]), linewidth=0)
    return fig

# ══════════════════════════════════════════════════════════════════════════════
# DATA
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_data(show_spinner="Loading…")
def load_data(file):
    xls = pd.ExcelFile(file, engine="openpyxl")
    out = {}
    for name in xls.sheet_names:
        try: out[name] = pd.read_excel(xls, sheet_name=name)
        except: pass
    return out

# ══════════════════════════════════════════════════════════════════════════════
# TOP NAVIGATION  — reliable active-tab highlighting via nth-child CSS
# ══════════════════════════════════════════════════════════════════════════════
def top_nav():
    page = st.session_state.get("page", "Overview")
    user = st.session_state.get("user", "")
    dark = st.session_state.get("dark", True)
    src  = st.session_state.get("data_source", "—")
    src_icon = "🔗" if src == "kobo" else "📊"
    labels    = [lbl for _, lbl in NAV_ITEMS]
    active_idx = labels.index(page) + 1 if page in labels else 1
    n_nav     = len(NAV_ITEMS)

    # ── SI branded topbar ─────────────────────────────────────────────────────
    st.markdown(f"""
    <div class='si-topbar'>
      <div class='si-logo-box'>IN</div>
      <div>
        <div class='si-brand'>INTERSOS
          <small>Mission Tchad · Information Management</small>
        </div>
      </div>
      <div style='margin-left:auto;display:flex;align-items:center;gap:10px;padding:2rem 2rem'>
      </div>
    </div>""", unsafe_allow_html=True)

    # ── Active tab nth-child injection ────────────────────────────────────────
    st.markdown(
        "<style>"
        f"div[data-testid='stHorizontalBlock']>div:nth-child({active_idx})"
        f">[data-testid='stButton'] button{{"
        f"color:#fff!important;"
        f"border-bottom:3px solid {SI_RED}!important;"
        f"font-weight:700!important;"
        f"background:rgba(227,0,27,0.13)!important;}}"
        f"div[data-testid='stHorizontalBlock']>div:nth-child({n_nav})"
        f">[data-testid='stButton'] button,"
        f"div[data-testid='stHorizontalBlock']>div:nth-child({n_nav})"
        f">[data-testid='stButton'] button{{"
        f"font-size:.78rem!important;border-radius:20px!important;padding:4rem 4rem!important;"
        f"background:rgba(255,255,255,0.08)!important;"
        f"border:1px solid rgba(255,255,255,0.15)!important;"
        f"color:rgba(255,255,255,.6)!important;padding:.35rem .8rem!important;}}"
        "</style>",
        unsafe_allow_html=True
    )

    # ── Nav button row ────────────────────────────────────────────────────────
    cols = st.columns([1]*n_nav + [0.55, 0.55])
    for i, (icon, lbl) in enumerate(NAV_ITEMS):
        with cols[i]:
            if st.button(f"{icon} {lbl}", key=f"nav_{lbl}",
                         use_container_width=True, help=lbl):
                if st.session_state.get("page") != lbl:
                    st.session_state["page"] = lbl
                    st.rerun()
    

    st.markdown("<div style='padding:1.4rem 2rem 0;'>", unsafe_allow_html=True)



# LOGIN - KOBO INTEGRATION BEFORE THIS
# ══════════════════════════════════════════════════════════════════════════════

KOBO_SERVERS = {
    "KoBoToolbox Global — kf.kobotoolbox.org": "https://kf.kobotoolbox.org",
    "OCHA Humanitarian — kobo.humanitarianresponse.info": "https://kobo.humanitarianresponse.info",
    "Custom server…": "custom",
}

def kobo_get_token(username, password, base_url):
    url = f"{base_url}/token/?format=json"
    try:
        r = requests.get(url, auth=(username, password), timeout=20)
        r.raise_for_status()
        return r.json().get("token"), None
    except Exception as e:
        code = getattr(e.response if hasattr(e,"response") else None,"status_code",None)
        if code == 401: return None, "Invalid KoBoToolbox username or password."
        return None, str(e)

def kobo_list_forms(token, base_url):
    url = f"{base_url}/api/v2/assets/?asset_type=survey&format=json"
    headers = {"Authorization": f"Token {token}"}
    try:
        r = requests.get(url, headers=headers, timeout=25)
        r.raise_for_status()
        data = r.json()
        forms = []
        for a in data.get("results",[]):
            forms.append({
                "uid":   a.get("uid",""),
                "name":  a.get("name","Unnamed"),
                "submissions": a.get("deployment__submission_count",0),
                "owner": a.get("owner__username",""),
                "modified": str(a.get("date_modified",""))[:10],
            })
        return forms, None
    except Exception as e:
        return None, str(e)

def kobo_download_data(token, uid, base_url):
    url = f"{base_url}/api/v2/assets/{uid}/data/?format=json&limit=30000"
    headers = {"Authorization": f"Token {token}"}
    try:
        r = requests.get(url, headers=headers, timeout=60)
        r.raise_for_status()
        results = r.json().get("results", [])
        if not results: return pd.DataFrame(), None
        df = pd.json_normalize(results)
        df.columns = [c.split("/")[-1].lstrip("_") for c in df.columns]
        df = df.loc[:, ~df.columns.duplicated()]
        return df, None
    except Exception as e:
        return None, str(e)

def datasource_page():
    inject_css(DARK)
    # Hide sidebar on datasource selection screen
    st.markdown("<style>[data-testid='stSidebar']{display:none!important;}</style>",
                unsafe_allow_html=True)
    th = DARK
    st.markdown("""
<style>
.ds-card{background:var(--panel);border:1px solid var(--border2);border-radius:14px;
  padding:1.8rem 2rem;text-align:center;transition:all .2s;}
.ds-card:hover{border-color:var(--accent);transform:translateY(-3px);box-shadow:0 12px 40px rgba(59,130,246,0.15);}
.ds-icon{font-size:2.6rem;margin-bottom:.8rem;}
.ds-title{font-size:1.1rem;font-weight:700;color:var(--text);margin-bottom:.4rem;}
.ds-desc{font-size:.82rem;color:var(--muted);line-height:1.6;}
.step-badge{display:inline-block;background:rgba(59,130,246,0.15);color:#60A5FA;
  font-size:.7rem;font-weight:700;letter-spacing:.08em;text-transform:uppercase;
  padding:3px 10px;border-radius:20px;margin-bottom:1.2rem;}
.kf-card{background:var(--panel);border:1px solid var(--border);border-radius:10px;
  padding:.85rem 1.1rem;margin-bottom:.45rem;display:flex;align-items:center;gap:12px;}
.kf-card:hover{border-color:var(--border2);}
.kf-name{font-size:.87rem;font-weight:600;color:var(--text);}
.kf-meta{font-size:.72rem;color:var(--muted);margin-top:2px;}
.kf-badge{font-size:.7rem;font-weight:700;padding:2px 8px;border-radius:20px;
  background:rgba(16,185,129,0.15);color:#34D399;margin-left:auto;white-space:nowrap;}
.step-row{display:flex;align-items:center;gap:8px;font-size:.8rem;color:var(--muted);margin-bottom:1.5rem;}
.sdot{width:24px;height:24px;border-radius:50%;display:flex;align-items:center;
  justify-content:center;font-size:.72rem;font-weight:700;flex-shrink:0;}
.sa{background:var(--accent);color:#fff;}
.sd{background:rgba(16,185,129,0.8);color:#fff;}
.sn{background:var(--dim);color:var(--muted);}
.sline{flex:1;height:1px;background:var(--border);}
</style>""", unsafe_allow_html=True)

    step = st.session_state.get("kobo_step", 0)

    def step_bar(current):
        labels = ["Source","Connect","Forms","Load"]
        html = "<div class='step-row'>"
        for i,s in enumerate(labels):
            cls = "sa" if i==current else ("sd" if i<current else "sn")
            wt  = "700" if i==current else "400"
            col = "var(--text)" if i==current else "var(--muted)"
            html += f"<div class='sdot {cls}'>{i+1}</div>"
            html += f"<span style='color:{col};font-weight:{wt};'>{s}</span>"
            if i<len(labels)-1: html += "<div class='sline'></div>"
        html += "</div>"
        st.markdown(html, unsafe_allow_html=True)

    # STEP 0: Choose source
    if step == 0:
        c1,c2,c3 = st.columns([1,2,1])
        with c2:
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown("<div class='step-badge'>Step 1 of 4 — Data Source</div>", unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:1.5rem;font-weight:800;color:{th['text']};margin-bottom:.3rem;font-family:Outfit,sans-serif;'>Choose your data source</div>",unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:.87rem;color:{th['muted']};margin-bottom:1.8rem;'>Select how to load program data into the dashboard.</div>",unsafe_allow_html=True)
            step_bar(0)
            ca,cb = st.columns(2)
            with ca:
                st.markdown("""<div class='ds-card'>
                  <div class='ds-icon'>📊</div>
                  <div class='ds-title'>Excel File</div>
                  <div class='ds-desc'>Upload an <b>.xlsx</b> file. Best for pre-processed or offline data.</div>
                </div>""",unsafe_allow_html=True)
                st.markdown("<br>",unsafe_allow_html=True)
                if st.button("Use Excel File →", use_container_width=True, key="src_excel"):
                    st.session_state.update(data_source="excel", kobo_step=10); st.rerun()
            with cb:
                st.markdown("""<div class='ds-card'>
                  <div class='ds-icon'>🔗</div>
                  <div class='ds-title'>KoBoToolbox Live</div>
                  <div class='ds-desc'>Connect directly to KoBoToolbox and pull real-time form submissions.</div>
                </div>""",unsafe_allow_html=True)
                st.markdown("<br>",unsafe_allow_html=True)
                if st.button("Connect to KoBoToolbox →", use_container_width=True, key="src_kobo"):
                    st.session_state.update(data_source="kobo", kobo_step=1); st.rerun()
            st.markdown(f"<br><div style='text-align:center;font-size:.72rem;color:{th['muted']};'>You can switch data source anytime from the dashboard.</div>",unsafe_allow_html=True)

    # STEP 10: Excel upload
    elif step == 10:
        c1,c2,c3 = st.columns([1,2,1])
        with c2:
            st.markdown("<br>",unsafe_allow_html=True)
            st.markdown("<div class='step-badge'>Excel Upload</div>",unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:1.5rem;font-weight:800;color:{th['text']};margin-bottom:1.2rem;font-family:Outfit,sans-serif;'>Upload Excel database</div>",unsafe_allow_html=True)
            step_bar(3)
            up = st.file_uploader("Select .xlsx file", type=["xlsx"])
            if up:
                with st.spinner("Loading sheets…"):
                    data = load_data(up)
                st.session_state["dfs"] = data
                st.success(f"✅ {len(data)} sheets: {', '.join(list(data.keys())[:5])}")
                st.markdown("<br>",unsafe_allow_html=True)
                if st.button("Open Dashboard →", use_container_width=True, key="excel_go"):
                    st.session_state["kobo_step"] = 99; st.rerun()
            st.markdown("<br>",unsafe_allow_html=True)
            if st.button("← Back", key="excel_back"):
                st.session_state["kobo_step"]=0; st.rerun()

    # STEP 1: KoBoToolbox credentials
    elif step == 1:
        c1,c2,c3 = st.columns([1,2,1])
        with c2:
            st.markdown("<br>",unsafe_allow_html=True)
            st.markdown("<div class='step-badge'>Step 2 of 4 — Connect</div>",unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:1.5rem;font-weight:800;color:{th['text']};margin-bottom:.3rem;font-family:Outfit,sans-serif;'>KoBoToolbox Login</div>",unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:.85rem;color:{th['muted']};margin-bottom:1.2rem;'>Your password is used only to retrieve a temporary API token and is never stored.</div>",unsafe_allow_html=True)
            step_bar(1)

            srv_choice = st.selectbox("KoBoToolbox Server", list(KOBO_SERVERS.keys()), key="kobo_srv")
            base_url   = KOBO_SERVERS[srv_choice]
            if base_url == "custom":
                base_url = st.text_input("Custom server URL", placeholder="https://your-kobo-server.org", key="kobo_cust")

            kobo_u = st.text_input("Username", placeholder="your_username", key="kobo_u")
            kobo_p = st.text_input("Password", type="password", placeholder="••••••••", key="kobo_p")

            st.markdown(f"<div style='background:rgba(59,130,246,0.07);border:1px solid rgba(59,130,246,0.18);border-radius:8px;padding:.65rem 1rem;margin:.7rem 0;font-size:.77rem;color:{th['muted']};'>🔒 Credentials are stored only in your browser session and cleared on logout.</div>",unsafe_allow_html=True)

            ca,cb = st.columns([3,1])
            with ca:
                go = st.button("Connect →", use_container_width=True, key="kobo_go")
            with cb:
                if st.button("← Back", key="k1b"): st.session_state["kobo_step"]=0; st.rerun()

            if go:
                if not kobo_u or not kobo_p: st.error("Enter your username and password."); return
                if not base_url or base_url=="custom": st.error("Enter a valid server URL."); return
                if not HAS_REQUESTS: st.error("Install requests: pip install requests"); return
                with st.spinner("Authenticating…"):
                    token, err = kobo_get_token(kobo_u, kobo_p, base_url)
                if err: st.error(f"❌ {err}")
                else:
                    st.session_state.update(kobo_token=token, kobo_base_url=base_url,
                                            kobo_username=kobo_u, kobo_step=2)
                    st.rerun()

    # STEP 2: Select forms
    elif step == 2:
        token    = st.session_state.get("kobo_token","")
        base_url = st.session_state.get("kobo_base_url","")
        kuser    = st.session_state.get("kobo_username","")

        c1,c2,c3 = st.columns([0.4,3,0.4])
        with c2:
            st.markdown("<br>",unsafe_allow_html=True)
            st.markdown("<div class='step-badge'>Step 3 of 4 — Select Forms</div>",unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:1.5rem;font-weight:800;color:{th['text']};margin-bottom:.3rem;font-family:Outfit,sans-serif;'>Select KoBoToolbox Forms</div>",unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:.84rem;color:{th['muted']};margin-bottom:1rem;'>Connected as <b style='color:{th['text']};'>{kuser}</b> · {base_url}</div>",unsafe_allow_html=True)
            step_bar(2)

            if "kobo_forms" not in st.session_state:
                with st.spinner("Fetching forms list…"):
                    forms, err = kobo_list_forms(token, base_url)
                if err: st.error(f"❌ {err}"); st.button("← Back",key="k2be",on_click=lambda:st.session_state.update(kobo_step=1)); return
                st.session_state["kobo_forms"] = forms or []

            forms = st.session_state.get("kobo_forms",[])
            if not forms:
                st.warning("No forms found. Check your credentials or KoBoToolbox account.")
                if st.button("← Back",key="k2bno"): st.session_state["kobo_step"]=1; st.rerun()
                return

            search   = st.text_input("🔍 Filter forms by name", label_visibility="collapsed", placeholder="Search…", key="fsearch")
            filtered = [f for f in forms if search.lower() in f["name"].lower()] if search else forms
            st.markdown(f"<div style='font-size:.77rem;color:{th['muted']};margin-bottom:.7rem;'>{len(filtered)} form(s) shown</div>",unsafe_allow_html=True)

            selected = st.session_state.get("kobo_selected",{})
            for form in filtered:
                uid  = form["uid"];  name = form["name"]
                subs = form["submissions"]; owner = form["owner"]; mod = form["modified"]
                cc, ci = st.columns([0.07, 0.93])
                with cc: checked = st.checkbox("", value=selected.get(uid,False), key=f"chk_{uid}"); selected[uid]=checked
                with ci:
                    fc = "#34D399" if subs>0 else "#64748B"
                    st.markdown(f"""<div class='kf-card'>
                      <div style='flex:1;'>
                        <div class='kf-name'>📋 {name}</div>
                        <div class='kf-meta'>UID: <code>{uid}</code> · Owner: {owner} · Modified: {mod}</div>
                      </div>
                      <div class='kf-badge' style='color:{fc};'>{subs:,} submissions</div>
                    </div>""",unsafe_allow_html=True)
            st.session_state["kobo_selected"] = selected

            n_sel = sum(1 for v in selected.values() if v)
            st.markdown("<br>",unsafe_allow_html=True)
            ca,cb,cc2 = st.columns([1,2,1])
            with ca:
                if st.button("← Back",key="k2b"): st.session_state.pop("kobo_forms",None); st.session_state["kobo_step"]=1; st.rerun()
            with cb:
                st.markdown(f"<div style='text-align:center;font-size:.82rem;color:{th['text']};padding:.5rem;'><b>{n_sel}</b> form(s) selected</div>",unsafe_allow_html=True)
            with cc2:
                if st.button(f"Load {n_sel} →", use_container_width=True, key="k2load", disabled=(n_sel==0)):
                    st.session_state["kobo_step"]=3; st.rerun()

    # STEP 3: Download & map
    elif step == 3:
        token    = st.session_state.get("kobo_token","")
        base_url = st.session_state.get("kobo_base_url","")
        forms    = st.session_state.get("kobo_forms",[])
        selected = st.session_state.get("kobo_selected",{})
        sel_uids = [u for u,v in selected.items() if v]
        fnames   = {f["uid"]:f["name"] for f in forms}

        c1,c2,c3 = st.columns([1,2,1])
        with c2:
            st.markdown("<br>",unsafe_allow_html=True)
            st.markdown("<div class='step-badge'>Step 4 of 4 — Loading Data</div>",unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:1.5rem;font-weight:800;color:{th['text']};margin-bottom:.3rem;font-family:Outfit,sans-serif;'>Downloading Submissions</div>",unsafe_allow_html=True)
            st.markdown(f"<div style='font-size:.84rem;color:{th['muted']};margin-bottom:1.2rem;'>Downloading {len(sel_uids)} form(s). This may take a moment.</div>",unsafe_allow_html=True)
            step_bar(3)

            pb   = st.progress(0)
            stat = st.empty()
            dfs_k = {}; errs = []; logs = []

            for i, uid in enumerate(sel_uids):
                nm = fnames.get(uid,uid)
                pb.progress(i/len(sel_uids), text=f"Downloading: {nm}…")
                stat.markdown(f"<div style='font-size:.81rem;color:{th['muted']};'>⏳ {nm}…</div>",unsafe_allow_html=True)
                df, err = kobo_download_data(token, uid, base_url)
                if err:   errs.append(nm); logs.append(f"❌ **{nm}** — {err}")
                else:
                    dfs_k[nm] = df if df is not None else pd.DataFrame()
                    logs.append(f"✅ **{nm}** — {len(df) if df is not None else 0:,} records, {len(df.columns) if df is not None else 0} fields")

            pb.progress(1.0, text="Done!")
            stat.empty()

            for log in logs: st.markdown(log)

            if dfs_k:
                DASH_SHEETS = ["— Keep as-is —","Beneficiary_Registration","WASH_Activities",
                               "Protection_Monitoring","Health_Activities","Indicator_Tracker"]
                st.markdown("<br>",unsafe_allow_html=True)
                st.markdown(f"<div style='font-weight:700;font-size:.9rem;color:{th['text']};margin-bottom:.7rem;'>Map forms to dashboard sections (optional)</div>",unsafe_allow_html=True)
                mappings = {}
                for fnm in dfs_k:
                    ca2,cb2 = st.columns([1.6,1])
                    with ca2: st.markdown(f"<div style='font-size:.84rem;font-weight:500;color:{th['text']};padding:.55rem 0;'>📋 {fnm}</div>",unsafe_allow_html=True)
                    with cb2: mappings[fnm] = st.selectbox("→",DASH_SHEETS,key=f"mp_{fnm}",label_visibility="collapsed")
                st.session_state["kobo_mappings"] = mappings

                st.markdown("<br>",unsafe_allow_html=True)
                ca3,cb3 = st.columns([1,2])
                with ca3:
                    if st.button("← Back",key="k3b"): st.session_state["kobo_step"]=2; st.rerun()
                with cb3:
                    if st.button("Open Dashboard →", use_container_width=True, key="k3go"):
                        final = {}
                        for fnm,df in dfs_k.items():
                            tgt = mappings.get(fnm,"— Keep as-is —")
                            final[tgt if tgt!="— Keep as-is —" else fnm] = df
                        st.session_state["dfs"] = final
                        st.session_state["kobo_step"] = 99
                        st.rerun()
            else:
                st.error("No data loaded. Check permissions on the selected forms.")
                if st.button("← Start over",key="k3rst"): st.session_state["kobo_step"]=0; st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# LOGIN — SI brand, pure Streamlit layout
# ══════════════════════════════════════════════════════════════════════════════
def login_page():
    # SUPPRIMER TOUS LES STYLES EXISTANTS
    st.markdown("""
    <style>
    /* RESET COMPLET - Supprime tous les styles Streamlit par défaut */
    [data-testid="stAppViewContainer"] {
        background: linear-gradient(160deg, #0E0004 0%, #180008 40%, #0E0004 100%) !important;
    }
    
    [data-testid="stHeader"] {
        display: none !important;
    }
    
    [data-testid="stToolbar"] {
        display: none !important;
    }
    
    .main .block-container {
        padding: 0 !important;
        max-width: 100% !important;
        margin: 0 !important;
    }
    
    div[data-testid="stVerticalBlock"] {
        gap: 0 !important;
        padding: 0 !important;
        margin: 0 !important;
    }
    
    div[data-testid="stHorizontalBlock"] {
        gap: 0 !important;
        margin: 0 !important;
        padding: 0 !important;
    }
    
    /* Supprimer tous les espaces par défaut */
    .element-container {
        margin: 0 !important;
        padding: 0 !important;
    }
    
    /* Reset des polices */
    * {
        font-family: 'Outfit', sans-serif !important;
    }
    
    /* Styles des inputs */
    div[data-testid="stTextInput"] {
        margin-bottom: 0.3rem !important;
    }
    
    div[data-testid="stTextInput"] > div > div > input {
        background: rgba(22,27,46,0.9) !important;
        border: 1px solid rgba(227,0,27,0.3) !important;
        border-radius: 10px !important;
        color: #F1F5F9 !important;
        padding: 0.7rem 1rem !important;
        font-size: 0.85rem !important;
    }
    
    div[data-testid="stTextInput"] > div > div > input:focus {
        border-color: #C1121F !important;
        box-shadow: 0 0 0 2px rgba(227,0,27,0.2) !important;
        outline: none !important;
    }
    
    div[data-testid="stTextInput"] > div > div > input::placeholder {
        color: #64748B !important;
        font-size: 0.8rem !important;
    }
    
    div[data-testid="stButton"] > button {
        background: #C1121F !important;
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        padding: 0.7rem 1.5rem !important;
        font-size: 0.9rem !important;
        font-weight: 700 !important;
        width: 100% !important;
        transition: all 0.2s ease !important;
    }
    
    div[data-testid="stButton"] > button:hover {
        background: #8B0D14 !important;
        transform: translateY(-1px) !important;
        box-shadow: 0 4px 12px rgba(227,0,27,0.4) !important;
    }
    
    .stAlert {
        background: rgba(227,0,27,0.1) !important;
        border-left: 3px solid #C1121F !important;
        border-radius: 8px !important;
        padding: 0.5rem 1rem !important;
        margin-top: 1rem !important;
    }
    </style>
    """, unsafe_allow_html=True)

    # ==================== TOP BAR ====================
    st.markdown(f"""
    <div style="background: rgba(0,0,0,0.6); border-bottom: 3px solid {SI_RED}; padding: 0.6rem 2rem; display: flex; align-items: center; gap: 14px; backdrop-filter: blur(12px); margin-bottom: 0;">
        <div style="width: 38px; height: 38px; background: {SI_RED}; border-radius: 8px; display: flex; align-items: center; justify-content: center; font-size: 1.2rem; font-weight: 900; color: white; flex-shrink: 0;">SI</div>
        <div>
            <div style="font-size: 0.9rem; font-weight: 900; color: white;">INTERSOS</div>
            <div style="font-size: 0.55rem; color: rgba(255,255,255,0.4); letter-spacing: 0.1em; text-transform: uppercase;">Mission Tchad · Information Management Platform</div>
        </div>
        <div style="margin-left: auto; background: {SI_RED}; color: white; font-size: 0.6rem; font-weight: 800; letter-spacing: 0.1em; text-transform: uppercase; padding: 4px 14px; border-radius: 20px;">Restricted Access</div>
    </div>
    """, unsafe_allow_html=True)

    # ==================== IMAGE SLIDER ====================
    imgs_html = "".join(f"<img src='{u}' alt='SI Tchad' style='height: 110px; width: 170px; object-fit: cover;'>" for u in SI_IMAGES * 3)
    st.markdown(f"""
    <div style="overflow: hidden; margin: 0; height: 110px;">
        <div style="display: flex; gap: 12px; animation: slideLogin 40s linear infinite; width: max-content; height: 110px;">
            {imgs_html}
        </div>
    </div>
    <style>
    @keyframes slideLogin {{
        0% {{ transform: translateX(0); }}
        100% {{ transform: translateX(-33.33%); }}
    }}
    </style>
    """, unsafe_allow_html=True)

    # ==================== MAIN CONTENT ====================
    # Utiliser un conteneur avec du padding
    st.markdown("<div style='padding: 2rem 2rem 2rem 2rem;'></div>", unsafe_allow_html=True)

    # Créer les colonnes avec un conteneur personnalisé
    col1, col2 = st.columns([1.2, 0.9], gap="large")

    # ==================== COLONNE GAUCHE ====================
    with col1:
        st.markdown(f"""
        <div>
            <div style='display: inline-block; background: {SI_RED}; color: white; font-size: 0.6rem; font-weight: 800; letter-spacing: 0.12em; text-transform: uppercase; padding: 4px 12px; border-radius: 4px; margin-bottom: 1.2rem;'>Mission Tchad 2025–2026</div>
            <h1 style='font-size: 2.2rem; font-weight: 900; color: white; margin: 0 0 0.8rem 0; line-height: 1.2;'>Information<br><span style='color: {SI_RED};'>Management</span><br>Dashboard</h1>
            <p style='font-size: 0.85rem; color: rgba(255,255,255,0.55); line-height: 1.6; margin: 0 0 1.8rem 0;'>Real-time monitoring of réponse humanitaire multi-secteur — Protection, Food Security, Shelter & NFI, and Cash & Voucher programs in Tchad.</p>
            <div style='display: flex; flex-direction: column; gap: 0.8rem;'>
                <div style='display: flex; align-items: center; gap: 12px;'><span style='background: rgba(59,130,246,0.15); padding: 6px 10px; border-radius: 8px;'>💧</span><span><strong style='color: #3B82F6;'>WASH</strong> — Water, Sanitation & Hygiene</span></div>
                <div style='display: flex; align-items: center; gap: 12px;'><span style='background: rgba(16,185,129,0.15); padding: 6px 10px; border-radius: 8px;'>🌾</span><span><strong style='color: #10B981;'>FSL</strong> — Food Security & Livelihoods</span></div>
                <div style='display: flex; align-items: center; gap: 12px;'><span style='background: rgba(245,158,11,0.15); padding: 6px 10px; border-radius: 8px;'>🏠</span><span><strong style='color: #F59E0B;'>Shelter & NFI</strong> — Non-Food Items</span></div>
                <div style='display: flex; align-items: center; gap: 12px;'><span style='background: rgba(227,0,27,0.15); padding: 6px 10px; border-radius: 8px;'>💵</span><span><strong style='color: {SI_RED}'>CVA</strong> — Cash & Voucher Assistance</span></div>
            </div>
        </div>
        """, unsafe_allow_html=True)

    # ==================== COLONNE DROITE ====================
    with col2:
        # Carte de login
        st.markdown(f"""
        <div style='background: rgba(28,33,48,0.95); border: 1px solid rgba(227,0,27,0.3); border-top: 4px solid {SI_RED}; border-radius: 16px; padding: 2rem 1.8rem; box-shadow: 0 25px 50px rgba(0,0,0,0.4);'>
            <div style='font-size: 1.4rem; font-weight: 900; color: white; margin-bottom: 0.3rem;'>Welcome back</div>
            <div style='font-size: 0.8rem; color: #94A3B8; margin-bottom: 1.8rem;'>Enter your credentials to access the dashboard</div>
        </div>
        """, unsafe_allow_html=True)
        
        # Espacement
        st.markdown("<div style='margin-top: 2rem;'></div>", unsafe_allow_html=True)
        
        # Username
        st.markdown("<div style='font-size: 0.7rem; font-weight: 700; color: #94A3B8; margin-bottom: 1rem;'>Username</div>", unsafe_allow_html=True)
        user = st.text_input("", placeholder="im_manager", label_visibility="collapsed", key="login_user")
        
        # Password
        st.markdown("<div style='font-size: 0.7rem; font-weight: 700; color: #94A3B8; margin-bottom: 1rem; margin-top: 1rem;'>Password</div>", unsafe_allow_html=True)
        pw = st.text_input("", type="password", placeholder="••••••••", label_visibility="collapsed", key="login_pw")
        
        # Espacement
        st.markdown("<div style='margin: 1rem 0 0.5rem 0;'></div>", unsafe_allow_html=True)
        
        # Bouton
        if st.button("Sign in →", use_container_width=True, key="login_btn"):
            if CREDENTIALS.get(user) == pw:
                st.session_state.update(auth=True, user=user, dark=True, page="Overview")
                st.rerun()
            else:
                st.error("❌ Invalid username or password. Please try again.")
        
        # Footer
        st.markdown("""
        <div style='text-align: center; margin-top: 1.8rem; padding-top: 1.2rem; border-top: 1px solid rgba(255,255,255,0.05);'>
            <div style='font-size: 0.6rem; color: rgba(255,255,255,0.2);'>INTERSOS © 2026 · Confidential</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("</div>", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# MAP PAGE  (enriched)
# ══════════════════════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════════════════════
# SUDAN GEOGRAPHIC CONSTANTS + SHAPEFILE
# ══════════════════════════════════════════════════════════════════════════════
CHAD_LAT_MIN,  CHAD_LAT_MAX  = 7.44,  23.45
CHAD_LON_MIN,  CHAD_LON_MAX  = 13.47, 24.00
CHAD_CENTER  = [15.0, 18.5]

# Chad province centroids (verified)
STATE_CENTROIDS = {
    "Lac":               (13.30, 14.15),
    "Sila":              (12.10, 21.30),
    "Ouaddaï":           (13.85, 20.85),
    "Wadi Fira":         (15.20, 21.50),
    "N'Djamena":         (12.12, 15.04),
    "Logone Oriental":   ( 8.60, 16.10),
    "Mayo-Kebbi Est":    ( 9.30, 15.30),
    "Moyen-Chari":       ( 9.00, 18.20),
    "Batha":             (13.80, 18.40),
    "Hadjer-Lamis":      (12.50, 16.00),
    "Logone Occidental": ( 8.80, 15.50),
    "Tandjilé":          ( 9.50, 16.60),
    "Guéra":             (11.50, 18.40),
    "Salamat":           (10.60, 20.30),
    "Mandoul":           ( 8.60, 17.90),
    "Ennedi Est":        (17.00, 22.80),
    "Ennedi Ouest":      (16.50, 21.50),
    "Tibesti":           (21.00, 17.50),
    "Borkou":            (17.85, 18.53),
    "Kanem":             (14.50, 15.50),
    "Barh El Gazel":     (14.00, 16.70),
    "Chari-Baguirmi":    (12.60, 15.50),
}

import random as _rnd, json as _json, os as _os
_rnd.seed(42)

# Load Tchad shapefile GeoJSON
_GEOJSON_PATH = "/home/claude/chad_states.geojson"
try:
    with open(_GEOJSON_PATH) as _f:
        SUDAN_GEOJSON = _json.load(_f)
except Exception:
    SUDAN_GEOJSON = None

# Pre-computed scatter points per state (within ~0.4° of centroid)
STATE_SCATTER = {
    s: [(lat + _rnd.uniform(-0.4, 0.4), lon + _rnd.uniform(-0.5, 0.5))
        for _ in range(12)]
    for s, (lat, lon) in STATE_CENTROIDS.items()
}

def _get_coords(row):
    """Return (lat, lon) for a row — real GPS if valid in Tchad, else state scatter."""
    try:
        lat = float(row.get("GPS_Latitude", "") or "")
        lon = float(row.get("GPS_Longitude", "") or "")
        if CHAD_LAT_MIN <= lat <= CHAD_LAT_MAX and CHAD_LON_MIN <= lon <= CHAD_LON_MAX:
            return lat, lon
    except (TypeError, ValueError):
        pass
    state = str(row.get("Province", ""))
    if state in STATE_SCATTER:
        idx = hash(str(row.get("Beneficiary_ID", row.get("_index", 0)))) % len(STATE_SCATTER[state])
        return STATE_SCATTER[state][idx]
    return (_rnd.uniform(8.5, 22.5), _rnd.uniform(14.0, 23.5))

def _enrich_gps(df):
    if df.empty: return df
    df = df.copy().reset_index(drop=True)
    df["_index"] = df.index
    coords = df.apply(_get_coords, axis=1)
    df["GPS_Latitude"]  = coords.apply(lambda x: x[0])
    df["GPS_Longitude"] = coords.apply(lambda x: x[1])
    return df

def _clean_gps(df):
    return _enrich_gps(df) if not df.empty else pd.DataFrame()

def _base_map(dark=True, zoom=5):
    tiles = "CartoDB dark_matter" if dark else "CartoDB positron"
    m = folium.Map(location=CHAD_CENTER, zoom_start=zoom, tiles=tiles,
                   attr="© CartoDB", prefer_canvas=True)
    MiniMap(toggle_display=True, position="bottomleft", tile_layer=tiles,
            zoom_level_offset=-6).add_to(m)
    return m

def _add_geojson_layer(m, data_col=None, state_values=None, dark=True,
                        fill_color="#3B82F6", fill_opacity=0.15,
                        line_color="#C1121F", line_weight=1.5):
    """Overlay Tchad states from GeoJSON with optional choropleth colouring."""
    if SUDAN_GEOJSON is None:
        return
    border_col = "rgba(227,0,27,0.7)" if dark else "#8B0D14"

    if state_values:
        # Choropleth mode
        folium.Choropleth(
            geo_data=SUDAN_GEOJSON,
            data=state_values,
            columns=["state","value"],
            key_on="feature.properties.name",
            fill_color="YlOrRd",
            fill_opacity=0.65,
            line_opacity=0.9,
            line_color=border_col,
            line_weight=line_weight,
            nan_fill_color="rgba(100,116,139,0.15)",
            nan_fill_opacity=0.3,
            legend_name=data_col or "Value",
            highlight=True,
        ).add_to(m)
    else:
        # Simple boundary overlay
        folium.GeoJson(
            SUDAN_GEOJSON,
            style_function=lambda f: {
                "fillColor":   fill_color,
                "color":       line_color,
                "weight":      line_weight,
                "fillOpacity": fill_opacity,
            },
            tooltip=folium.GeoJsonTooltip(
                fields=["name"],
                aliases=["Province:"],
                style="font-family:Outfit,sans-serif;font-size:12px;"
            )
        ).add_to(m)

def _add_state_labels(m, dark=True):
    bg  = "rgba(14,17,32,0.85)" if dark else "rgba(255,255,255,0.90)"
    tc  = "#f1f5f9" if dark else "#0f172a"
    bd  = "rgba(227,0,27,0.4)"
    for state, (lat, lon) in STATE_CENTROIDS.items():
        folium.Marker(
            location=[lat, lon],
            icon=folium.DivIcon(
                html=(f"<div style='background:{bg};color:{tc};padding:2px 7px;"
                      f"border-radius:8px;font-family:Outfit,sans-serif;font-size:9.5px;"
                      f"font-weight:700;white-space:nowrap;border:1px solid {bd};"
                      f"box-shadow:0 2px 8px rgba(0,0,0,0.35);'>{state}</div>"),
                icon_size=(150, 20), icon_anchor=(75, 10)
            )
        ).add_to(m)

def _map_legend(m, items, title="Legend", dark=True):
    bg = "rgba(14,17,32,0.93)" if dark else "rgba(255,255,255,0.93)"
    tc = "#f1f5f9" if dark else "#1e293b"
    sc = "#94a3b8" if dark else "#64748b"
    html = (f"<div style='position:fixed;bottom:32px;right:16px;z-index:9999;"
            f"background:{bg};border:1px solid rgba(227,0,27,0.25);border-radius:10px;"
            f"padding:11px 15px;font-family:Outfit,sans-serif;min-width:145px;'>"
            f"<div style='font-weight:800;font-size:9.5px;color:{tc};margin-bottom:7px;"
            f"letter-spacing:.09em;text-transform:uppercase;'>{title}</div>")
    for label, color in items:
        html += (f"<div style='display:flex;align-items:center;gap:7px;margin-bottom:4px;'>"
                 f"<span style='width:10px;height:10px;border-radius:50%;background:{color};"
                 f"display:inline-block;flex-shrink:0;border:1px solid rgba(255,255,255,0.2);'></span>"
                 f"<span style='font-size:10px;color:{sc};'>{label}</span></div>")
    html += "</div>"
    m.get_root().html.add_child(folium.Element(html))

# ── MAP 1: GeoJSON boundaries + beneficiary clusters + heatmap ───────────────
def map_beneficiaries(df, dark=True):
    dfm = _clean_gps(df)
    m   = _base_map(dark)
    _add_geojson_layer(m, dark=dark)          # Tchad state boundaries
    _add_state_labels(m, dark)

    if not dfm.empty:
        HeatMap(dfm[["GPS_Latitude","GPS_Longitude"]].values.tolist(),
                radius=18, blur=22, min_opacity=0.18, max_zoom=10).add_to(m)

        cl = MarkerCluster(options={"maxClusterRadius":55,
                                    "showCoverageOnHover":False,
                                    "spiderfyOnMaxZoom":True}).add_to(m)
        for _, row in dfm.iterrows():
            sec   = str(row.get("Sector","")) if has(df,"Sector") else ""
            color = SC.get(sec, "#64748B")
            pop   = (f"<div style='font-family:Outfit,sans-serif;font-size:12px;min-width:175px;'>"
                     f"<b>{row.get('Beneficiary_ID','—')}</b><br>"
                     f"State: {row.get('State','—')}<br>"
                     f"<b style='color:{color};'>{sec}</b><br>"
                     f"Displacement: {row.get('Displacement_Status','—')}<br>"
                     f"Vulnerability: {row.get('Vulnerability_Level','—')}</div>")
            folium.CircleMarker(
                location=[row["GPS_Latitude"], row["GPS_Longitude"]],
                radius=5, color=color, fill=True, fill_color=color,
                fill_opacity=0.82, weight=0.7,
                popup=folium.Popup(pop, max_width=240),
                tooltip=f"{row.get('Locality','—')} · {sec}"
            ).add_to(cl)

    _map_legend(m, list(SC.items()), "Sectors", dark)
    return m

# ── MAP 2: Choropleth — beneficiary density by state ─────────────────────────
def map_choropleth_density(df, dark=True):
    m = _base_map(dark)
    if has(df,"State") and not df.empty:
        counts = df.groupby("Province").size().reset_index(name="value")
        counts.columns = ["state","value"]
        _add_geojson_layer(m, data_col="Beneficiaries", state_values=counts,
                           dark=dark, line_weight=2.0)
    else:
        _add_geojson_layer(m, dark=dark)
    _add_state_labels(m, dark)
    return m

# ── MAP 3: Sector bubbles per state over shapefile ───────────────────────────
def map_sector_coverage(df, dark=True):
    m = _base_map(dark)
    _add_geojson_layer(m, dark=dark)
    _add_state_labels(m, dark)

    OFFSETS = [(0.0,0.0),(0.28,0.0),(-0.28,0.0),(0.0,0.32),(0.0,-0.32)]
    if has(df,"State") and has(df,"Sector"):
        for state, (lat, lon) in STATE_CENTROIDS.items():
            sub   = df[df["Province"]==state] if has(df,"State") else pd.DataFrame()
            total = len(sub)
            if total == 0: continue
            for i,(sec,color) in enumerate(SC.items()):
                n = slen(sub,"Sector",sec)
                if n == 0: continue
                pct    = n / max(total,1)
                radius = max(8, min(42, pct * 58))
                dlat,dlon = OFFSETS[i % len(OFFSETS)]
                mlat = max(CHAD_LAT_MIN+0.1, min(CHAD_LAT_MAX-0.1, lat+dlat))
                mlon = max(CHAD_LON_MIN+0.1, min(CHAD_LON_MAX-0.1, lon+dlon))
                folium.CircleMarker(
                    location=[mlat, mlon], radius=radius,
                    color=color, fill=True, fill_color=color,
                    fill_opacity=0.62, weight=1.5,
                    tooltip=f"{state} · {sec}: {n:,} ({pct:.0%})",
                    popup=folium.Popup(
                        f"<b>{state}</b><br><b style='color:{color};'>{sec}</b><br>"
                        f"{n:,} beneficiaries ({pct:.0%})", max_width=210)
                ).add_to(m)

    _map_legend(m, list(SC.items()), "Sectors", dark)
    return m

# ── MAP 4: Displacement status over shapefile ─────────────────────────────────
def map_displacement(df, dark=True):
    dfm = _clean_gps(df)
    m   = _base_map(dark)
    _add_geojson_layer(m, dark=dark)
    _add_state_labels(m, dark)

    DISP_COLORS = {"IDP":"#EF4444","Refugee":"#3B82F6",
                   "Host Community":"#10B981","Returnee":"#F59E0B"}
    if not dfm.empty and has(dfm,"Displacement_Status"):
        for disp,color in DISP_COLORS.items():
            sub = dfm[dfm["Displacement_Status"]==disp]
            if sub.empty: continue
            cl = MarkerCluster(name=disp,
                               options={"maxClusterRadius":40,"showCoverageOnHover":False}).add_to(m)
            for _,row in sub.iterrows():
                folium.CircleMarker(
                    location=[row["GPS_Latitude"],row["GPS_Longitude"]],
                    radius=5,color=color,fill=True,fill_color=color,
                    fill_opacity=0.82,weight=0.7,
                    tooltip=f"{row.get('Locality','—')} · {disp}",
                    popup=folium.Popup(
                        f"<b>{disp}</b><br>State: {row.get('State','—')}<br>"
                        f"Locality: {row.get('Locality','—')}",max_width=200)
                ).add_to(cl)

    folium.LayerControl(position="topright").add_to(m)
    _map_legend(m, list(DISP_COLORS.items()), "Displacement", dark)
    return m

# ── MAP 5: Vulnerability choropleth ──────────────────────────────────────────
def map_vulnerability(df, dark=True):
    dfm = _clean_gps(df)
    m   = _base_map(dark)

    VULN_COLORS = {"Extremely Vulnerable":"#EF4444",
                   "Vulnerable":"#F59E0B",
                   "Moderately Vulnerable":"#10B981"}
    VULN_WEIGHTS = {"Extremely Vulnerable":1.0,"Vulnerable":0.6,"Moderately Vulnerable":0.3}

    if not dfm.empty and has(dfm,"Vulnerability_Level"):
        for vl,color in VULN_COLORS.items():
            sub = dfm[dfm["Vulnerability_Level"]==vl]
            if sub.empty: continue
            w = VULN_WEIGHTS[vl]
            data = [[r["GPS_Latitude"],r["GPS_Longitude"],w] for _,r in sub.iterrows()]
            HeatMap(data,radius=16,blur=18,min_opacity=0.22,
                    gradient={"0.3":color+"55","0.6":color+"AA","1.0":color}).add_to(m)

        # State circles coloured by dominant vulnerability
        for state,(lat,lon) in STATE_CENTROIDS.items():
            sub = dfm[dfm["State"]==state] if has(dfm,"State") else pd.DataFrame()
            if sub.empty: continue
            dom = sub["Vulnerability_Level"].mode()
            if dom.empty: continue
            col = VULN_COLORS.get(dom.iloc[0],"#64748B")
            n   = len(sub)
            folium.CircleMarker(
                location=[lat,lon],
                radius=max(12,min(40,(n/max(len(dfm),1))**0.5*55)),
                color=col,fill=True,fill_color=col,fill_opacity=0.25,weight=2,
                tooltip=f"{state} — dominant: {dom.iloc[0]} ({n:,})"
            ).add_to(m)

    _add_geojson_layer(m,dark=dark,fill_opacity=0.05,line_weight=1.2)
    _add_state_labels(m,dark)
    _map_legend(m,list(VULN_COLORS.items()),"Vulnerability",dark)
    return m

def page_map(dfs):
    ph("Couverture Géographique", "Tchad shapefile maps · choropleth · clusters · heatmap · spatial analysis")
    df = dfs.get("Beneficiary_Registration", pd.DataFrame())
    if df.empty: st.info("Load the Excel database to view geographic data."); return
    th   = TH()
    dark = st.session_state.get("dark", True)

    # ── Filters ───────────────────────────────────────────────────────────────
    c1,c2,c3,c4 = st.columns(4)
    with c1: s1 = st.selectbox("State",        sopts(df,"Province"),               key="m_s")
    with c2: s2 = st.selectbox("Sector",       sopts(df,"Sector"),              key="m_sec")
    with c3: s3 = st.selectbox("Displacement", sopts(df,"Displacement_Status"), key="m_d")
    with c4: s4 = st.selectbox("Vulnerability",sopts(df,"Vulnerability_Level"), key="m_v")

    df_f = sfilt(sfilt(sfilt(sfilt(df.copy(),"Province",s1),"Sector",s2),
                       "Displacement_Status",s3),"Vulnerability_Level",s4)
    nb  = len(df_f)
    fem = slen(df_f,"Sex","Female")

    # ── KPIs ──────────────────────────────────────────────────────────────────
    c1,c2,c3,c4,c5,c6 = st.columns(6)
    c1.markdown(kpi("Beneficiaries", N(nb),              "in selection", "red",   "👥"), unsafe_allow_html=True)
    c2.markdown(kpi("States",        N(suniq(df_f,"Province")), "covered",  "blue",  "🗺️"), unsafe_allow_html=True)
    c3.markdown(kpi("Localities",    N(suniq(df_f,"Locality")),"zones",  "teal",  "📌"), unsafe_allow_html=True)
    c4.markdown(kpi("Female",        f"{100*fem/nb:.0f}%" if nb else "—","","purple","♀"), unsafe_allow_html=True)
    c5.markdown(kpi("IDP",           N(slen(df_f,"Displacement_Status","IDP")),"","amber","🏕️"), unsafe_allow_html=True)
    c6.markdown(kpi("Active",        N(slen(df_f,"Statut_Enregistrement","Active")),"","green","✅"), unsafe_allow_html=True)
    st.markdown("<br>", unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════════════════
    # MAP 1 — Tchad shapefile + beneficiary clusters + heatmap  (FULL WIDTH)
    # ════════════════════════════════════════════════════════════════════════
    sh("Map 1 — Beneficiary Distribution on Tchad Shapefile (Clusters + Heatmap)")
    col_m1, col_r1 = st.columns([2.5, 1])
    with col_m1:
        st.markdown(f"<div style='font-size:.78rem;color:{th['muted']};margin-bottom:.5rem;'>"
                    f"Tchad state boundaries from official GeoJSON · {nb:,} beneficiaries mapped "
                    f"· Click clusters to zoom · Click markers for details.</div>", unsafe_allow_html=True)
        st_folium(map_beneficiaries(df_f, dark), use_container_width=True,
                  height=460, returned_objects=[], key="map1")
    with col_r1:
        if has(df_f,"State") and nb > 0:
            by_s = df_f.groupby("Province").size().reset_index(name="n").sort_values("n")
            fig = px.bar(by_s, x="n", y="State", orientation="h", color="n",
                         color_continuous_scale=[th["panel"],"#C1121F"],
                         title="Beneficiaries by State")
            fig.update_layout(coloraxis_showscale=False)
            pc(T(fig, h=220, leg=False))
        if has(df_f,"Sector") and nb > 0:
            by_sec = df_f.groupby("Sector").size().reset_index(name="n")
            fig = px.pie(by_sec, values="n", names="Sector", hole=0.56,
                         color="Sector", color_discrete_map=SC, title="By Sector")
            fig.update_traces(textinfo="none",
                              marker=dict(line=dict(color=th["panel"], width=2)))
            pc(T(fig, h=220))

    st.markdown("<br>", unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════════════════
    # MAP 2 — Choropleth: beneficiary density by state
    # ════════════════════════════════════════════════════════════════════════
    sh("Map 2 — Choropleth: Beneficiary Density by State")
    col_m2, col_r2 = st.columns([2.5, 1])
    with col_m2:
        st.markdown(f"<div style='font-size:.78rem;color:{th['muted']};margin-bottom:.5rem;'>"
                    f"State shading intensity ∝ number of registered beneficiaries. "
                    f"Hover over each state for details.</div>", unsafe_allow_html=True)
        st_folium(map_choropleth_density(df_f, dark), use_container_width=True,
                  height=430, returned_objects=[], key="map2")
    with col_r2:
        if has(df_f,"State") and nb > 0:
            st_d = df_f.groupby("Province").size().reset_index(name="Count").sort_values("Count", ascending=False)
            txt_col = th["text"]; muted_col = th["muted"]
            st.markdown(f"<div style='font-size:.8rem;font-weight:700;color:{txt_col};margin-bottom:.5rem;'>Top States by Volume</div>", unsafe_allow_html=True)
            for _, row in st_d.head(8).iterrows():
                pct = row["Count"] / nb * 100
                st.markdown(
                    f"<div style='margin-bottom:6px;'>"
                    f"<div style='display:flex;justify-content:space-between;font-size:.77rem;"
                    f"color:{txt_col};margin-bottom:2px;'>"
                    f"<span>{row['State']}</span><span>{row['Count']:,}</span></div>"
                    f"<div style='background:rgba(128,128,128,0.12);border-radius:3px;height:5px;'>"
                    f"<div style='width:{pct:.1f}%;height:100%;background:{SI_RED};border-radius:3px;'></div></div>"
                    f"</div>",
                    unsafe_allow_html=True
                )

    st.markdown("<br>", unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════════════════
    # MAP 3 & 4 — Sector coverage + Displacement (side by side)
    # ════════════════════════════════════════════════════════════════════════
    sh("Map 3 & 4 — Sector Coverage · Displacement Status")
    col_m3, col_m4 = st.columns(2)
    with col_m3:
        st.markdown(f"<div style='font-size:.77rem;color:{th['muted']};margin-bottom:.4rem;'>"
                    f"Bubble size = sector share within each state · GeoJSON boundaries shown.</div>",
                    unsafe_allow_html=True)
        st_folium(map_sector_coverage(df_f, dark), use_container_width=True,
                  height=380, returned_objects=[], key="map3")
    with col_m4:
        st.markdown(f"<div style='font-size:.77rem;color:{th['muted']};margin-bottom:.4rem;'>"
                    f"Colour = displacement category (IDP/Refugee/Host/Returnee) · Clustered markers.</div>",
                    unsafe_allow_html=True)
        st_folium(map_displacement(df_f, dark), use_container_width=True,
                  height=380, returned_objects=[], key="map4")

    st.markdown("<br>", unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════════════════
    # MAP 5 — Vulnerability heatmap (full width)
    # ════════════════════════════════════════════════════════════════════════
    sh("Map 5 — Vulnerability Heatmap on Tchad Boundaries")
    col_m5, col_r5 = st.columns([2.5, 1])
    with col_m5:
        st.markdown(f"<div style='font-size:.77rem;color:{th['muted']};margin-bottom:.4rem;'>"
                    f"Red = extremely vulnerable · Amber = vulnerable · Green = moderately vulnerable.</div>",
                    unsafe_allow_html=True)
        st_folium(map_vulnerability(df_f, dark), use_container_width=True,
                  height=430, returned_objects=[], key="map5")
    with col_r5:
        if has(df_f,"Vulnerability_Level") and nb > 0:
            vc = df_f["Vulnerability_Level"].value_counts().reset_index()
            vc.columns = ["Level","Count"]
            cmap = {"Extremely Vulnerable":"#EF4444",
                    "Vulnerable":"#F59E0B",
                    "Moderately Vulnerable":"#10B981"}
            fig = px.bar(vc, x="Count", y="Level", orientation="h",
                         color="Level", color_discrete_map=cmap,
                         title="Vulnerability Breakdown")
            fig.update_layout(showlegend=False)
            pc(T(fig, h=220, leg=False))
        if has(df_f,"Displacement_Status") and nb > 0:
            dc = df_f["Displacement_Status"].value_counts().reset_index()
            dc.columns = ["Status","Count"]
            cmap2 = {"IDP":"#EF4444","Refugee":"#3B82F6",
                     "Host Community":"#10B981","Returnee":"#F59E0B"}
            fig = px.pie(dc, values="Count", names="Status", hole=0.55,
                         color="Status", color_discrete_map=cmap2,
                         title="Displacement Status")
            fig.update_traces(textinfo="none",
                              marker=dict(line=dict(color=th["panel"],width=2)))
            pc(T(fig, h=220))

    st.markdown("<br>", unsafe_allow_html=True)

    # ════════════════════════════════════════════════════════════════════════
    # SPATIAL ANALYSIS CHARTS
    # ════════════════════════════════════════════════════════════════════════
    sh("Spatial Analysis — Cross-Dimensional Charts")
    c1,c2,c3 = st.columns(3)
    with c1:
        if has(df_f,"State") and has(df_f,"Sector") and nb > 0:
            ct = df_f.groupby(["Province","Sector"]).size().reset_index(name="Count")
            fig = px.bar(ct, x="State", y="Count", color="Sector", barmode="stack",
                         color_discrete_map=SC, title="Beneficiaries by State & Sector")
            pc(T(fig, h=290))
    with c2:
        if has(df_f,"Age") and nb > 0:
            da = df_f.copy()
            da["Age"] = pd.to_numeric(da["Age"], errors="coerce")
            da = da.dropna(subset=["Age"])
            bins=[0,5,18,35,50,120]; labs=["0–4","5–17","18–34","35–49","50+"]
            da["Age_Group"] = pd.cut(da["Age"], bins=bins, labels=labs, right=False)
            ag = da.groupby("Age_Group", observed=True).size().reset_index(name="Count")
            fig = px.bar(ag, x="Age_Group", y="Count", color="Age_Group",
                         color_discrete_sequence=COLORS, title="Age Group Distribution")
            fig.update_layout(showlegend=False)
            pc(T(fig, h=290))
    with c3:
        if has(df_f,"Registration_Date") and nb > 0:
            dt = df_f.copy()
            dt["Registration_Date"] = pd.to_datetime(dt["Registration_Date"], errors="coerce")
            dt = dt.dropna(subset=["Registration_Date"])
            dt["Month"] = dt["Registration_Date"].dt.to_period("M").astype(str)
            trend = dt.groupby("Month").size().reset_index(name="Count")
            fig = px.line(trend, x="Month", y="Count", markers=True,
                          title="Registration Timeline")
            fig.update_traces(line_color=SI_RED, marker_color=SI_RED2, line_width=2)
            pc(T(fig, h=290))

    sh("Coverage Heatmap — State × Sector")
    if has(df_f,"State") and has(df_f,"Sector") and nb > 0:
        heat = df_f.groupby(["Province","Sector"]).size().reset_index(name="Count")
        fig = px.density_heatmap(heat, x="Sector", y="State", z="Count",
                                 color_continuous_scale="Reds",
                                 title="Beneficiary Density by State & Sector")
        fig.update_layout(coloraxis_colorbar=dict(tickfont=dict(color=th["fontc"])))
        pc(T(fig, h=320, leg=False))


# ══════════════════════════════════════════════════════════════════════════════
# OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
def page_overview(dfs):
    th = TH()
    df_b = dfs.get("Beneficiary_Registration", pd.DataFrame())
    df_w = dfs.get("WASH_Activities",           pd.DataFrame())
    df_f = dfs.get("Protection_Monitoring",          pd.DataFrame())
    df_c = dfs.get("Health_Activities",        pd.DataFrame())
    df_i = dfs.get("Indicator_Tracker",         pd.DataFrame())

    tot  = len(df_b); act = slen(df_b,"Statut_Enregistrement","Active")
    wr   = int(ssum(dfs.get("Protection_Monitoring",pd.DataFrame()),"Beneficiaires_Atteints"))
    fhh  = int(ssum(dfs.get("Health_Activities",pd.DataFrame()),"Consultations_Total"))
    paid = sfilt(df_c,"Statut","Paid")
    usd  = paid["Admissions"].sum() if has(paid,"Admissions") and len(paid)>0 else 0
    on_t = slen(df_i,"Status","En bonne voie"); ti = len(df_i)

    # ── SI HERO BANNER ────────────────────────────────────────────────────────
    imgs_html = "".join(f"<img src='{u}' alt='SI'>" for u in SI_IMAGES * 2)
    st.markdown(f"""
    <div class='si-hero'>
      <img src='{SI_IMAGES[0]}' alt='SI Tchad humanitarian response'>
      <div class='si-hero-content'>
        <div class='si-tag'>Mission Tchad · {datetime.now().strftime("%B %Y")}</div>
        <div class='si-hero-title'>Program Overview Dashboard</div>
        <div class='si-hero-sub'>
          Multi-sector humanitarian response ·
          {tot:,} beneficiaries registered ·
          {suniq(df_b,"Province")} states covered
        </div>
      </div>
      <div class='si-hero-bar'></div>
    </div>
    <div class='img-slider'>
      <div class='img-track'>{imgs_html}</div>
    </div>
    """, unsafe_allow_html=True)


    sh("Programme Scale — Last 90 days vs Previous 90 days")

    # Period comparison
    ben_cur, ben_prv  = _split_periods(df_b, "Registration_Date")
    wash_cur, wash_prv = _split_periods(df_w, "Report_Date")
    fsl_cur, fsl_prv  = _split_periods(df_f, "Distribution_Date")
    cva_cur, cva_prv  = _split_periods(df_c, "Transfer_Date")

    tot_cur  = len(ben_cur);  tot_prv  = len(ben_prv)
    wr_cur   = int(ssum(wash_cur,"Beneficiaires_Atteints"))
    wr_prv   = int(ssum(wash_prv,"Beneficiaires_Atteints"))
    fhh_cur  = int(ssum(fsl_cur,"Beneficiaires_Atteints"))
    fhh_prv  = int(ssum(fsl_prv,"Beneficiaires_Atteints"))
    paid_cur = sfilt(cva_cur,"Statut","Paid") if not cva_cur.empty else pd.DataFrame()
    paid_prv = sfilt(cva_prv,"Statut","Paid") if not cva_prv.empty else pd.DataFrame()
    usd_cur  = paid_cur["Admissions"].sum() if has(paid_cur,"Admissions") and len(paid_cur)>0 else 0
    usd_prv  = paid_prv["Admissions"].sum() if has(paid_prv,"Admissions") and len(paid_prv)>0 else 0

    c1,c2,c3,c4,c5,c6 = st.columns(6)
    c1.markdown(kpi_with_delta("Beneficiaries", N(tot), tot, tot_prv if tot_prv>0 else None,
        f"{act:,} active", "blue","👤"), unsafe_allow_html=True)
    c2.markdown(kpi_with_delta("Bénéf. Atteints", N(wr), wr, wr_prv if wr_prv>0 else None,
        "individuals","teal","💧"), unsafe_allow_html=True)
    c3.markdown(kpi_with_delta("Cas Protection", N(fhh), fhh, fhh_prv if fhh_prv>0 else None,
        "households","green","🌾"), unsafe_allow_html=True)
    c4.markdown(kpi_with_delta("Consultations", f"${N(usd)}", usd, usd_prv if usd_prv>0 else None,
        "USD disbursed","amber","💵"), unsafe_allow_html=True)
    c5.markdown(kpi("Provinces", str(suniq(df_b,"Province")), "couvertes","purple","🗺️"), unsafe_allow_html=True)
    c6.markdown(kpi("En bonne voie", f"{on_t}/{ti}", "indicators","green","📊",
        f"↑ {on_t/ti:.0%}" if ti else None,"up"), unsafe_allow_html=True)

    # Alert engine (from sidebar engine, shown inline on overview)
    active_alerts = run_alert_engine(dfs)
    critical_alerts = [a for a in active_alerts if a["level"] in ("critical","warning")]
    if critical_alerts:
        sh("⚡ Program Alerts")
        for a in critical_alerts:
            render_alert(a)


    sh("Beneficiary Profile")
    c1,c2,c3 = st.columns([1.1,1,1])
    with c1:
        if not df_b.empty and has(df_b,"Sector"):
            d = df_b.groupby("Sector").size().reset_index(name="n")
            fig = px.pie(d, values="n", names="Sector", hole=0.58,
                         color="Sector", color_discrete_map=SC, title="Beneficiaries by Sector")
            fig.update_traces(textinfo="percent", textfont_size=10,
                              marker=dict(line=dict(color=th["bg"],width=2)))
            pc(T(fig,h=300))
    with c2:
        if not df_b.empty and has(df_b,"Displacement_Status"):
            d = df_b["Displacement_Status"].value_counts().reset_index(); d.columns=["Status","Count"]
            fig = px.bar(d, x="Count", y="Status", orientation="h",
                         color="Count", color_continuous_scale=["#1E3A6E","#3B82F6"],
                         title="Displacement Status")
            fig.update_layout(coloraxis_showscale=False)
            pc(T(fig,h=300,leg=False))
    with c3:
        if not df_b.empty and has(df_b,"Sex") and has(df_b,"Vulnerability_Level"):
            ct = df_b.groupby(["Vulnerability_Level","Sex"]).size().reset_index(name="n")
            fig = px.bar(ct, x="Vulnerability_Level", y="n", color="Sex", barmode="stack",
                         color_discrete_map={"Female":"#EC4899","Male":"#3B82F6"},
                         title="Vulnerability × Sex")
            pc(T(fig,h=300))

    sh("Registration Trend")
    if not df_b.empty and has(df_b,"Registration_Date"):
        df_t = df_b.copy()
        df_t["Registration_Date"] = pd.to_datetime(df_t["Registration_Date"],errors="coerce")
        df_t = df_t.dropna(subset=["Registration_Date"])
        df_t["Month"] = df_t["Registration_Date"].dt.to_period("M").astype(str)
        if has(df_t,"Sector"):
            trend = df_t.groupby(["Month","Sector"]).size().reset_index(name="Count")
            fig = px.area(trend, x="Month", y="Count", color="Sector",
                          color_discrete_map=SC, title="Monthly Registrations by Sector")
            fig.update_traces(line_width=2)
            pc(T(fig,h=290))

    if not df_b.empty and has(df_b,"State") and has(df_b,"Sector"):
        sh("Coverage Heatmap — State × Sector")
        heat = df_b.groupby(["Province","Sector"]).size().reset_index(name="Count")
        fig = px.density_heatmap(heat, x="Sector", y="State", z="Count",
                                 color_continuous_scale="Blues",
                                 title="Beneficiaries by State × Sector")
        fig.update_layout(coloraxis_colorbar=dict(tickfont=dict(color=th["fontc"])))
        pc(T(fig,h=280,leg=False))

# ══════════════════════════════════════════════════════════════════════════════
# WASH
# ══════════════════════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════════════
# PAGE — PROTECTION
# ══════════════════════════════════════════════════════════════════════════════
def page_protection(dfs):
    ph("Protection", "Monitoring de la protection — VBG · ENA · PSS · Assistance Juridique")
    df = dfs.get("Protection_Monitoring", pd.DataFrame())
    if df.empty: st.info("Chargez la base de données INTERSOS pour voir les données protection."); return
    th = TH()

    # Filters
    c1,c2,c3 = st.columns(3)
    with c1: s1 = st.selectbox("Province", sopts(df,"Province"),          key="pr_p")
    with c2: s2 = st.selectbox("Type",     sopts(df,"Type_Intervention"), key="pr_t")
    with c3: s3 = st.selectbox("Donateur", sopts(df,"Donateur"),          key="pr_d")
    df_f = filter_df(df, Province=s1, Type_Intervention=s2, Donateur=s3)
    nb = len(df_f)

    # KPIs
    cible  = int(ssum(df_f,"Cible_Beneficiaires"))
    att    = int(ssum(df_f,"Beneficiaires_Atteints"))
    fem    = int(ssum(df_f,"Femmes_Atteintes"))
    vbg    = int(ssum(df_f,"Cas_VBG"))
    ena    = int(ssum(df_f,"Cas_ENA"))
    pss    = int(ssum(df_f,"Seances_PSS"))
    cov    = att/max(cible,1)

    c1,c2,c3,c4,c5,c6 = st.columns(6)
    c1.markdown(kpi("Bénéficiaires",N(att),      f"{cov:.0%} couverture","red",   "🛡️"),unsafe_allow_html=True)
    c2.markdown(kpi("Femmes Atteintes",N(fem),   f"{fem/max(att,1):.0%}","purple","♀"),  unsafe_allow_html=True)
    c3.markdown(kpi("Cas VBG",        N(vbg),    "pris en charge",       "amber", "⚠️"), unsafe_allow_html=True)
    c4.markdown(kpi("ENA",            N(ena),    "enfants non accomp.",  "blue",  "👶"), unsafe_allow_html=True)
    c5.markdown(kpi("Séances PSS",    N(pss),    "psychosocial",         "green", "🧠"), unsafe_allow_html=True)
    c6.markdown(kpi("Provinces",      N(suniq(df_f,"Province")),"couvertes","teal","🗺️"),unsafe_allow_html=True)

    sh("Couverture Protection par Province")
    c1,c2 = st.columns(2)
    with c1:
        if has(df_f,"Province") and nb>0:
            by_p = df_f.groupby("Province")[["Cible_Beneficiaires","Beneficiaires_Atteints"]].sum().reset_index()
            fig = px.bar(by_p,x="Province",y=["Cible_Beneficiaires","Beneficiaires_Atteints"],
                         barmode="group",color_discrete_map={
                             "Cible_Beneficiaires":"rgba(193,18,31,0.25)",
                             "Beneficiaires_Atteints":"#C1121F"},
                         labels={"value":"Bénéficiaires","variable":""})
            pc(T(fig,h=300))
    with c2:
        if has(df_f,"Type_Intervention") and nb>0:
            by_t = df_f.groupby("Type_Intervention")["Beneficiaires_Atteints"].sum().reset_index()
            by_t = by_t.sort_values("Beneficiaires_Atteints",ascending=True)
            fig = px.bar(by_t,x="Beneficiaires_Atteints",y="Type_Intervention",
                         orientation="h",color="Beneficiaires_Atteints",
                         color_continuous_scale=["#FEE2E2","#C1121F"])
            fig.update_layout(coloraxis_showscale=False)
            pc(T(fig,h=300,leg=False))

    sh("Cas Spéciaux — VBG & ENA par Province")
    c1,c2,c3 = st.columns(3)
    with c1:
        if has(df_f,"Province") and nb>0:
            vbg_p = df_f.groupby("Province")["Cas_VBG"].sum().reset_index().sort_values("Cas_VBG",ascending=False)
            fig = px.bar(vbg_p,x="Province",y="Cas_VBG",
                         color="Cas_VBG",color_continuous_scale=["#FEF3C7","#C1121F"],
                         title="Cas VBG par Province")
            fig.update_layout(coloraxis_showscale=False)
            pc(T(fig,h=260,leg=False))
    with c2:
        if has(df_f,"Province") and nb>0:
            ena_p = df_f.groupby("Province")["Cas_ENA"].sum().reset_index()
            fig = px.pie(ena_p,values="Cas_ENA",names="Province",hole=0.55,
                         color_discrete_sequence=COLORS,title="Répartition ENA")
            fig.update_traces(textinfo="none",marker=dict(line=dict(color=th["panel"],width=2)))
            pc(T(fig,h=260))
    with c3:
        if has(df_f,"Cas_Complexes") and nb>0:
            comp_p = df_f.groupby("Province")["Cas_Complexes"].sum().reset_index()
            fig = px.bar(comp_p,x="Province",y="Cas_Complexes",
                         color="Cas_Complexes",color_continuous_scale=["#EFF6FF","#3B82F6"],
                         title="Cas Complexes")
            fig.update_layout(coloraxis_showscale=False)
            pc(T(fig,h=260,leg=False))


# ══════════════════════════════════════════════════════════════════════════════
# PAGE — SANTÉ
# ══════════════════════════════════════════════════════════════════════════════
def page_sante(dfs):
    ph("Santé Primaire", "Consultations · Santé Reproductive · Vaccination · Références")
    df = dfs.get("Health_Activities", pd.DataFrame())
    if df.empty: st.info("Chargez la base de données INTERSOS pour voir les données santé."); return
    th = TH()

    c1,c2,c3 = st.columns(3)
    with c1: s1 = st.selectbox("Province",    sopts(df,"Province"),    key="sa_p")
    with c2: s2 = st.selectbox("Activité",    sopts(df,"Type_Activite"),key="sa_a")
    with c3: s3 = st.selectbox("Fonctionnalité",sopts(df,"Fonctionnalite"),key="sa_f")
    df_f = filter_df(df, Province=s1, Type_Activite=s2, Fonctionnalite=s3)
    nb = len(df_f)

    cons  = int(ssum(df_f,"Consultations_Total"))
    cpn   = int(ssum(df_f,"Consultations_CPN"))
    vacc  = int(ssum(df_f,"Vaccinations"))
    ref_  = int(ssum(df_f,"Referrals"))
    mal   = int(ssum(df_f,"Cas_Malaria"))
    acc   = int(ssum(df_f,"Accouchements_Assistes"))

    c1,c2,c3,c4,c5,c6 = st.columns(6)
    c1.markdown(kpi("Consultations",   N(cons), "ambulatoires","blue",  "🏥"),unsafe_allow_html=True)
    c2.markdown(kpi("CPN",             N(cpn),  "soins prénataux","teal","🤰"),unsafe_allow_html=True)
    c3.markdown(kpi("Vaccinations",    N(vacc), "administrées",  "green","💉"),unsafe_allow_html=True)
    c4.markdown(kpi("Accouchements",   N(acc),  "assistés",      "purple","👶"),unsafe_allow_html=True)
    c5.markdown(kpi("Cas Malaria",     N(mal),  "traités",       "amber","🦟"),unsafe_allow_html=True)
    c6.markdown(kpi("Références",      N(ref_), "vers hôpitaux", "red",  "🚑"),unsafe_allow_html=True)

    sh("Performance Santé par Province et Activité")
    c1,c2 = st.columns(2)
    with c1:
        if has(df_f,"Province") and nb>0:
            by_p = df_f.groupby("Province")["Consultations_Total"].sum().reset_index().sort_values("Consultations_Total")
            fig = px.bar(by_p,x="Consultations_Total",y="Province",orientation="h",
                         color="Consultations_Total",color_continuous_scale=["#EFF6FF","#3B82F6"])
            fig.update_layout(coloraxis_showscale=False)
            pc(T(fig,h=300,leg=False))
    with c2:
        if has(df_f,"Type_Activite") and nb>0:
            by_a = df_f.groupby("Type_Activite")["Consultations_Total"].sum().reset_index()
            fig = px.pie(by_a,values="Consultations_Total",names="Type_Activite",
                         hole=0.55,color_discrete_sequence=COLORS)
            fig.update_traces(textinfo="none",marker=dict(line=dict(color=th["panel"],width=2)))
            pc(T(fig,h=300))

    sh("Couverture vs Cible Mensuelle")
    if has(df_f,"Taux_Couverture_Pct") and nb>0:
        df_cov = df_f[df_f["Taux_Couverture_Pct"].notna()].copy()
        df_cov["Taux_Pct"] = pd.to_numeric(df_cov["Taux_Couverture_Pct"],errors="coerce")*100
        df_cov = df_cov.groupby("Province")["Taux_Pct"].mean().reset_index()
        fig = px.bar(df_cov,x="Province",y="Taux_Pct",
                     color="Taux_Pct",
                     color_continuous_scale=["#FEE2E2","#FEF3C7","#D1FAE5"],
                     title="Taux de couverture moyen (%)")
        fig.add_hline(y=80,line_dash="dash",line_color="#C1121F",
                      annotation_text="Seuil 80%",annotation_position="top right")
        fig.update_layout(coloraxis_showscale=False)
        pc(T(fig,h=300,leg=False))


# ══════════════════════════════════════════════════════════════════════════════
# PAGE — NUTRITION
# ══════════════════════════════════════════════════════════════════════════════
def page_nutrition(dfs):
    ph("Nutrition", "CRENAS · CRENI · ANJE · Dépistage MUAC · Taux de Guérison")
    df = dfs.get("Nutrition_Program", pd.DataFrame())
    if df.empty:
        # Fallback: try to use Health sheet for some nutrition data
        df = dfs.get("Health_Activities", pd.DataFrame())
        if df.empty:
            st.info("Chargez la base de données INTERSOS pour voir les données nutrition.")
            return
    th = TH()

    c1,c2 = st.columns(2)
    with c1: s1 = st.selectbox("Province", sopts(df,"Province"), key="nu_p")
    with c2: s2 = st.selectbox("Programme",sopts(df,"Type_Programme") if has(df,"Type_Programme") else ["All"],key="nu_t")
    df_f = filter_df(df, Province=s1, **({"Type_Programme":s2} if has(df,"Type_Programme") and s2!="All" else {}))
    nb = len(df_f)

    adm   = int(ssum(df_f,"Admissions"))        if has(df_f,"Admissions")        else 0
    guer  = int(ssum(df_f,"Sorties_Gueris"))    if has(df_f,"Sorties_Gueris")    else 0
    dec_  = int(ssum(df_f,"Sorties_Deces"))     if has(df_f,"Sorties_Deces")     else 0
    aban  = int(ssum(df_f,"Sorties_Abandon"))   if has(df_f,"Sorties_Abandon")   else 0
    muac  = int(ssum(df_f,"MUAC_Rouge"))        if has(df_f,"MUAC_Rouge")        else 0
    enc   = int(ssum(df_f,"En_Cours_Traitement")) if has(df_f,"En_Cours_Traitement") else 0
    tg    = guer/max(adm,1)*100

    c1,c2,c3,c4,c5,c6 = st.columns(6)
    c1.markdown(kpi("Admissions",       N(adm),  "CRENAS/CRENI","red",   "📥"),unsafe_allow_html=True)
    c2.markdown(kpi("Guéris",           N(guer), f"{tg:.0f}%",  "green", "✅"),unsafe_allow_html=True)
    c3.markdown(kpi("Décès",            N(dec_), f"{dec_/max(adm,1):.1%}","amber","💔"),unsafe_allow_html=True)
    c4.markdown(kpi("Abandons",         N(aban), f"{aban/max(adm,1):.1%}","purple","🚪"),unsafe_allow_html=True)
    c5.markdown(kpi("MUAC Rouge",       N(muac), "MAS sévère",  "blue",  "📏"),unsafe_allow_html=True)
    c6.markdown(kpi("En Traitement",    N(enc),  "actifs",      "teal",  "🏥"),unsafe_allow_html=True)

    sh("Résultats Nutritionnels par Province")
    c1,c2 = st.columns(2)
    with c1:
        if has(df_f,"Province") and adm>0:
            by_p = df_f.groupby("Province")[["Admissions","Sorties_Gueris","Sorties_Deces"]].sum().reset_index()
            fig = px.bar(by_p,x="Province",
                         y=["Admissions","Sorties_Gueris","Sorties_Deces"],
                         barmode="group",
                         color_discrete_map={
                             "Admissions":"#3B82F6",
                             "Sorties_Gueris":"#10B981",
                             "Sorties_Deces":"#C1121F"})
            pc(T(fig,h=300))
    with c2:
        if has(df_f,"Taux_Guerison_Pct") and nb>0:
            tg_p = df_f.groupby("Province")["Taux_Guerison_Pct"].mean()*100
            tg_p = tg_p.reset_index(); tg_p.columns=["Province","Taux_Guerison"]
            fig = px.bar(tg_p,x="Province",y="Taux_Guerison",
                         color="Taux_Guerison",
                         color_continuous_scale=["#FEE2E2","#D1FAE5"])
            fig.add_hline(y=75,line_dash="dash",line_color="#C1121F",
                          annotation_text="Seuil SPHERE 75%")
            fig.update_layout(coloraxis_showscale=False)
            pc(T(fig,h=300,leg=False))


def page_wash(dfs):
    ph("WASH Monitoring","Water, Sanitation & Hygiene — activity performance and gender disaggregation")
    df = dfs.get("WASH_Activities",pd.DataFrame())
    if df.empty: st.info("No WASH_Monitoring sheet found."); return
    th = TH()
    c1,c2 = st.columns(2)
    with c1: s1 = st.selectbox("State",         sopts(df,"Province"),        key="w_s")
    with c2: s2 = st.selectbox("Activity Type", sopts(df,"Activity_Type"),key="w_a")
    df_f = sfilt(sfilt(df.copy(),"Province",s1),"Activity_Type",s2)
    reached = int(ssum(df_f,"Beneficiaires_Atteints"))
    target  = int(ssum(df_f,"Target_Beneficiaries")) or 1
    pct = reached/target
    sh("Key Indicators")
    c1,c2,c3,c4,c5 = st.columns(5)
    dd="up" if pct>=0.8 else "mid" if pct>=0.6 else "down"
    c1.markdown(kpi("Reached",    N(reached),                     f"of {N(target)} target","blue",  "💧",f"{pct:.0%} coverage",dd),unsafe_allow_html=True)
    c2.markdown(kpi("Water",      N(int(ssum(df_f,"Water_Volume_Liters")))+" L","litres",  "teal",  "🚰"),unsafe_allow_html=True)
    c3.markdown(kpi("Latrines",   N(int(ssum(df_f,"Latrine_Units_Built"))),    "built",    "green", "🏗️"),unsafe_allow_html=True)
    c4.markdown(kpi("Hyg Kits",   N(int(ssum(df_f,"Hygiene_Kits_Dist"))),      "distrib.", "amber", "🧴"),unsafe_allow_html=True)
    c5.markdown(kpi("NFI Kits",   N(int(ssum(df_f,"NFI_Kits_Dist"))),          "distrib.", "purple","📦"),unsafe_allow_html=True)

    sh("Activity Performance")
    c1,c2 = st.columns(2)
    with c1:
        if has(df_f,"Activity_Type"):
            grp = df_f.groupby("Activity_Type")[["Target_Beneficiaries","Beneficiaires_Atteints"]].sum().reset_index()
            fig = go.Figure()
            fig.add_bar(name="Target", x=grp["Activity_Type"],y=grp["Target_Beneficiaries"],
                        marker_color="rgba(59,130,246,0.2)",marker_line_color="#3B82F6",marker_line_width=1)
            fig.add_bar(name="Reached",x=grp["Activity_Type"],y=grp["Beneficiaires_Atteints"],marker_color="#3B82F6")
            fig.update_layout(barmode="group",title="Target vs Reached")
            pc(T(fig,h=320))
    with c2:
        if has(df_f,"Functionality_Status"):
            func=df_f["Functionality_Status"].value_counts().reset_index(); func.columns=["Status","Count"]
            cmap={"Fully Functional":"#10B981","Partially Functional":"#F59E0B","Non-Functional":"#EF4444","Under Construction":"#8B5CF6"}
            fig=px.pie(func,values="Count",names="Status",hole=0.55,color="Status",color_discrete_map=cmap,title="Infrastructure Functionality")
            fig.update_traces(textinfo="percent+label",textfont_size=9,marker=dict(line=dict(color=th["bg"],width=2)))
            pc(T(fig,h=320))

    sh("Gender Disaggregation")
    if has(df_f,"Reached_Female") and has(df_f,"Reached_Male"):
        c1,c2,c3 = st.columns(3)
        tf=df_f["Reached_Female"].sum(); tm=df_f["Reached_Male"].sum()
        with c1:
            fig=go.Figure(go.Bar(x=["Female","Male"],y=[tf,tm],marker_color=["#EC4899","#3B82F6"],
                text=[N(tf),N(tm)],textposition="outside",textfont=dict(color=th["text"],size=11)))
            fig.update_layout(title="Sex Disaggregation — Total")
            pc(T(fig,h=280,leg=False))
        with c2:
            if has(df_f,"State"):
                sg=df_f.groupby("Province")[["Reached_Female","Reached_Male"]].sum().reset_index()
                fig=px.bar(sg,x="State",y=["Reached_Female","Reached_Male"],barmode="stack",
                    color_discrete_map={"Reached_Female":"#EC4899","Reached_Male":"#3B82F6"},title="Sex by State")
                pc(T(fig,h=280))
        with c3:
            if has(df_f,"Reached_Children_U5"):
                tc5=df_f["Reached_Children_U5"].sum()
                fig=go.Figure(go.Pie(labels=["Children U5","Adults"],values=[tc5,max(0,(tf+tm)-tc5)],hole=0.55,
                    marker=dict(colors=["#F59E0B","#3B82F6"],line=dict(color=th["bg"],width=2))))
                fig.update_layout(title="Children U5 vs Adults")
                pc(T(fig,h=280))

    sh("Reporting Compliance")
    if has(df_f,"Cluster_Reporting"):
        c1,c2 = st.columns(2)
        with c1:
            cr=df_f["Cluster_Reporting"].value_counts().reset_index(); cr.columns=["Status","Count"]
            fig=px.pie(cr,values="Count",names="Status",hole=0.55,color="Status",
                color_discrete_map={"Reported":"#10B981","Not reported":"#EF4444"},title="Cluster Reporting")
            fig.update_traces(marker=dict(line=dict(color=th["bg"],width=2)),textinfo="percent")
            pc(T(fig,h=270))
        with c2:
            if has(df_f,"State"):
                g2=df_f.groupby(["Province","Cluster_Reporting"]).size().reset_index(name="n")
                fig=px.bar(g2,x="State",y="n",color="Cluster_Reporting",
                    color_discrete_map={"Reported":"#10B981","Not reported":"#EF4444"},title="Reporting by State")
                pc(T(fig,h=270))

# ══════════════════════════════════════════════════════════════════════════════
# FSL
# ══════════════════════════════════════════════════════════════════════════════
def page_fsl_legacy(dfs):
    ph("Food Security & Livelihoods","Distribution tracking — commodities, pipeline, donor coverage")
    df=dfs.get("Protection_Monitoring",pd.DataFrame())
    if df.empty: st.info("No FSL_Distribution sheet found."); return
    th=TH()
    c1,c2,c3=st.columns(3)
    with c1: s1=st.selectbox("State",     sopts(df,"Province"),         key="f_s")
    with c2: s2=st.selectbox("Commodity", sopts(df,"Commodity_Type"),key="f_c")
    with c3: s3=st.selectbox("Donor",     sopts(df,"Donor"),         key="f_d")
    df_f=sfilt(sfilt(sfilt(df.copy(),"Province",s1),"Commodity_Type",s2),"Donor",s3)
    hht=int(ssum(df_f,"Cible_Beneficiaires")); hhr=int(ssum(dfs.get("Health_Activities",pd.DataFrame()),"Consultations_Total"))
    qp=ssum(df_f,"Quantity_Planned"); qd=ssum(df_f,"Quantity_Distributed")
    fem=int(ssum(df_f,"Female_HHH_Reached")); pb=slen(df_f,"Type_Programme","Pipeline break")
    cov=hhr/hht if hht else 0
    sh("Key Indicators")
    c1,c2,c3,c4,c5=st.columns(5)
    dd="up" if cov>=0.8 else "mid" if cov>=0.6 else "down"
    c1.markdown(kpi("HH Reached",  N(hhr), f"of {N(hht)} targeted","green", "🏠",f"{cov:.0%} coverage",dd),unsafe_allow_html=True)
    c2.markdown(kpi("Qty Dist.",   N(qd),  f"of {N(qp)} planned",  "blue",  "📦"),unsafe_allow_html=True)
    c3.markdown(kpi("Female HHH",  N(fem), f"{fem/hhr:.0%}" if hhr else "—","purple","♀"),unsafe_allow_html=True)
    c4.markdown(kpi("Pipeline ⚠️", str(pb),"supply breaks",        "red" if pb else "green","⚠️"),unsafe_allow_html=True)
    c5.markdown(kpi("Records",     N(len(df_f)),"in selection",    "teal", "📋"),unsafe_allow_html=True)

    sh("Distribution Analysis")
    c1,c2,c3=st.columns(3)
    with c1:
        if has(df_f,"Commodity_Type"):
            comm=df_f.groupby("Commodity_Type")["Beneficiaires_Atteints"].sum().reset_index()
            comm=comm.sort_values("Beneficiaires_Atteints",ascending=True).tail(10); comm.columns=["Commodity","Beneficiaires_Atteints"]
            fig=px.bar(comm,x="Beneficiaires_Atteints",y="Commodity",orientation="h",color="Beneficiaires_Atteints",
                color_continuous_scale=["#14532D","#34D399"],title="HH Reached by Commodity")
            fig.update_layout(coloraxis_showscale=False)
            pc(T(fig,h=350,leg=False))
    with c2:
        if has(df_f,"Type_Programme"):
            pipe=df_f["Type_Programme"].value_counts().reset_index(); pipe.columns=["Status","Count"]
            cmap={"In stock":"#10B981","Low stock":"#F59E0B","Pipeline break":"#EF4444","Awaiting delivery":"#8B5CF6"}
            fig=px.pie(pipe,values="Count",names="Status",hole=0.55,color="Status",color_discrete_map=cmap,title="Pipeline Status")
            fig.update_traces(marker=dict(line=dict(color=th["bg"],width=2)),textinfo="percent")
            pc(T(fig,h=350))
    with c3:
        if has(df_f,"Donor"):
            don=df_f.groupby("Donor")["Beneficiaires_Atteints"].sum().reset_index()
            fig=px.bar(don,x="Donor",y="Beneficiaires_Atteints",color="Donor",color_discrete_sequence=COLORS,
                title="HH Reached by Donor",text="Beneficiaires_Atteints")
            fig.update_traces(texttemplate="%{text:,.0f}",textposition="outside",textfont=dict(color=th["text"],size=9))
            fig.update_layout(showlegend=False)
            pc(T(fig,h=350,leg=False))

    sh("Distribution Timeline")
    if has(df_f,"Distribution_Date"):
        df_t=df_f.copy(); df_t["Distribution_Date"]=pd.to_datetime(df_t["Distribution_Date"],errors="coerce")
        df_t=df_t.dropna(subset=["Distribution_Date"]); df_t["Month"]=df_t["Distribution_Date"].dt.to_period("M").astype(str)
        if has(df_t,"Donor"):
            monthly=df_t.groupby(["Month","Donor"])["Beneficiaires_Atteints"].sum().reset_index()
            fig=px.area(monthly,x="Month",y="Beneficiaires_Atteints",color="Donor",title="Monthly HH Reached by Donor",color_discrete_sequence=COLORS)
            fig.update_traces(line_width=1.5); pc(T(fig,h=270))

    sh("Post-Distribution Monitoring")
    c1,c2=st.columns(2)
    with c1:
        if has(df_f,"Beneficiary_Satisfaction"):
            sat=df_f["Beneficiary_Satisfaction"].value_counts().reset_index(); sat.columns=["Level","Count"]
            cmap={"Above 80%":"#10B981","60–80%":"#F59E0B","Below 60%":"#EF4444","N/A":"#475569"}
            fig=px.bar(sat,x="Level",y="Count",color="Level",color_discrete_map=cmap,title="Beneficiary Satisfaction",text="Count")
            fig.update_traces(textposition="outside",textfont=dict(color=th["text"])); fig.update_layout(showlegend=False)
            pc(T(fig,h=270))
    with c2:
        if has(df_f,"Post_Dist_Monitor"):
            pdm=df_f["Post_Dist_Monitor"].value_counts().reset_index(); pdm.columns=["Status","Count"]
            fig=px.pie(pdm,values="Count",names="Status",hole=0.5,title="PDM Completion",color_discrete_sequence=COLORS)
            fig.update_traces(marker=dict(line=dict(color=th["bg"],width=2)),textinfo="percent")
            pc(T(fig,h=270))

# ══════════════════════════════════════════════════════════════════════════════
# CVA
# ══════════════════════════════════════════════════════════════════════════════
def page_cva_legacy(dfs):
    ph("Cash & Voucher Assistance","Transfer tracking — payment status, modalities, beneficiary profile")
    df=dfs.get("Health_Activities",pd.DataFrame())
    if df.empty: st.info("No CVA_Cash_Transfers sheet found."); return
    th=TH()
    c1,c2,c3=st.columns(3)
    with c1: s1=st.selectbox("State",         sopts(df,"Province"),        key="c_s")
    with c2: s2=st.selectbox("Transfer Type", sopts(df,"Transfer_Type"),key="c_t")
    with c3: s3=st.selectbox("Round",         sopts(df,"Round"),        key="c_r")
    df_f=sfilt(sfilt(sfilt(df.copy(),"Province",s1),"Transfer_Type",s2),"Round",s3)
    paid_df=sfilt(df_f,"Statut","Paid"); pend_df=sfilt(df_f,"Statut","Pending"); fail_df=sfilt(df_f,"Statut","Failed")
    usd=paid_df["Admissions"].sum() if has(paid_df,"Admissions") and len(paid_df)>0 else 0
    avg_v=paid_df["Admissions"].mean() if has(paid_df,"Admissions") and len(paid_df)>0 else 0
    fem_pct=slen(df_f,"Female_Headed_HH","Yes")/len(df_f) if len(df_f)>0 and has(df_f,"Female_Headed_HH") else 0
    sh("Key Indicators")
    c1,c2,c3,c4,c5,c6=st.columns(6)
    c1.markdown(kpi("Total Paid", f"${N(usd)}","USD",            "blue",  "💵"),unsafe_allow_html=True)
    c2.markdown(kpi("Paid",       N(len(paid_df)),"completed",   "green", "✅"),unsafe_allow_html=True)
    c3.markdown(kpi("Pending",    N(len(pend_df)),"awaiting",    "amber", "⏳"),unsafe_allow_html=True)
    c4.markdown(kpi("Failed",     N(len(fail_df)),"to check",    "red" if len(fail_df) else "green","❌"),unsafe_allow_html=True)
    c5.markdown(kpi("Avg Value",  f"${avg_v:.0f}","per HH",      "teal",  "📊"),unsafe_allow_html=True)
    c6.markdown(kpi("Female HHH", f"{fem_pct:.0%}","",           "purple","♀"),unsafe_allow_html=True)
    if len(fail_df)>0:
        st.markdown(f"<div class='alert-banner'><span>❌ <b>{len(fail_df)}</b> failed transfers — verification required before next round.</span></div>",unsafe_allow_html=True)
    sh("Transfer Analysis")
    c1,c2=st.columns(2)
    with c1:
        if has(df_f,"Statut"):
            sc=df_f["Statut"].value_counts().reset_index(); sc.columns=["Status","Count"]
            cmap={"Paid":"#10B981","Pending":"#F59E0B","Failed":"#EF4444","Cancelled":"#8B5CF6"}
            fig=px.pie(sc,values="Count",names="Status",hole=0.58,color="Status",color_discrete_map=cmap,title="Transfer Status")
            fig.update_traces(marker=dict(line=dict(color=th["bg"],width=2)),textinfo="percent")
            pc(T(fig,h=310))
    with c2:
        if has(df_f,"Payment_Method") and has(df_f,"Admissions"):
            pm=df_f.groupby("Payment_Method")["Admissions"].sum().reset_index().sort_values("Admissions",ascending=True)
            fig=px.bar(pm,x="Admissions",y="Payment_Method",orientation="h",color="Admissions",
                color_continuous_scale=["#1E3A6E","#60A5FA"],title="USD by Payment Method",text="Admissions")
            fig.update_traces(texttemplate="$%{text:,.0f}",textposition="outside",textfont=dict(color=th["text"],size=9))
            fig.update_layout(coloraxis_showscale=False); pc(T(fig,h=310,leg=False))
    sh("Transfer Timeline")
    if has(df_f,"Transfer_Date") and has(df_f,"Admissions"):
        df_t=df_f.copy(); df_t["Transfer_Date"]=pd.to_datetime(df_t["Transfer_Date"],errors="coerce")
        df_t=df_t.dropna(subset=["Transfer_Date"]); df_t["Month"]=df_t["Transfer_Date"].dt.to_period("M").astype(str)
        if has(df_t,"Statut"):
            agg=df_t.groupby(["Month","Statut"])["Admissions"].sum().reset_index()
            cmap={"Paid":"#10B981","Pending":"#F59E0B","Failed":"#EF4444","Cancelled":"#8B5CF6"}
            fig=px.bar(agg,x="Month",y="Admissions",color="Statut",color_discrete_map=cmap,title="Monthly Volume (USD)")
            pc(T(fig,h=270))
    sh("Beneficiary Profile")
    c1,c2=st.columns(2)
    with c1:
        if has(df_f,"Transfer_Type") and has(df_f,"Admissions"):
            tt=df_f.groupby("Transfer_Type")["Admissions"].sum().reset_index(); tt.columns=["Type","Total"]
            fig=px.bar(tt,x="Type",y="Total",color="Type",color_discrete_sequence=COLORS,title="USD by Transfer Type",text="Total")
            fig.update_traces(texttemplate="%{text:,.0f}",textposition="outside",textfont=dict(color=th["text"],size=9)); fig.update_layout(showlegend=False)
            pc(T(fig,h=270))
    with c2:
        if has(df_f,"State") and has(df_f,"Admissions"):
            st_g=df_f.groupby("Province")["Admissions"].sum().reset_index()
            fig=px.pie(st_g,values="Admissions",names="State",hole=0.5,title="Value by State",color_discrete_sequence=COLORS)
            fig.update_traces(marker=dict(line=dict(color=th["bg"],width=2)),textinfo="percent")
            pc(T(fig,h=270))

# ══════════════════════════════════════════════════════════════════════════════
# INDICATORS
# ══════════════════════════════════════════════════════════════════════════════
def page_ind(dfs):
    ph("Indicator Tracker","Results-based management — progress vs annual targets by sector")
    df=dfs.get("Indicator_Tracker",pd.DataFrame())
    if df.empty: st.info("No Indicator_Tracker sheet found."); return
    th=TH()
    on_t=slen(df,"Status","En bonne voie"); at_r=slen(df,"Status","À risque"); off=slen(df,"Status","Hors trajectoire"); tot=len(df)
    c1,c2,c3,c4=st.columns(4)
    c1.markdown(kpi("Total",    str(tot),f"indicators",       "blue",  "📋"),unsafe_allow_html=True)
    c2.markdown(kpi("En bonne voie", str(on_t),f"{on_t/tot:.0%}", "green", "✅",f"↑ {on_t/tot:.0%}","up"),unsafe_allow_html=True)
    c3.markdown(kpi("At Risk",  str(at_r),f"{at_r/tot:.0%}", "amber", "⚠️"),unsafe_allow_html=True)
    c4.markdown(kpi("Off Track",str(off), f"{off/tot:.0%}",  "red",   "❌"),unsafe_allow_html=True)
    sh("Performance Overview")
    c1,c2=st.columns([1,2])
    with c1:
        fig=go.Figure(go.Pie(labels=["En bonne voie","At Risk","Off Track"],values=[on_t,at_r,off],hole=0.62,
            marker=dict(colors=["#10B981","#F59E0B","#EF4444"],line=dict(color=th["bg"],width=3)),textinfo="none"))
        fig.add_annotation(text=f"<b>{tot}</b>",x=0.5,y=0.5,font_size=20,showarrow=False,font_color=th["text"])
        fig.update_layout(title="Overall Status")
        pc(T(fig,h=280))
    with c2:
        if has(df,"Sector") and has(df,"Status"):
            perf=df.groupby(["Sector","Status"]).size().reset_index(name="Count")
            cmap={"En bonne voie":"#10B981","À risque":"#F59E0B","Hors trajectoire":"#EF4444"}
            fig=px.bar(perf,x="Sector",y="Count",color="Status",color_discrete_map=cmap,barmode="stack",title="Status by Sector")
            pc(T(fig,h=280))
    if has(df,"Sector"):
        for sec in sorted(df["Sector"].dropna().unique()):
            df_sec=df[df["Sector"]==sec]; sh(f"Sector — {sec}")
            for _,row in df_sec.iterrows():
                ind=str(row.get("Indicateur","")); unit=str(row.get("Unit",""))
                target=row.get("Cible_Annuelle",None); cumul=row.get("Cumul",None)
                q1=float(row.get("Q1 Achieved",0) or 0); q2=float(row.get("Q2 Achieved",0) or 0)
                q3=float(row.get("Q3 Achieved",0) or 0); q4=float(row.get("Q4 (Partial)",0) or 0)
                status=str(row.get("Status",""))
                has_t=pd.notna(target) and str(target) not in ["","nan","None"]
                if has_t:
                    try:
                        t=float(str(target).replace(",","").replace("%","")); cv=float(str(cumul).replace(",","").replace("%","")) if pd.notna(cumul) else 0
                        pct=min(cv/t,1.0) if t>0 else 0
                    except: t,cv,pct=0,0,0
                    bc="#10B981" if pct>=0.8 else "#F59E0B" if pct>=0.6 else "#EF4444"
                    st.markdown(f"""<div class='ind-card'>
                      <div style='display:flex;justify-content:space-between;align-items:flex-start;margin-bottom:7px;'>
                        <div><div class='ind-name'>{ind}</div><div class='ind-vals'>Target: {N(t)} {unit} · Achieved: {N(cv)} · {pct*100:.1f}%</div></div>
                        <div>{bdg(status)}</div>
                      </div>
                      <div class='pbar-wrap'><div class='pbar' style='width:{pct*100:.1f}%;background:{bc};'></div></div>
                      <div style='display:flex;gap:16px;margin-top:5px;'>
                        <span class='ind-vals'>Q1: {N(q1)}</span><span class='ind-vals'>Q2: {N(q2)}</span>
                        <span class='ind-vals'>Q3: {N(q3)}</span><span class='ind-vals'>Q4: {N(q4)}</span>
                        <span class='ind-vals' style='margin-left:auto;color:#60A5FA;'>Cumul: {N(q1+q2+q3+q4)}</span>
                      </div>
                    </div>""",unsafe_allow_html=True)
                else:
                    st.markdown(f"""<div class='ind-card' style='display:flex;justify-content:space-between;align-items:center;'>
                      <div><div class='ind-name'>{ind}</div><div class='ind-vals'>Count — no annual target</div></div>
                      <div style='display:flex;align-items:center;gap:10px;'>
                        <span style='font-size:1.1rem;font-weight:700;color:{th["text"]};'>{N(cumul) if pd.notna(cumul) else "—"} <span style='font-size:.7rem;color:#64748B;'>{unit}</span></span>
                        {bdg(status)}
                      </div>
                    </div>""",unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# RAW DATA
# ══════════════════════════════════════════════════════════════════════════════
def page_raw(dfs):
    ph("Raw Data Explorer","Browse, filter and export program datasets")
    sheet=st.selectbox("Select dataset",list(dfs.keys()))
    df=dfs[sheet]
    c1,c2=st.columns([3,1])
    with c1: st.markdown(f"<div style='padding:.5rem 0;font-size:.8rem;color:#64748B;'>{len(df):,} rows · {len(df.columns)} columns</div>",unsafe_allow_html=True)
    with c2: st.download_button("⬇ Download CSV",df.to_csv(index=False).encode(),file_name=f"{sheet}.csv",mime="text/csv",use_container_width=True)
    num_cols=df.select_dtypes(include="number").columns.tolist()
    if num_cols:
        sh("Quick Statistics")
        st.dataframe(df[num_cols].describe().round(2),use_container_width=True)
    sh("Data Table")
    st.dataframe(df,use_container_width=True,height=500)

# ══════════════════════════════════════════════════════════════════════════════
# AUTOMATIC REPORT GENERATOR
# ══════════════════════════════════════════════════════════════════════════════
def build_pdf_report(dfs):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm,
                            leftMargin=2.2*cm, rightMargin=2.2*cm)
    styles = getSampleStyleSheet()
    SI_BLUE = rl_colors.HexColor("#1A3A6B")
    SI_ACCENT = rl_colors.HexColor("#3B82F6")
    SI_LIGHT = rl_colors.HexColor("#EFF6FF")
    SI_GREEN = rl_colors.HexColor("#10B981")
    SI_AMBER = rl_colors.HexColor("#F59E0B")
    SI_RED   = rl_colors.HexColor("#EF4444")
    GRAY     = rl_colors.HexColor("#64748B")
    LIGHTGRAY= rl_colors.HexColor("#F8FAFC")

    H1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=18, textColor=SI_BLUE,
                          spaceAfter=6, fontName="Helvetica-Bold")
    H2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13, textColor=SI_BLUE,
                          spaceBefore=14, spaceAfter=4, fontName="Helvetica-Bold")
    H3 = ParagraphStyle("h3", parent=styles["Heading3"], fontSize=11, textColor=SI_ACCENT,
                          spaceBefore=10, spaceAfter=4, fontName="Helvetica-Bold")
    BODY = ParagraphStyle("body", parent=styles["Normal"], fontSize=10, textColor=rl_colors.HexColor("#1E293B"),
                           spaceAfter=6, leading=15, fontName="Helvetica")
    SMALL= ParagraphStyle("small",parent=styles["Normal"], fontSize=8.5, textColor=GRAY,
                           spaceAfter=4, leading=12, fontName="Helvetica")
    CENTER = ParagraphStyle("center",parent=BODY, alignment=TA_CENTER)
    RIGHT  = ParagraphStyle("right", parent=SMALL,alignment=TA_RIGHT)

    story = []
    now = datetime.now().strftime("%B %d, %Y")
    page_w = A4[0] - 4.2*cm

    def hr(): return HRFlowable(width="100%",thickness=1,color=rl_colors.HexColor("#E2E8F0"),spaceAfter=8,spaceBefore=4)
    def sp(n=8): return Spacer(1, n)

    def kpi_table(rows):
        col_w = page_w / len(rows)
        cells = []
        for label, value, sub in rows:
            cells.append([
                Paragraph(label.upper(), ParagraphStyle("kl",parent=SMALL,textColor=GRAY,fontSize=7.5,fontName="Helvetica-Bold",letterSpacing=0.5)),
                Paragraph(value, ParagraphStyle("kv",parent=styles["Normal"],fontSize=18,fontName="Helvetica-Bold",textColor=SI_BLUE)),
                Paragraph(sub, SMALL),
            ])
        tdata = [cells[i] for i in range(len(cells))]
        # Transpose: one row per column
        tdata_t = [[cells[i][j] for i in range(len(cells))] for j in range(3)]
        widths = [col_w]*len(rows)
        t = Table(tdata_t, colWidths=widths)
        t.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,-1),SI_LIGHT),
            ("ROWBACKGROUNDS",(0,0),(-1,-1),[SI_LIGHT,SI_LIGHT]),
            ("BOX",(0,0),(-1,-1),0.5,rl_colors.HexColor("#BFDBFE")),
            ("INNERGRID",(0,0),(-1,-1),0.3,rl_colors.HexColor("#DBEAFE")),
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
            ("TOPPADDING",(0,0),(-1,-1),8),
            ("BOTTOMPADDING",(0,0),(-1,-1),8),
            ("LEFTPADDING",(0,0),(-1,-1),10),
            ("RIGHTPADDING",(0,0),(-1,-1),6),
        ]))
        return t

    def data_table(headers, rows, col_widths=None):
        data = [headers] + rows
        cw = col_widths or [page_w/len(headers)]*len(headers)
        t = Table(data, colWidths=cw, repeatRows=1)
        ts = TableStyle([
            ("BACKGROUND",(0,0),(-1,0),SI_BLUE),
            ("TEXTCOLOR",(0,0),(-1,0),rl_colors.white),
            ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
            ("FONTSIZE",(0,0),(-1,0),9),
            ("ALIGN",(0,0),(-1,0),"CENTER"),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[LIGHTGRAY,rl_colors.white]),
            ("FONTSIZE",(0,1),(-1,-1),8.5),
            ("FONTNAME",(0,1),(-1,-1),"Helvetica"),
            ("ALIGN",(1,1),(-1,-1),"CENTER"),
            ("ALIGN",(0,1),(0,-1),"LEFT"),
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
            ("TOPPADDING",(0,0),(-1,-1),5),
            ("BOTTOMPADDING",(0,0),(-1,-1),5),
            ("LEFTPADDING",(0,0),(-1,-1),8),
            ("RIGHTPADDING",(0,0),(-1,-1),8),
            ("BOX",(0,0),(-1,-1),0.5,rl_colors.HexColor("#CBD5E1")),
            ("LINEBELOW",(0,0),(-1,0),1,rl_colors.HexColor("#1E40AF")),
        ])
        t.setStyle(ts)
        return t

    # ── COVER PAGE ──────────────────────────────────────────────────────────
    story.append(sp(60))
    story.append(Paragraph("INTERSOS", ParagraphStyle("cover_org",parent=styles["Normal"],
        fontSize=11,textColor=SI_ACCENT,fontName="Helvetica-Bold",alignment=TA_CENTER,letterSpacing=1.5)))
    story.append(sp(8))
    story.append(Paragraph("Mission Tchad", ParagraphStyle("cover_mission",parent=styles["Normal"],
        fontSize=28,textColor=SI_BLUE,fontName="Helvetica-Bold",alignment=TA_CENTER)))
    story.append(Paragraph("Rapport de Programme — Gestion de l'Information", ParagraphStyle("cover_sub",parent=styles["Normal"],
        fontSize=15,textColor=GRAY,fontName="Helvetica",alignment=TA_CENTER,spaceAfter=4)))
    story.append(sp(6))
    story.append(HRFlowable(width="60%",thickness=2,color=SI_ACCENT,hAlign="CENTER",spaceAfter=14,spaceBefore=4))
    story.append(Paragraph(f"Generated: {now}", ParagraphStyle("cover_date",parent=styles["Normal"],
        fontSize=10,textColor=GRAY,fontName="Helvetica",alignment=TA_CENTER)))
    story.append(Paragraph("Période de rapport : Janvier 2024 – Décembre 2025", ParagraphStyle("cover_period",parent=styles["Normal"],
        fontSize=10,textColor=GRAY,fontName="Helvetica",alignment=TA_CENTER)))
    story.append(sp(30))
    story.append(Paragraph("⚠ CONFIDENTIAL — For internal use only", ParagraphStyle("cover_conf",parent=styles["Normal"],
        fontSize=9,textColor=rl_colors.HexColor("#DC2626"),fontName="Helvetica-Bold",alignment=TA_CENTER)))
    story.append(PageBreak())

    # ── EXECUTIVE SUMMARY ───────────────────────────────────────────────────
    story.append(Paragraph("1. Executive Summary", H1))
    story.append(hr())
    df_b = dfs.get("Beneficiary_Registration",pd.DataFrame())
    df_w = dfs.get("WASH_Activities",pd.DataFrame())
    df_f = dfs.get("Protection_Monitoring",pd.DataFrame())
    df_c = dfs.get("Health_Activities",pd.DataFrame())
    df_i = dfs.get("Indicator_Tracker",pd.DataFrame())

    tot  = len(df_b)
    act  = slen(df_b,"Statut_Enregistrement","Active")
    wr   = int(ssum(dfs.get("Protection_Monitoring",pd.DataFrame()),"Beneficiaires_Atteints"))
    fhh  = int(ssum(dfs.get("Health_Activities",pd.DataFrame()),"Consultations_Total"))
    paid = sfilt(df_c,"Statut","Paid")
    usd  = paid["Admissions"].sum() if has(paid,"Admissions") and len(paid)>0 else 0
    on_t = slen(df_i,"Status","En bonne voie"); at_r=slen(df_i,"Status","À risque"); off_t=slen(df_i,"Status","Hors trajectoire"); ti=len(df_i)
    states_n = suniq(df_b,"Province")

    story.append(Paragraph(
        f"This report presents the consolidated program performance of the INTERSOS (SI) Mission Tchad "
        f"for the period January 2025 to April 2026. The mission is operational across <b>{states_n} states</b> of Tchad, "
        f"providing multi-sector humanitarian assistance in WASH, Food Security & Livelihoods (FSL), Shelter & NFI, "
        f"and Cash & Voucher Assistance (CVA).", BODY))
    story.append(Paragraph(
        f"As of the reporting date, the mission has registered a total of <b>{tot:,} beneficiaries</b>, "
        f"of whom <b>{act:,} ({act/tot:.0%})</b> are currently active in the program. "
        f"WASH activities have cumulatively reached <b>{wr:,} individuals</b>, while FSL distributions "
        f"have covered <b>{fhh:,} households</b>. Cash transfers totalling <b>${usd:,.0f}</b> have been disbursed "
        f"to eligible households.", BODY))
    story.append(sp(6))
    story.append(kpi_table([
        ("Total Beneficiaries", f"{tot:,}", f"{act:,} active ({act/tot:.0%})" if tot else ""),
        ("WASH Individuals", f"{wr:,}", "cumulative reached"),
        ("FSL Households", f"{fhh:,}", "all distributions"),
        ("Cash Disbursed", f"${usd:,.0f}", "USD paid out"),
        ("States Covered", str(states_n), "operational areas"),
    ]))
    story.append(sp(12))

    # Indicator status narrative
    story.append(Paragraph(
        f"Of the <b>{ti}</b> program indicators tracked in the results framework, <b>{on_t}</b> are currently "
        f"<b style='color:green;'>On Track</b> ({on_t/ti:.0%}), <b>{at_r}</b> are <b>At Risk</b> ({at_r/ti:.0%}), "
        f"and <b>{off_t}</b> are <b>Off Track</b> ({off_t/ti:.0%}). "
        f"Immediate attention is required for off-track indicators, particularly in sectors experiencing supply chain disruptions "
        f"or low field coverage.", BODY) if ti else Paragraph("No indicator data available.", BODY))

    story.append(PageBreak())

    # ── SECTOR: BENEFICIARIES ───────────────────────────────────────────────
    story.append(Paragraph("2. Beneficiary Registration", H1))
    story.append(hr())
    if not df_b.empty:
        story.append(Paragraph("2.1 Overview", H2))
        story.append(Paragraph(
            f"The beneficiary registry contains <b>{tot:,} records</b> across {suniq(df_b,'State')} states and "
            f"{suniq(df_b,'Locality')} localities. The dominant displacement categories are IDPs and host community members.",BODY))
        # Breakdown table
        if has(df_b,"Sector"):
            sec_cnt = df_b["Sector"].value_counts().reset_index(); sec_cnt.columns=["Sector","Count"]
            sec_cnt["% of Total"] = (sec_cnt["Count"]/tot*100).map(lambda x:f"{x:.1f}%")
            headers=["Sector","Beneficiaries","% of Total"]
            rows=[[s,f"{c:,}",p] for s,c,p in sec_cnt.values.tolist()]
            story.append(Paragraph("Table 1 — Beneficiaries by Sector",H3))
            story.append(data_table(headers,rows,[8*cm,4*cm,4*cm]))
            story.append(sp(8))
        if has(df_b,"Displacement_Status"):
            disp=df_b["Displacement_Status"].value_counts().reset_index(); disp.columns=["Status","Count"]
            disp["% of Total"]=(disp["Count"]/tot*100).map(lambda x:f"{x:.1f}%")
            story.append(Paragraph("Table 2 — Displacement Status Breakdown",H3))
            story.append(data_table(["Displacement Status","Count","% of Total"],
                [[s,f"{c:,}",p] for s,c,p in disp.values.tolist()],[7*cm,4*cm,4*cm]))
            story.append(sp(8))
        if has(df_b,"Sex"):
            fem=slen(df_b,"Sex","Female"); male=tot-fem
            story.append(Paragraph("2.2 Gender & Vulnerability", H2))
            story.append(Paragraph(f"Of all registered beneficiaries, <b>{fem:,} ({fem/tot:.0%})</b> are female and "
                                    f"<b>{male:,} ({male/tot:.0%})</b> are male. The program is aligned with SI's "
                                    f"gender mainstreaming policy, targeting female-headed households as a priority group.", BODY))

    story.append(PageBreak())

    # ── SECTOR: WASH ────────────────────────────────────────────────────────
    story.append(Paragraph("3. WASH Program Performance", H1))
    story.append(hr())
    if not df_w.empty:
        reached=int(ssum(dfs.get("Protection_Monitoring",pd.DataFrame()),"Beneficiaires_Atteints")); target=int(ssum(df_w,"Target_Beneficiaries")) or 1
        pct=reached/target; wv=int(ssum(df_w,"Water_Volume_Liters")); lat=int(ssum(df_w,"Latrine_Units_Built")); hyg=int(ssum(df_w,"Hygiene_Kits_Dist"))
        story.append(Paragraph(
            f"The WASH program has reached <b>{reached:,} individuals</b> against a target of <b>{target:,}</b>, "
            f"representing an overall coverage rate of <b>{pct:.0%}</b>. "
            f"Water distribution totalled <b>{wv:,} litres</b>, <b>{lat:,} latrine units</b> were constructed or rehabilitated, "
            f"and <b>{hyg:,} hygiene kits</b> were distributed to households in need.", BODY))
        story.append(kpi_table([
            ("Individuals Reached",f"{reached:,}",f"{pct:.0%} of target"),
            ("Water Distributed",f"{wv:,} L","litres"),
            ("Latrines Built",f"{lat:,}","units"),
            ("Hygiene Kits",f"{hyg:,}","distributed"),
        ]))
        story.append(sp(10))
        if has(df_w,"Activity_Type"):
            grp=df_w.groupby("Activity_Type")[["Target_Beneficiaries","Beneficiaires_Atteints"]].sum().reset_index()
            grp["Coverage"]=grp.apply(lambda r: f"{r['Reached_Beneficiaries']/r['Target_Beneficiaries']:.0%}" if r["Target_Beneficiaries"]>0 else "N/A",axis=1)
            story.append(Paragraph("Table 3 — WASH Performance by Activity Type",H3))
            rows=[[r["Activity_Type"],f"{int(r['Target_Beneficiaries']):,}",f"{int(r['Reached_Beneficiaries']):,}",r["Coverage"]] for _,r in grp.iterrows()]
            story.append(data_table(["Activity Type","Target","Reached","Coverage Rate"],rows,[6*cm,3*cm,3*cm,3*cm]))
    else:
        story.append(Paragraph("No WASH data available in the uploaded database.",BODY))

    story.append(PageBreak())

    # ── SECTOR: FSL ─────────────────────────────────────────────────────────
    story.append(Paragraph("4. Food Security & Livelihoods (FSL)", H1))
    story.append(hr())
    if not df_f.empty:
        hht=int(ssum(df_f,"Cible_Beneficiaires")); hhr=int(ssum(dfs.get("Health_Activities",pd.DataFrame()),"Consultations_Total"))
        qp=ssum(df_f,"Quantity_Planned"); qd=ssum(df_f,"Quantity_Distributed")
        fem_fsl=int(ssum(df_f,"Female_HHH_Reached")); pb=slen(df_f,"Type_Programme","Pipeline break")
        cov=hhr/hht if hht else 0
        story.append(Paragraph(
            f"FSL distributions have reached <b>{hhr:,} households</b> out of a planned target of <b>{hht:,}</b> "
            f"(coverage rate: <b>{cov:.0%}</b>). Female-headed households accounted for <b>{fem_fsl:,}</b> of beneficiary households "
            f"({fem_fsl/hhr:.0%} of reached). Total commodity volume distributed stands at <b>{qd:,.1f}</b> units "
            f"against a plan of <b>{qp:,.1f}</b> units.", BODY))
        if pb:
            story.append(Paragraph(f"⚠ Supply Chain Alert: <b>{pb} pipeline breaks</b> were recorded during the reporting period. "
                                    f"Program teams should coordinate with the logistics cluster to address supply disruptions.",BODY))
        story.append(kpi_table([
            ("HH Reached",f"{hhr:,}",f"{cov:.0%} of target"),
            ("Qty Distributed",f"{qd:,.0f}","units"),
            ("Female HHH",f"{fem_fsl:,}",f"{fem_fsl/hhr:.0%} of reached" if hhr else ""),
            ("Pipeline Breaks",str(pb),"supply alerts"),
        ]))
        story.append(sp(10))
        if has(df_f,"Commodity_Type"):
            comm=df_f.groupby("Commodity_Type")["Beneficiaires_Atteints"].sum().reset_index().sort_values("Beneficiaires_Atteints",ascending=False)
            story.append(Paragraph("Table 4 — HH Reached by Commodity Type",H3))
            rows=[[r["Commodity_Type"],f"{int(r['HH_Reached']):,}"] for _,r in comm.iterrows()]
            story.append(data_table(["Commodity","HH Reached"],rows,[9*cm,5*cm]))
    else:
        story.append(Paragraph("No FSL data available in the uploaded database.",BODY))

    story.append(PageBreak())

    # ── SECTOR: CVA ─────────────────────────────────────────────────────────
    story.append(Paragraph("5. Cash & Voucher Assistance (CVA)", H1))
    story.append(hr())
    if not df_c.empty:
        paid=sfilt(df_c,"Statut","Paid"); pend=sfilt(df_c,"Statut","Pending"); fail=sfilt(df_c,"Statut","Failed")
        usd=paid["Admissions"].sum() if has(paid,"Admissions") and len(paid)>0 else 0
        avg_v=paid["Admissions"].mean() if has(paid,"Admissions") and len(paid)>0 else 0
        fem_cva=slen(df_c,"Female_Headed_HH","Yes")
        story.append(Paragraph(
            f"A total of <b>{len(paid):,} cash transfers</b> have been successfully paid, disbursing <b>${usd:,.0f} USD</b> "
            f"to eligible households, with an average transfer value of <b>${avg_v:,.0f}</b> per household. "
            f"Female-headed households represent <b>{fem_cva:,}</b> beneficiaries "
            f"({fem_cva/len(df_c):.0%} of all transfers).", BODY))
        if len(fail)>0:
            story.append(Paragraph(f"⚠ Transfer Alert: <b>{len(fail)} transfers</b> failed and require investigation before the next payment cycle.",BODY))
        story.append(kpi_table([
            ("Total Paid",f"${usd:,.0f}","USD disbursed"),
            ("Paid Transfers",f"{len(paid):,}","completed"),
            ("Pending",f"{len(pend):,}","awaiting"),
            ("Failed",f"{len(fail):,}","to investigate"),
        ]))
        story.append(sp(10))
        if has(df_c,"Payment_Method") and has(df_c,"Admissions"):
            pm=df_c.groupby("Payment_Method")["Admissions"].sum().reset_index().sort_values("Admissions",ascending=False)
            story.append(Paragraph("Table 5 — USD Disbursed by Payment Method",H3))
            rows=[[r["Payment_Method"],f"${r['Transfer_Value_USD']:,.0f}"] for _,r in pm.iterrows()]
            story.append(data_table(["Payment Method","Total USD"],rows,[9*cm,5*cm]))
    else:
        story.append(Paragraph("No CVA data available.",BODY))

    story.append(PageBreak())

    # ── INDICATOR TRACKER ───────────────────────────────────────────────────
    story.append(Paragraph("6. Program Indicators — Results Framework", H1))
    story.append(hr())
    if not df_i.empty:
        story.append(Paragraph(
            f"The program results framework tracks <b>{ti}</b> indicators across four sectors. "
            f"As of the reporting date, <b>{on_t} indicators ({on_t/ti:.0%})</b> are on track, "
            f"<b>{at_r} ({at_r/ti:.0%})</b> are at risk, and <b>{off_t} ({off_t/ti:.0%})</b> are off track.", BODY))
        story.append(sp(6))
        # Full indicator table
        headers=["Indicateur","Unit","Cible_Annuelle","Cumul","% vs Target","Status"]
        rows=[]
        for _,row in df_i.iterrows():
            ind=str(row.get("Indicateur",""))[:55]; unit=str(row.get("Unit",""))
            targ=row.get("Cible_Annuelle",None); cumul=row.get("Cumul",None); status=str(row.get("Status",""))
            try:
                t=float(str(targ).replace(",","").replace("%","")) if pd.notna(targ) and str(targ) not in ["","nan"] else None
                cv=float(str(cumul).replace(",","").replace("%","")) if pd.notna(cumul) and str(cumul) not in ["","nan"] else 0
                pct_s=f"{cv/t:.0%}" if t and t>0 else "N/A"
                targ_s=f"{t:,.0f}" if t else "N/A"
                cumul_s=f"{cv:,.0f}"
            except: targ_s=str(targ); cumul_s=str(cumul); pct_s="N/A"
            rows.append([ind,unit,targ_s,cumul_s,pct_s,status])
        t_table=data_table(headers,rows,[5.2*cm,1.8*cm,2.2*cm,2.2*cm,2.2*cm,2*cm])
        story.append(t_table)
    else:
        story.append(Paragraph("No indicator data available.",BODY))

    story.append(PageBreak())

    # ── RECOMMENDATIONS ──────────────────────────────────────────────────────
    story.append(Paragraph("7. Recommendations", H1))
    story.append(hr())
    recs = [
        ("Data Quality & Completeness",
         "Ensure all field teams submit data within 48 hours of each activity. Implement automated daily quality checks through the dashboard to flag missing or inconsistent records before they reach program reports."),
        ("WASH Coverage",
         f"Coverage stands at {int(ssum(df_w,'Reached_Beneficiaries'))/max(int(ssum(df_w,'Target_Beneficiaries')),1):.0%} of target. Field teams in underperforming states should be prioritized for supervisory support and resource reallocation."),
        ("CVA Failed Transfers",
         f"{slen(df_c,'Transfer_Status','Failed')} transfers have failed and must be investigated. Coordinate with payment agents to identify root causes (account errors, connectivity issues) and reprocess eligible cases before the next payment round."),
        ("FSL Pipeline Monitoring",
         f"{slen(df_f,'Pipeline_Status','Pipeline break')} pipeline breaks were recorded. The logistics team should increase monitoring frequency and establish a minimum buffer stock threshold to prevent distribution gaps."),
        ("Off-Track Indicators",
         f"{slen(df_i,'Status','Off track')} indicators are off track. Program managers should review implementation plans and propose corrective measures in the next Monthly Program Review (MPR)."),
        ("MEAL Capacity",
         "Consider a refresher training for field data collectors on KoboToolbox protocols and GPS accuracy. High-quality georeferenced data is essential for accurate coverage mapping and cluster reporting compliance."),
    ]
    for i,(title,text) in enumerate(recs,1):
        story.append(Paragraph(f"7.{i} {title}",H3))
        story.append(Paragraph(text,BODY))

    story.append(PageBreak())

    # ── ANNEXES ──────────────────────────────────────────────────────────────
    story.append(Paragraph("Annex — Data Sources & Methodology", H1))
    story.append(hr())
    annex_data = [
        ["Dataset","Sheet Name","Records","Key Variables"],
        ["Beneficiary Registry","Beneficiary_Registration",f"{len(df_b):,}","ID, State, Sector, Sex, Age, GPS"],
        ["WASH Activities","WASH_Activities",f"{len(df_w):,}","Activity Type, Target, Reached, Functionality"],
        ["FSL Distributions","Protection_Monitoring",f"{len(df_f):,}","Commodity, HH Reached, Pipeline, Donor"],
        ["CVA Transfers","Health_Activities",f"{len(df_c):,}","Transfer Type, Value USD, Status, Method"],
        ["Indicator Tracker","Indicator_Tracker",f"{len(df_i):,}","Indicator, Target, Q1–Q4, Status"],
    ]
    t_annex=Table(annex_data,colWidths=[4.5*cm,4.5*cm,2.5*cm,5.1*cm])
    t_annex.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(-1,0),SI_BLUE),("TEXTCOLOR",(0,0),(-1,0),rl_colors.white),
        ("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),("FONTSIZE",(0,0),(-1,0),9),
        ("ROWBACKGROUNDS",(0,1),(-1,-1),[LIGHTGRAY,rl_colors.white]),
        ("FONTSIZE",(0,1),(-1,-1),8.5),("FONTNAME",(0,1),(-1,-1),"Helvetica"),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8),
        ("BOX",(0,0),(-1,-1),0.5,rl_colors.HexColor("#CBD5E1")),
    ]))
    story.append(t_annex)
    story.append(sp(16))
    story.append(Paragraph(
        "All data was collected through KoboToolbox electronic data collection forms and processed using Python/pandas. "
        "Geographic data was captured via GPS-enabled devices. "
        "Indicator progress is calculated against annual targets agreed with donors at project inception.",SMALL))
    story.append(sp(20))
    story.append(HRFlowable(width="100%",thickness=1,color=rl_colors.HexColor("#E2E8F0"),spaceAfter=8))
    story.append(Paragraph(f"INTERSOS — Mission Tchad IM Unit | Report generated {now} | CONFIDENTIAL",
        ParagraphStyle("footer",parent=styles["Normal"],fontSize=8,textColor=GRAY,alignment=TA_CENTER,fontName="Helvetica")))

    doc.build(story)
    return buf.getvalue()

def build_text_report(dfs):
    now = datetime.now().strftime("%B %d, %Y")
    df_b=dfs.get("Beneficiary_Registration",pd.DataFrame())
    df_w=dfs.get("WASH_Activities",pd.DataFrame())
    df_f=dfs.get("Protection_Monitoring",pd.DataFrame())
    df_c=dfs.get("Health_Activities",pd.DataFrame())
    df_i=dfs.get("Indicator_Tracker",pd.DataFrame())
    tot=len(df_b); act=slen(df_b,"Statut_Enregistrement","Active")
    wr=int(ssum(dfs.get("Protection_Monitoring",pd.DataFrame()),"Beneficiaires_Atteints")); fhh=int(ssum(dfs.get("Health_Activities",pd.DataFrame()),"Consultations_Total"))
    paid=sfilt(df_c,"Statut","Paid")
    usd=paid["Admissions"].sum() if has(paid,"Admissions") and len(paid)>0 else 0
    on_t=slen(df_i,"Status","En bonne voie"); ti=len(df_i)
    lines = [
        "="*70,
        "INTERSOS — SUDAN MISSION",
        "INFORMATION MANAGEMENT PROGRAM REPORT",
        f"Generated: {now}",
        "CONFIDENTIAL — FOR INTERNAL USE ONLY",
        "="*70, "",
        "1. EXECUTIVE SUMMARY",
        "-"*40,
        f"Total Beneficiaries:    {tot:,} ({act:,} active — {act/tot:.0%})" if tot else "No beneficiary data.",
        f"WASH Individuals:       {wr:,} reached",
        f"FSL Households:         {fhh:,} reached",
        f"Cash Disbursed:         ${usd:,.0f} USD",
        f"States Covered:         {suniq(df_b,'State')}",
        f"Indicators On Track:    {on_t}/{ti} ({on_t/ti:.0%})" if ti else "No indicator data.",
        "",
        "2. WASH", "-"*40,
        f"Reached: {wr:,} | Target: {int(ssum(df_w,'Target_Beneficiaries')):,} | Coverage: {wr/max(int(ssum(df_w,'Target_Beneficiaries')),1):.0%}",
        f"Water: {int(ssum(df_w,'Water_Volume_Liters')):,} L | Latrines: {int(ssum(df_w,'Latrine_Units_Built')):,} | Hygiene Kits: {int(ssum(df_w,'Hygiene_Kits_Dist')):,}",
        "",
        "3. FSL", "-"*40,
        f"HH Reached: {fhh:,} / {int(ssum(df_f,'HH_Targeted')):,} | Pipeline Breaks: {slen(df_f,'Pipeline_Status','Pipeline break')}",
        f"Female-headed HH: {int(ssum(df_f,'Female_HHH_Reached')):,}",
        "",
        "4. CVA", "-"*40,
        f"Paid: {len(paid):,} transfers | ${usd:,.0f} USD | Failed: {slen(df_c,'Transfer_Status','Failed')}",
        "",
        "5. INDICATORS", "-"*40,
        f"On Track: {on_t} | At Risk: {slen(df_i,'Status','At risk')} | Off Track: {slen(df_i,'Status','Off track')}",
        "",
        "="*70,
        "END OF REPORT",
        "="*70,
    ]
    return "\n".join(lines)

# ══════════════════════════════════════════════════════════════════════════════
# WORD REPORT BUILDER  (python-docx + kaleido for chart images)
# ══════════════════════════════════════════════════════════════════════════════
def build_word_report(dfs):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import plotly.express as px
        import plotly.graph_objects as go
    except ImportError as e:
        raise RuntimeError(f"Missing library: {e}. Run: pip install python-docx kaleido")

    # ── Colour palette ────────────────────────────────────────────────────────
    C_NAVY   = RGBColor(0x1A, 0x3A, 0x6B)
    C_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
    C_ACCENT = RGBColor(0x3B, 0x82, 0xF6)
    C_GREEN  = RGBColor(0x10, 0xB9, 0x81)
    C_AMBER  = RGBColor(0xF5, 0x9E, 0x0B)
    C_RED    = RGBColor(0xEF, 0x44, 0x44)
    C_GRAY   = RGBColor(0x64, 0x74, 0x8B)
    C_LGRAY  = RGBColor(0xF1, 0xF5, 0xF9)
    C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    CHART_COLORS = ["#3B82F6","#10B981","#F59E0B","#EF4444","#8B5CF6","#14B8A6","#EC4899","#F97316"]
    PLOT_BG = "rgba(0,0,0,0)"

    doc  = Document()
    now  = datetime.now().strftime("%B %d, %Y")
    buf  = io.BytesIO()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── Helper: set cell shading ──────────────────────────────────────────────
    def shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color.replace("#","").upper())
        tcPr.append(shd)

    # ── Helper: paragraph style ───────────────────────────────────────────────
    def para(text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6, italic=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color or C_GRAY
        return p

    def heading(text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = color or C_NAVY
            run.font.bold = True
        p.paragraph_format.space_before = Pt(10 if level > 1 else 18)
        p.paragraph_format.space_after  = Pt(4)
        return p

    def hr_line():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(8)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"),  "6")
        bottom.set(qn("w:space"),"1")
        bottom.set(qn("w:color"),"2E75B6")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def body_text(text, space_after=6):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        return p

    def caption(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(10)
        for run in p.runs:
            run.font.size = Pt(9)
            run.italic = True
            run.font.color.rgb = C_GRAY

    # ── Helper: data table ────────────────────────────────────────────────────
    def add_table(headers, rows, col_widths=None):
        n_cols = len(headers)
        t = doc.add_table(rows=1 + len(rows), cols=n_cols)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        page_w = section.page_width - section.left_margin - section.right_margin
        default_w = int(page_w / n_cols)
        widths = col_widths or [default_w]*n_cols

        # Header row
        hrow = t.rows[0]
        for i, (hdr, w) in enumerate(zip(headers, widths)):
            cell = hrow.cells[i]
            cell.width = w
            shade_cell(cell, "1A3A6B")
            run = cell.paragraphs[0].add_run(hdr)
            run.font.bold  = True
            run.font.size  = Pt(9)
            run.font.color.rgb = C_WHITE
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Data rows
        for r_idx, row_data in enumerate(rows):
            row_obj = t.rows[r_idx + 1]
            bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                cell = row_obj.cells[c_idx]
                cell.width = widths[c_idx]
                shade_cell(cell, bg)
                run = cell.paragraphs[0].add_run(str(val))
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1E,0x29,0x3B)
                cell.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        doc.add_paragraph()

    # ── Helper: render Plotly fig → bytes → inline image ─────────────────────
    def add_chart(fig, width_inches=6.0, height_px=320, caption_text=""):
        fig.update_layout(
            paper_bgcolor="white", plot_bgcolor="white",
            font=dict(family="Arial", color="#334155", size=11),
            margin=dict(l=40,r=40,t=40,b=40),
            title_font=dict(size=13, color="#1A3A6B", family="Arial"),
            legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(color="#334155",size=10),
                        orientation="h", yanchor="bottom", y=1.02, xanchor="left", x=0),
            colorway=CHART_COLORS,
        )
        fig.update_xaxes(showgrid=True, gridcolor="#E2E8F0", zeroline=False,
                         linecolor="#CBD5E1", tickfont=dict(size=10,color="#64748B"))
        fig.update_yaxes(showgrid=True, gridcolor="#E2E8F0", zeroline=False,
                         tickfont=dict(size=10,color="#64748B"))
        try:
            img_bytes = fig.to_image(format="png", width=int(width_inches*120),
                                     height=height_px, scale=2)
            img_buf = io.BytesIO(img_bytes)
            doc.add_picture(img_buf, width=Inches(width_inches))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption_text:
                caption(caption_text)
        except Exception as e:
            body_text(f"[Chart not available: {e}]")

    # ── Helper: KPI summary table ─────────────────────────────────────────────
    def add_kpi_row(items):
        """items = list of (label, value, sub)"""
        t = doc.add_table(rows=3, cols=len(items))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        page_w = section.page_width - section.left_margin - section.right_margin
        col_w  = int(page_w / len(items))
        for c_idx, (label, value, sub) in enumerate(items):
            shade_cell(t.rows[0].cells[c_idx], "EFF6FF")
            shade_cell(t.rows[1].cells[c_idx], "EFF6FF")
            shade_cell(t.rows[2].cells[c_idx], "EFF6FF")
            for row in t.rows:
                row.cells[c_idx].width = col_w
            # Label
            r0 = t.rows[0].cells[c_idx].paragraphs[0].add_run(label.upper())
            r0.font.size = Pt(7.5); r0.font.color.rgb = C_GRAY; r0.font.bold = True
            t.rows[0].cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Value
            r1 = t.rows[1].cells[c_idx].paragraphs[0].add_run(value)
            r1.font.size = Pt(18); r1.font.color.rgb = C_NAVY; r1.font.bold = True
            t.rows[1].cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Sub
            r2 = t.rows[2].cells[c_idx].paragraphs[0].add_run(sub)
            r2.font.size = Pt(8.5); r2.font.color.rgb = C_GRAY
            t.rows[2].cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # ═════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_org.add_run("INTERSOS")
    r.font.size = Pt(14); r.font.color.rgb = C_ACCENT; r.font.bold = True

    doc.add_paragraph()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("Mission Tchad")
    r.font.size = Pt(32); r.font.color.rgb = C_NAVY; r.font.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run("Rapport de Programme — Gestion de l'Information")
    r.font.size = Pt(16); r.font.color.rgb = C_GRAY

    hr_line()

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_date.add_run(f"Generated: {now}     |     Reporting Period: January 2025 – April 2026")
    r.font.size = Pt(10); r.font.color.rgb = C_GRAY

    doc.add_paragraph()
    doc.add_paragraph()
    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_conf.add_run("⚠  CONFIDENTIAL — For internal use only")
    r.font.size = Pt(10); r.font.color.rgb = C_RED; r.font.bold = True

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # DATA
    # ═════════════════════════════════════════════════════════════════════════
    df_b = dfs.get("Beneficiary_Registration", pd.DataFrame())
    df_w = dfs.get("WASH_Activities",           pd.DataFrame())
    df_f = dfs.get("Protection_Monitoring",          pd.DataFrame())
    df_c = dfs.get("Health_Activities",        pd.DataFrame())
    df_i = dfs.get("Indicator_Tracker",         pd.DataFrame())

    tot   = len(df_b)
    act   = slen(df_b,"Statut_Enregistrement","Active")
    wr    = int(ssum(dfs.get("Protection_Monitoring",pd.DataFrame()),"Beneficiaires_Atteints"))
    fhh   = int(ssum(dfs.get("Health_Activities",pd.DataFrame()),"Consultations_Total"))
    paid  = sfilt(df_c,"Statut","Paid")
    usd   = paid["Admissions"].sum() if has(paid,"Admissions") and len(paid)>0 else 0
    on_t  = slen(df_i,"Status","En bonne voie")
    at_r  = slen(df_i,"Status","À risque")
    off_t = slen(df_i,"Status","Hors trajectoire")
    ti    = len(df_i)

    # ═════════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ═════════════════════════════════════════════════════════════════════════
    heading("1. Executive Summary", 1)
    hr_line()
    body_text(
        f"This report presents the consolidated program performance of INTERSOS (SI) Mission Tchad "
        f"for the period January 2025 to April 2026. The mission operates across {suniq(df_b,'State')} states of Tchad, "
        f"providing multi-sector humanitarian assistance in WASH, Food Security & Livelihoods (FSL), Shelter & NFI, "
        f"and Cash & Voucher Assistance (CVA).", space_after=8)
    body_text(
        f"As of the reporting date, {tot:,} beneficiaries have been registered ({act:,} active). "
        f"WASH activities reached {wr:,} individuals; FSL distributions covered {fhh:,} households; "
        f"cash transfers totalling ${usd:,.0f} USD were disbursed.", space_after=12)

    add_kpi_row([
        ("Total Beneficiaries", f"{tot:,}", f"{act:,} active ({act/tot:.0%})" if tot else ""),
        ("WASH Individuals", f"{wr:,}", "cumulative reached"),
        ("FSL Households", f"{fhh:,}", "all distributions"),
        ("Cash Disbursed", f"${usd:,.0f}", "USD paid out"),
        ("States Covered", str(suniq(df_b,"Province")), "operational areas"),
    ])

    if ti:
        body_text(
            f"Of {ti} tracked indicators: {on_t} On Track ({on_t/ti:.0%}), "
            f"{at_r} At Risk ({at_r/ti:.0%}), {off_t} Off Track ({off_t/ti:.0%}).", space_after=4)

    # Indicator status chart
    if ti > 0:
        fig = go.Figure(go.Pie(
            labels=["En bonne voie","At Risk","Off Track"],
            values=[on_t, at_r, off_t], hole=0.55,
            marker=dict(colors=["#10B981","#F59E0B","#EF4444"],
                        line=dict(color="white",width=2)),
            textinfo="percent+label", textfont_size=11,
        ))
        fig.update_layout(title="Overall Indicator Status", showlegend=False)
        add_chart(fig, width_inches=3.5, height_px=260,
                  caption_text="Figure 1 — Overall indicator performance status")

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 2. BENEFICIARY REGISTRATION
    # ═════════════════════════════════════════════════════════════════════════
    heading("2. Enregistrement des Bénéficiaires", 1)
    hr_line()
    if not df_b.empty:
        heading("2.1 Overview", 2)
        body_text(
            f"The beneficiary registry contains {tot:,} records across "
            f"{suniq(df_b,'State')} states and {suniq(df_b,'Locality')} localities.")

        # Table: by sector
        if has(df_b,"Sector"):
            sc = df_b["Sector"].value_counts().reset_index(); sc.columns=["Sector","Count"]
            sc["% of Total"]=(sc["Count"]/tot*100).map(lambda x:f"{x:.1f}%")
            heading("Table 1 — Beneficiaries by Sector",3)
            add_table(["Sector","Count","% of Total"],
                      [[r["Sector"],f"{r['Count']:,}",r["% of Total"]] for _,r in sc.iterrows()],
                      col_widths=[Cm(8),Cm(3.5),Cm(3.5)])
            # Pie chart
            fig = px.pie(sc, values="Count", names="Sector", hole=0.50,
                         color="Sector",
                         color_discrete_map={"WASH":"#3B82F6","FSL":"#10B981",
                                             "Shelter & NFI":"#F59E0B","Cash & Voucher Assistance":"#EF4444"},
                         title="Beneficiaries by Sector")
            fig.update_traces(textinfo="percent+label", textfont_size=11,
                              marker=dict(line=dict(color="white",width=2)))
            add_chart(fig, width_inches=4.5, height_px=280,
                      caption_text="Figure 2 — Beneficiary distribution by sector")

        # Table: displacement
        if has(df_b,"Displacement_Status"):
            disp=df_b["Displacement_Status"].value_counts().reset_index(); disp.columns=["Status","Count"]
            disp["% of Total"]=(disp["Count"]/tot*100).map(lambda x:f"{x:.1f}%")
            heading("Table 2 — Displacement Status",3)
            add_table(["Displacement Status","Count","% of Total"],
                      [[r["Status"],f"{r['Count']:,}",r["% of Total"]] for _,r in disp.iterrows()],
                      col_widths=[Cm(7),Cm(3.5),Cm(3.5)])
            fig = px.bar(disp, x="Count", y="Status", orientation="h",
                         color="Count", color_continuous_scale=["#1E3A6E","#3B82F6"],
                         title="Beneficiaries by Displacement Status")
            fig.update_layout(coloraxis_showscale=False)
            add_chart(fig, width_inches=5.5, height_px=260,
                      caption_text="Figure 3 — Displacement status breakdown")

        heading("2.2 Gender & Vulnerability", 2)
        if has(df_b,"Sex"):
            fem  = slen(df_b,"Sex","Female")
            male = tot - fem
            body_text(f"Female beneficiaries: {fem:,} ({fem/tot:.0%}).  Male: {male:,} ({male/tot:.0%}).")
            if has(df_b,"Vulnerability_Level"):
                vd = df_b.groupby(["Vulnerability_Level","Sex"]).size().reset_index(name="n")
                fig = px.bar(vd, x="Vulnerability_Level", y="n", color="Sex", barmode="group",
                             color_discrete_map={"Female":"#EC4899","Male":"#3B82F6"},
                             title="Vulnerability Level × Sex")
                add_chart(fig, width_inches=5.5, height_px=270,
                          caption_text="Figure 4 — Vulnerability level by sex")

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 2B. GEOGRAPHIC MAPS  (Folium → PNG via selenium/screenshot fallback)
    # We render static Plotly maps as proxy since Folium needs a browser
    # ═════════════════════════════════════════════════════════════════════════
    heading("2B. Geographic Coverage — Tchad State Maps", 1)
    hr_line()
    body_text(
        "The following maps illustrate the geographic distribution of program beneficiaries "
        "across Tchad's states, using official administrative boundaries. "
        "Map 1 shows beneficiary counts per state; Map 2 shows sector composition; "
        "Map 3 shows vulnerability concentration.", space_after=10)

    # Map proxy 1: Choropleth-style bar chart (state counts)
    if not df_b.empty and has(df_b,"State"):
        counts = df_b.groupby("Province").size().reset_index(name="Count").sort_values("Count", ascending=True)
        fig = px.bar(counts, x="Count", y="State", orientation="h",
                     color="Count", color_continuous_scale=["#FFEAEA","#C1121F"],
                     title="Map 1 — Beneficiary Count by Tchad State")
        fig.update_layout(coloraxis_showscale=True,
                          coloraxis_colorbar=dict(title="Beneficiaries", tickfont=dict(size=9)))
        add_chart(fig, width_inches=6.2, height_px=380,
                  caption_text="Figure M1 — Beneficiary distribution across Tchad states (shapefile reference)")

    # Map proxy 2: Grouped bar — Sector × State
    if not df_b.empty and has(df_b,"State") and has(df_b,"Sector"):
        cs = df_b.groupby(["Province","Sector"]).size().reset_index(name="Count")
        top_states = df_b.groupby("Province").size().nlargest(8).index.tolist()
        cs = cs[cs["State"].isin(top_states)]
        fig = px.bar(cs, x="Count", y="State", color="Sector", orientation="h",
                     barmode="stack",
                     color_discrete_map={"WASH":"#3B82F6","FSL":"#10B981",
                                         "Shelter & NFI":"#F59E0B",
                                         "Cash & Voucher Assistance":"#C1121F"},
                     title="Map 2 — Sector Coverage by State (Top 8 states)")
        add_chart(fig, width_inches=6.2, height_px=360,
                  caption_text="Figure M2 — Sector distribution per state (shapefile reference)")

    # Map proxy 3: Vulnerability treemap
    if not df_b.empty and has(df_b,"State") and has(df_b,"Vulnerability_Level"):
        vt = df_b.groupby(["Province","Vulnerability_Level"]).size().reset_index(name="Count")
        fig = px.treemap(vt, path=["State","Vulnerability_Level"], values="Count",
                         color="Vulnerability_Level",
                         color_discrete_map={"Extremely Vulnerable":"#EF4444",
                                             "Vulnerable":"#F59E0B",
                                             "Moderately Vulnerable":"#10B981"},
                         title="Map 3 — Vulnerability Level by State")
        add_chart(fig, width_inches=6.2, height_px=380,
                  caption_text="Figure M3 — Vulnerability concentration across Tchad states")

    # Map proxy 4: Scatter geo plot using Plotly
    if not df_b.empty and has(df_b,"State"):
        # Build state-level summary for scatter
        state_summary = df_b.groupby("Province").agg(
            Count=("State","count"),
            **({} if not has(df_b,"Sector") else {})
        ).reset_index()
        # Add centroids
        state_summary["Lat"] = state_summary["State"].map(
            lambda s: STATE_CENTROIDS.get(s, (14.5, 29.5))[0])
        state_summary["Lon"] = state_summary["State"].map(
            lambda s: STATE_CENTROIDS.get(s, (14.5, 29.5))[1])
        fig = px.scatter_geo(
            state_summary, lat="Lat", lon="Lon", size="Count",
            text="State", color="Count",
            color_continuous_scale=["#FFEAEA","#C1121F"],
            size_max=45,
            scope="africa",
            center={"lat":15.0,"lon":29.5},
            title="Map 4 — Geographic Bubble Map: Beneficiaries by State",
            projection="natural earth",
        )
        fig.update_geos(
            fitbounds="locations", resolution=50,
            showcoastlines=True, coastlinecolor="#CBD5E1",
            showland=True, landcolor="#F8FAFC",
            showocean=True, oceancolor="#EFF6FF",
            showframe=True, framecolor="#94A3B8", framewidth=1,
            showcountries=True, countrycolor="#94A3B8",
            showsubunits=True, subunitcolor="#C1121F", subunitwidth=1.5,
        )
        fig.update_traces(textposition="top center", textfont=dict(size=8))
        fig.update_layout(coloraxis_showscale=False, showlegend=False)
        add_chart(fig, width_inches=6.2, height_px=420,
                  caption_text="Figure M4 — Proportional bubble map: beneficiary volume per state")

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 3. WASH
    # ═════════════════════════════════════════════════════════════════════════
    heading("3. EHA/WASH — Eau, Assainissement, Hygiène", 1)
    hr_line()
    if not df_w.empty:

        reached = int(ssum(dfs.get("Protection_Monitoring",pd.DataFrame()),"Beneficiaires_Atteints"))
        target  = int(ssum(df_w,"Target_Beneficiaries")) or 1
        pct     = reached/target
        wv  = int(ssum(df_w,"Water_Volume_Liters"))
        lat = int(ssum(df_w,"Latrine_Units_Built"))
        hyg = int(ssum(df_w,"Hygiene_Kits_Dist"))
        nfi = int(ssum(df_w,"NFI_Kits_Dist"))

        body_text(
            f"WASH activities reached {reached:,} individuals against a target of {target:,} "
            f"(coverage: {pct:.0%}). Water distributed: {wv:,} litres. "
            f"Latrines constructed: {lat:,}. Hygiene kits: {hyg:,}. NFI kits: {nfi:,}.")
        add_kpi_row([
            ("Individuals Reached", f"{reached:,}", f"{pct:.0%} of target"),
            ("Water", f"{wv:,} L", "distributed"),
            ("Latrines", f"{lat:,}", "constructed"),
            ("Hygiene Kits", f"{hyg:,}", "distributed"),
        ])

        # Activity performance chart
        if has(df_w,"Activity_Type"):
            grp = df_w.groupby("Activity_Type")[["Target_Beneficiaries","Beneficiaires_Atteints"]].sum().reset_index()
            fig = go.Figure()
            fig.add_bar(name="Target", x=grp["Activity_Type"], y=grp["Target_Beneficiaries"],
                        marker_color="rgba(59,130,246,0.25)", marker_line_color="#3B82F6", marker_line_width=1.5)
            fig.add_bar(name="Reached",x=grp["Activity_Type"], y=grp["Beneficiaires_Atteints"],
                        marker_color="#3B82F6")
            fig.update_layout(barmode="group", title="Target vs Reached by Activity Type")
            add_chart(fig, width_inches=6.0, height_px=300,
                      caption_text="Figure 5 — WASH target vs reached by activity")

            heading("Table 3 — WASH Performance by Activity",3)
            grp["Coverage%"] = grp.apply(lambda r: f"{r['Reached_Beneficiaries']/r['Target_Beneficiaries']:.0%}"
                                          if r["Target_Beneficiaries"]>0 else "N/A", axis=1)
            add_table(["Activity Type","Target","Reached","Coverage"],
                      [[r["Activity_Type"],f"{int(r['Target_Beneficiaries']):,}",
                        f"{int(r['Reached_Beneficiaries']):,}",r["Coverage%"]] for _,r in grp.iterrows()],
                      col_widths=[Cm(6.5),Cm(3),Cm(3),Cm(2.5)])

        # Functionality pie
        if has(df_w,"Functionality_Status"):
            func = df_w["Functionality_Status"].value_counts().reset_index(); func.columns=["Status","Count"]
            cmap = {"Fully Functional":"#10B981","Partially Functional":"#F59E0B",
                    "Non-Functional":"#EF4444","Under Construction":"#8B5CF6"}
            fig = px.pie(func, values="Count", names="Status", hole=0.50,
                         color="Status", color_discrete_map=cmap, title="Infrastructure Functionality")
            fig.update_traces(textinfo="percent+label", marker=dict(line=dict(color="white",width=2)))
            add_chart(fig, width_inches=4.5, height_px=270,
                      caption_text="Figure 6 — WASH infrastructure functionality status")

        # Sex disaggregation
        if has(df_w,"Reached_Female") and has(df_w,"Reached_Male"):
            tf = df_w["Reached_Female"].sum(); tm = df_w["Reached_Male"].sum()
            body_text(f"Gender breakdown — Female: {tf:,.0f} ({tf/(tf+tm):.0%}), Male: {tm:,.0f} ({tm/(tf+tm):.0%}).")
            if has(df_w,"State"):
                sg = df_w.groupby("Province")[["Reached_Female","Reached_Male"]].sum().reset_index()
                fig = px.bar(sg, x="State", y=["Reached_Female","Reached_Male"], barmode="stack",
                             color_discrete_map={"Reached_Female":"#EC4899","Reached_Male":"#3B82F6"},
                             title="WASH — Sex Disaggregation by State")
                add_chart(fig, width_inches=6.0, height_px=280,
                          caption_text="Figure 7 — WASH sex disaggregation by state")
    else:
        body_text("No WASH data available.")

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 4. FSL
    # ═════════════════════════════════════════════════════════════════════════
    heading("4. Protection — VBG, ENA, PSS", 1)
    hr_line()
    if not df_f.empty:
        hht = int(ssum(df_f,"Cible_Beneficiaires"))
        hhr = int(ssum(dfs.get("Health_Activities",pd.DataFrame()),"Consultations_Total"))
        qp  = ssum(df_f,"Quantity_Planned")
        qd  = ssum(df_f,"Quantity_Distributed")
        fem_fsl = int(ssum(df_f,"Female_HHH_Reached"))
        pb  = slen(df_f,"Type_Programme","Pipeline break")
        cov = hhr/hht if hht else 0

        body_text(
            f"FSL distributions reached {hhr:,} households of a target of {hht:,} "
            f"(coverage: {cov:.0%}). Female-headed households: {fem_fsl:,} ({fem_fsl/hhr:.0%} of reached). "
            f"Quantity distributed: {qd:,.0f} units of {qp:,.0f} planned.")
        if pb:
            body_text(f"⚠  Supply Chain Alert: {pb} pipeline breaks recorded during the reporting period.")
        add_kpi_row([
            ("HH Reached", f"{hhr:,}", f"{cov:.0%} of target"),
            ("Qty Distributed", f"{qd:,.0f}", "units"),
            ("Female HHH", f"{fem_fsl:,}", f"{fem_fsl/hhr:.0%}" if hhr else ""),
            ("Pipeline Breaks", str(pb), "alerts"),
        ])

        # Commodity chart
        if has(df_f,"Commodity_Type"):
            comm = df_f.groupby("Commodity_Type")["Beneficiaires_Atteints"].sum().reset_index()
            comm = comm.sort_values("Beneficiaires_Atteints",ascending=True).tail(10); comm.columns=["Commodity","Beneficiaires_Atteints"]
            fig = px.bar(comm, x="Beneficiaires_Atteints", y="Commodity", orientation="h",
                         color="Beneficiaires_Atteints", color_continuous_scale=["#14532D","#34D399"],
                         title="HH Reached by Commodity Type")
            fig.update_layout(coloraxis_showscale=False)
            add_chart(fig, width_inches=6.0, height_px=320,
                      caption_text="Figure 8 — Households reached by commodity type")
            heading("Table 4 — FSL by Commodity Type",3)
            add_table(["Commodity","HH Reached"],
                      [[r["Commodity"],f"{int(r['HH_Reached']):,}"] for _,r in comm.iterrows()],
                      col_widths=[Cm(10),Cm(5)])

        # Pipeline status
        if has(df_f,"Type_Programme"):
            pipe = df_f["Type_Programme"].value_counts().reset_index(); pipe.columns=["Status","Count"]
            cmap = {"In stock":"#10B981","Low stock":"#F59E0B",
                    "Pipeline break":"#EF4444","Awaiting delivery":"#8B5CF6"}
            fig = px.pie(pipe, values="Count", names="Status", hole=0.50,
                         color="Status", color_discrete_map=cmap, title="Pipeline Status")
            fig.update_traces(textinfo="percent+label", marker=dict(line=dict(color="white",width=2)))
            add_chart(fig, width_inches=4.0, height_px=260,
                      caption_text="Figure 9 — FSL supply chain pipeline status")

        # Donor chart
        if has(df_f,"Donor"):
            don = df_f.groupby("Donor")["Beneficiaires_Atteints"].sum().reset_index()
            fig = px.bar(don, x="Donor", y="Beneficiaires_Atteints", color="Donor",
                         color_discrete_sequence=CHART_COLORS, title="HH Reached by Donor",
                         text="Beneficiaires_Atteints")
            fig.update_traces(texttemplate="%{text:,.0f}", textposition="outside")
            fig.update_layout(showlegend=False)
            add_chart(fig, width_inches=5.5, height_px=270,
                      caption_text="Figure 10 — HH reached by donor")

        # PDM satisfaction
        if has(df_f,"Beneficiary_Satisfaction"):
            sat = df_f["Beneficiary_Satisfaction"].value_counts().reset_index(); sat.columns=["Level","Count"]
            cmap = {"Above 80%":"#10B981","60–80%":"#F59E0B","Below 60%":"#EF4444","N/A":"#94A3B8"}
            fig = px.bar(sat, x="Level", y="Count", color="Level",
                         color_discrete_map=cmap, title="Beneficiary Satisfaction (PDM)", text="Count")
            fig.update_traces(textposition="outside"); fig.update_layout(showlegend=False)
            add_chart(fig, width_inches=5.0, height_px=260,
                      caption_text="Figure 11 — Post-Distribution Monitoring: satisfaction")
    else:
        body_text("No FSL data available.")

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 5. CVA
    # ═════════════════════════════════════════════════════════════════════════
    heading("5. Santé — Consultations & Activités Médicales", 1)
    hr_line()
    if not df_c.empty:
        paid_df = sfilt(df_c,"Statut","Paid")
        pend_df = sfilt(df_c,"Statut","Pending")
        fail_df = sfilt(df_c,"Statut","Failed")
        usd2    = paid_df["Admissions"].sum() if has(paid_df,"Admissions") and len(paid_df)>0 else 0
        avg_v   = paid_df["Admissions"].mean() if has(paid_df,"Admissions") and len(paid_df)>0 else 0
        fem_cva = slen(df_c,"Female_Headed_HH","Yes")

        body_text(
            f"{len(paid_df):,} transfers paid totalling ${usd2:,.0f} USD (avg ${avg_v:,.0f}/HH). "
            f"Pending: {len(pend_df):,}. Failed: {len(fail_df):,}. "
            f"Female-headed HH: {fem_cva:,} ({fem_cva/len(df_c):.0%}).")
        if len(fail_df) > 0:
            body_text(f"⚠  {len(fail_df)} failed transfers require investigation before the next payment round.")
        add_kpi_row([
            ("Total Paid", f"${usd2:,.0f}", "USD"),
            ("Paid Transfers", f"{len(paid_df):,}", "completed"),
            ("Pending", f"{len(pend_df):,}", "awaiting"),
            ("Failed", f"{len(fail_df):,}", "to check"),
        ])

        # Status donut
        if has(df_c,"Statut"):
            sc = df_c["Statut"].value_counts().reset_index(); sc.columns=["Status","Count"]
            cmap={"Paid":"#10B981","Pending":"#F59E0B","Failed":"#EF4444","Cancelled":"#8B5CF6"}
            fig = px.pie(sc, values="Count", names="Status", hole=0.55,
                         color="Status", color_discrete_map=cmap, title="Transfer Status")
            fig.update_traces(textinfo="percent+label", marker=dict(line=dict(color="white",width=2)))
            add_chart(fig, width_inches=4.0, height_px=260,
                      caption_text="Figure 12 — CVA transfer status breakdown")

        # Payment method bar
        if has(df_c,"Payment_Method") and has(df_c,"Admissions"):
            pm = df_c.groupby("Payment_Method")["Admissions"].sum().reset_index()
            pm = pm.sort_values("Admissions",ascending=True)
            fig = px.bar(pm, x="Admissions", y="Payment_Method", orientation="h",
                         color="Admissions", color_continuous_scale=["#1E3A6E","#60A5FA"],
                         title="USD Disbursed by Payment Method", text="Admissions")
            fig.update_traces(texttemplate="$%{text:,.0f}", textposition="outside")
            fig.update_layout(coloraxis_showscale=False)
            add_chart(fig, width_inches=5.5, height_px=280,
                      caption_text="Figure 13 — USD disbursed by payment method")

        # Transfer timeline
        if has(df_c,"Transfer_Date") and has(df_c,"Admissions"):
            df_tl = df_c.copy()
            df_tl["Transfer_Date"] = pd.to_datetime(df_tl["Transfer_Date"],errors="coerce")
            df_tl = df_tl.dropna(subset=["Transfer_Date"])
            df_tl["Month"] = df_tl["Transfer_Date"].dt.to_period("M").astype(str)
            if has(df_tl,"Statut"):
                agg = df_tl.groupby(["Month","Statut"])["Admissions"].sum().reset_index()
                cmap={"Paid":"#10B981","Pending":"#F59E0B","Failed":"#EF4444","Cancelled":"#8B5CF6"}
                fig = px.bar(agg, x="Month", y="Admissions", color="Statut",
                             color_discrete_map=cmap, title="Monthly Transfer Volume (USD)")
                add_chart(fig, width_inches=6.0, height_px=280,
                          caption_text="Figure 14 — Monthly CVA transfer volume by status")

        heading("Table 5 — USD Disbursed by Payment Method",3)
        if has(df_c,"Payment_Method") and has(df_c,"Admissions"):
            pm2 = df_c.groupby("Payment_Method")["Admissions"].sum().reset_index()
            pm2 = pm2.sort_values("Admissions",ascending=False)
            add_table(["Payment Method","Total USD"],
                      [[r["Payment_Method"],f"${r['Transfer_Value_USD']:,.0f}"] for _,r in pm2.iterrows()],
                      col_widths=[Cm(9),Cm(6)])
    else:
        body_text("No CVA data available.")

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 6. PROGRAM INDICATORS
    # ═════════════════════════════════════════════════════════════════════════
    heading("6. Nutrition — CRENAS, CRENI, ANJE", 1)
    hr_line()
    if not df_i.empty:
        body_text(
            f"The results framework tracks {ti} indicators. "
            f"On Track: {on_t} ({on_t/ti:.0%})  At Risk: {at_r} ({at_r/ti:.0%})  "
            f"Off Track: {off_t} ({off_t/ti:.0%}).")

        # Status by sector chart
        if has(df_i,"Sector") and has(df_i,"Statut"):
            perf = df_i.groupby(["Sector","Status"]).size().reset_index(name="Count")
            cmap = {"En bonne voie":"#10B981","À risque":"#F59E0B","Hors trajectoire":"#EF4444"}
            fig = px.bar(perf, x="Sector", y="Count", color="Status",
                         color_discrete_map=cmap, barmode="stack",
                         title="Indicator Status by Sector")
            add_chart(fig, width_inches=6.0, height_px=280,
                      caption_text="Figure 15 — Indicator performance by sector")

        # Full indicator table
        heading("Table 6 — Full Indicator Tracking Table",3)
        ind_rows = []
        for _, row in df_i.iterrows():
            ind  = str(row.get("Indicateur",""))[:60]
            unit = str(row.get("Unit",""))
            targ = row.get("Cible_Annuelle",None)
            cum  = row.get("Cumul",None)
            stat = str(row.get("Status",""))
            try:
                t = float(str(targ).replace(",","").replace("%","")) if pd.notna(targ) and str(targ) not in ["","nan"] else None
                c = float(str(cum).replace(",","").replace("%",""))  if pd.notna(cum)  and str(cum)  not in ["","nan"] else 0
                ts = f"{t:,.0f}" if t else "N/A"
                cs = f"{c:,.0f}"
                ps = f"{c/t:.0%}" if t and t>0 else "N/A"
            except: ts = str(targ); cs = str(cum); ps = "N/A"
            ind_rows.append([ind, unit, ts, cs, ps, stat])
        add_table(
            ["Indicateur","Unit","Target","Achieved","% vs Target","Status"],
            ind_rows,
            col_widths=[Cm(5.8),Cm(1.5),Cm(1.8),Cm(1.8),Cm(2),Cm(2.1)]
        )
    else:
        body_text("No indicator data available.")

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # 7. RECOMMENDATIONS
    # ═════════════════════════════════════════════════════════════════════════
    heading("7. Indicateurs de Performance du Programme", 1)
    hr_line()
    recs = [
        ("Data Quality", "Enforce 48-hour data submission deadlines and implement automated daily quality checks to catch missing or inconsistent records before reporting."),
        ("WASH Coverage", f"Current coverage is {int(ssum(df_w,'Reached_Beneficiaries'))/max(int(ssum(df_w,'Target_Beneficiaries')),1):.0%}. Underperforming states require supervisory support and resource reallocation."),
        ("CVA Failed Transfers", f"{slen(df_c,'Transfer_Status','Failed')} transfers failed. Coordinate with payment agents to identify root causes and reprocess eligible cases before the next round."),
        ("FSL Pipeline", f"{slen(df_f,'Pipeline_Status','Pipeline break')} pipeline breaks recorded. Establish minimum buffer stock thresholds and increase monitoring frequency."),
        ("Off-Track Indicators", f"{slen(df_i,'Status','Off track')} indicators are off track. Program managers must review plans and propose corrective measures at the next Monthly Program Review."),
        ("MEAL Capacity", "Conduct a refresher training for field data collectors on KoBoToolbox protocols, GPS accuracy, and data quality standards."),
    ]
    for i,(title,text) in enumerate(recs,1):
        heading(f"7.{i} {title}", 2)
        body_text(text, space_after=8)

    doc.add_page_break()

    # ═════════════════════════════════════════════════════════════════════════
    # ANNEX
    # ═════════════════════════════════════════════════════════════════════════
    heading("Annex — Data Sources & Methodology", 1)
    hr_line()
    add_table(
        ["Dataset","Sheet Name","Records","Key Variables"],
        [
            ["Beneficiary Registry","Beneficiary_Registration",f"{len(df_b):,}","ID, State, Sector, Sex, Age, GPS"],
            ["WASH Activities","WASH_Activities",f"{len(df_w):,}","Activity, Target, Reached, Functionality"],
            ["FSL Distributions","Protection_Monitoring",f"{len(df_f):,}","Commodity, HH Reached, Pipeline, Donor"],
            ["CVA Transfers","Health_Activities",f"{len(df_c):,}","Type, Value USD, Status, Method"],
            ["Indicator Tracker","Indicator_Tracker",f"{len(df_i):,}","Indicator, Target, Q1–Q4, Status"],
        ],
        col_widths=[Cm(4),Cm(4),Cm(2.5),Cm(5)]
    )
    body_text(
        "All data collected via KoBoToolbox electronic forms and processed with Python/pandas. "
        "Geographic data captured via GPS-enabled Android devices. "
        "Indicator targets agreed with donors at project inception.", space_after=10)

    # Footer line
    hr_line()
    p_footer = doc.add_paragraph(
        f"INTERSOS — Mission Tchad IM Unit | Report generated {now} | CONFIDENTIAL")
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_footer.runs:
        run.font.size = Pt(8); run.font.color.rgb = C_GRAY; run.italic = True

    doc.save(buf)
    return buf.getvalue()


def page_report(dfs):
    ph("Générateur de Rapport", "Générer un rapport complet du programme — PDF, Word, or Text")
    if not dfs:
        st.info("Load the Excel database first to generate a report.")
        return

    th = TH()
    sh("Report Configuration")
    c1, c2, c3 = st.columns(3)
    with c1:
        report_type = st.selectbox("Report Format", [
            "📄 Word (.docx) — Full report with charts",
            "🖨️ PDF — Full report",
            "📝 Plain Text — Summary",
        ])
    with c2:
        st.markdown(f"<div style='padding:.5rem 0;font-size:.82rem;color:{th['muted']};'>Report language</div>", unsafe_allow_html=True)
        st.markdown(f"<div style='font-weight:600;color:{th['text']};'>English</div>", unsafe_allow_html=True)
    with c3:
        st.markdown(f"<div style='padding:.5rem 0;font-size:.82rem;color:{th['muted']};'>Generation date</div>", unsafe_allow_html=True)
        st.markdown(f"<div style='font-weight:600;color:{th['text']};'>{datetime.now().strftime('%B %d, %Y')}</div>", unsafe_allow_html=True)

    # Preview KPIs
    sh("Key Figures — Report Preview")
    df_b = dfs.get("Beneficiary_Registration", pd.DataFrame())
    df_w = dfs.get("WASH_Activities",           pd.DataFrame())
    df_f = dfs.get("Protection_Monitoring",          pd.DataFrame())
    df_c = dfs.get("Health_Activities",        pd.DataFrame())
    df_i = dfs.get("Indicator_Tracker",         pd.DataFrame())
    tot  = len(df_b)
    wr   = int(ssum(df_w, "Beneficiaires_Atteints"))
    fhh  = int(ssum(df_f, "Beneficiaires_Atteints"))
    paid = sfilt(df_c, "Statut", "Paid")
    usd  = paid["Admissions"].sum() if has(paid, "Admissions") and len(paid) > 0 else 0
    on_t = slen(df_i, "Status", "En bonne voie"); ti = len(df_i)

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.markdown(kpi("Beneficiaries", N(tot),   "registered",  "blue",  "👤"), unsafe_allow_html=True)
    c2.markdown(kpi("Bénéf. Atteints",  N(wr),    "individuals", "teal",  "💧"), unsafe_allow_html=True)
    c3.markdown(kpi("Cas Protection",        N(fhh),   "households",  "green", "🌾"), unsafe_allow_html=True)
    c4.markdown(kpi("Consultations",     f"${N(usd)}", "USD",     "amber", "💵"), unsafe_allow_html=True)
    c5.markdown(kpi("Indicators",    f"{on_t}/{ti}", "on track", "green", "📊"), unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    sh("Report Contents")
    st.markdown(f"""<div style='background:{th["panel"]};border:1px solid {th["border"]};border-radius:10px;padding:1.2rem 1.5rem;'>
      <div style='font-size:.85rem;color:{th["text"]};line-height:2.2;'>
        ✅ &nbsp;Cover page · Mission header · Date · Confidentiality notice<br>
        ✅ &nbsp;Executive summary — consolidated KPIs table<br>
        ✅ &nbsp;Section 2 — Beneficiary Registration (tables + <b>pie chart, bar charts</b>)<br>
        ✅ &nbsp;Section 3 — WASH Monitoring (performance table + <b>target vs reached chart, functionality pie</b>)<br>
        ✅ &nbsp;Section 4 — FSL Distribution (commodity table + <b>pipeline chart, coverage bar</b>)<br>
        ✅ &nbsp;Section 5 — CVA Cash Transfers (transfer table + <b>status donut, payment method bar</b>)<br>
        ✅ &nbsp;Section 6 — Program Indicators (full results framework + <b>status bar chart</b>)<br>
        ✅ &nbsp;Section 7 — Recommendations (data-driven)<br>
        ✅ &nbsp;Annex — Data sources & methodology
      </div>
    </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("🚀 Generate Report", use_container_width=False):
        fname_date = datetime.now().strftime("%Y%m%d")
        if "Word" in report_type:
            with st.spinner("Building Word report with charts… this may take 20–30 seconds."):
                try:
                    docx_bytes = build_word_report(dfs)
                    st.download_button(
                        "⬇ Download Word Report (.docx)", docx_bytes,
                        file_name=f"SI_Tchad_IM_Report_{fname_date}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    st.success("✅ Word report generated successfully!")
                except Exception as e:
                    st.error(f"Word generation error: {e}")
                    import traceback; st.code(traceback.format_exc())

        elif "PDF" in report_type:
            if HAS_PDF:
                with st.spinner("Building PDF report…"):
                    try:
                        pdf_bytes = build_pdf_report(dfs)
                        st.download_button(
                            "⬇ Download PDF Report", pdf_bytes,
                            file_name=f"SI_Tchad_IM_Report_{fname_date}.pdf",
                            mime="application/pdf",
                        )
                        st.success("✅ PDF report generated!")
                    except Exception as e:
                        st.error(f"PDF error: {e}")
            else:
                st.warning("ReportLab not installed. Run: pip install reportlab")
        else:
            txt = build_text_report(dfs)
            st.download_button(
                "⬇ Download Text Report", txt.encode(),
                file_name=f"SI_Tchad_IM_Report_{fname_date}.txt",
                mime="text/plain",
            )
            st.success("✅ Text report ready!")
            st.code(txt, language=None)



# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
def main():
    # 1. Not logged in → login page
    if not st.session_state.get("auth"):
        login_page()
        return

    # 2. Record login time (once)
    if "login_time" not in st.session_state:
        st.session_state["login_time"] = datetime.now().strftime("%d %b %Y %H:%M")

    # 3. Logged in but data source not yet selected → datasource page
    kobo_step = st.session_state.get("kobo_step", 0)
    if kobo_step != 99:
        inject_css(DARK)
        datasource_page()
        return

    # 4. Full dashboard
    inject_css(TH())
    top_nav()

    page = st.session_state.get("page", "Overview")
    dfs  = st.session_state.get("dfs", {})
    th   = TH()

    # 5. Render sidebar (profile + alerts + thresholds)
    render_sidebar(dfs)

    if not dfs and page != "Report":
        st.markdown(f"""<div style='text-align:center;padding:5rem 2rem;
          background:{th["panel"]};border:1px dashed {th["border2"]};
          border-radius:14px;margin-top:2rem;'>
          <div style='font-size:2.5rem;margin-bottom:1rem;'>📂</div>
          <div style='font-size:1.1rem;font-weight:700;color:{th["text"]};margin-bottom:.5rem;'>No data loaded</div>
          <div style='font-size:.85rem;color:{th["muted"]}; '>
            Go back to the data source page to load your data.
          </div>
        </div>""", unsafe_allow_html=True)
        if st.button("← Change data source"):
            st.session_state["kobo_step"] = 0
            st.session_state.pop("dfs", None); st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
        return

    # 6. Global alert banner (critical only, compact)
    if dfs:
        alerts = run_alert_engine(dfs)
        criticals = [a for a in alerts if a["level"] == "critical"]
        if criticals:
            with st.expander(
                f"🚨 {len(criticals)} Critical Alert{'s' if len(criticals)>1 else ''} — click to review",
                expanded=False
            ):
                for a in criticals:
                    render_alert(a)

    # 7. Dispatch to page
    dispatch = {
        "Overview":   page_overview,
        "Carte":      page_map,
        "Protection": page_protection,
        "Santé":      page_sante,
        "Nutrition":  page_nutrition,
        "EHA/WASH":   page_wash,
        "Indicateurs":page_ind,
        "Données":    page_raw,
        "Rapport":    page_report,
    }
    dispatch.get(page, page_overview)(dfs)
    st.markdown("</div>", unsafe_allow_html=True)

if __name__ == "__main__":
    main()